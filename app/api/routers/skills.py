"""Skills and prompt packs router — skill packages, prompt packs, skill store."""
from __future__ import annotations

import logging
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Query, Body

from ..deps.hub import get_hub
from ..deps.auth import CurrentUser, get_current_user

logger = logging.getLogger("tudouclaw.api.skills")

router = APIRouter(prefix="/api/portal", tags=["skills"])


@router.get("/tools")
async def list_tool_catalog(
    user: CurrentUser = Depends(get_current_user),
):
    """Return the full tool catalog as OpenAI function-calling schemas.

    Powers the agent Settings modal "工具权限" section so new tools
    automatically appear as checkboxes. Returns unfiltered catalog —
    per-agent allow/deny lists are applied at execution time.
    """
    try:
        from ... import tools as _tools_mod
        if hasattr(_tools_mod, "tool_registry"):
            defs = _tools_mod.tool_registry.get_definitions()
            return {"tools": defs, "count": len(defs)}
        return {"tools": [], "count": 0}
    except Exception as e:
        logger.exception("list_tool_catalog failed")
        raise HTTPException(500, str(e))


def _get_skill_or_404(hub, skill_id: str):
    """Get skill package or raise 404."""
    try:
        skill = hub.get_skill_package(skill_id) if hasattr(hub, "get_skill_package") else None
        if not skill:
            raise HTTPException(status_code=404, detail=f"Skill package '{skill_id}' not found")
        return skill
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Skill packages
# ---------------------------------------------------------------------------

@router.get("/skill-pkgs")
async def list_skill_packages(
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """List all installed skill packages."""
    try:
        skills = hub.list_skill_packages() if hasattr(hub, "list_skill_packages") else []
        skills_list = [s.to_dict() if hasattr(s, "to_dict") else s for s in skills]
        return {"skill_packages": skills_list}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/skill-pkgs/{skill_id}")
async def get_skill_package(
    skill_id: str,
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Get skill package detail."""
    try:
        skill = _get_skill_or_404(hub, skill_id)
        data = skill.to_dict() if hasattr(skill, "to_dict") else skill
        return data if isinstance(data, dict) else {"data": data}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/skill-pkgs/install")
async def install_skill_package(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Install a skill package."""
    try:
        skill_id = body.get("skill_id", "")
        if not skill_id:
            raise HTTPException(400, "Missing skill_id")

        if hasattr(hub, "install_skill_package"):
            result = hub.install_skill_package(skill_id, body)
            return {"ok": True, "result": result}

        return {"ok": True}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/skills/install-from-url-async")
async def install_skill_from_url_async(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Kick off URL-based skill install in a background thread.

    Returns ``{"install_id": "..."}`` immediately. The UI polls
    ``GET /api/portal/skills/install-progress/{install_id}`` for live
    progress + final result.

    This is the UX-friendly variant of ``/install-from-url`` (which is
    synchronous and blocks for 5-10s).
    """
    import threading
    import os
    from ...skills import install_progress as _ip
    from ...skills.url_installer import install_skill_from_url, InstallFromUrlError

    url = (body.get("url") or "").strip()
    if not url:
        raise HTTPException(400, "url is required")
    overwrite = bool(body.get("overwrite", False))

    skill_registry = getattr(hub, "skill_registry", None)
    if skill_registry is None:
        raise HTTPException(500, "skill_registry not initialized")
    data_dir = getattr(hub, "_data_dir", None) or os.path.expanduser("~/.tudou_claw")
    catalog_dir = os.path.join(data_dir, "skill_catalog")
    actor = getattr(user, "user_id", "") or "unknown"

    install_id = _ip.start(source_url=url)

    def _bg_install() -> None:
        try:
            result = install_skill_from_url(
                url, catalog_dir=catalog_dir,
                skill_registry=skill_registry,
                skill_store=getattr(hub, "skill_store", None),
                installed_by=actor, overwrite=overwrite,
                progress_id=install_id,
            )
            _ip.complete(install_id, success=True, result=result)
        except InstallFromUrlError as e:
            _ip.complete(install_id, success=False, error=str(e))
        except Exception as e:
            logger.exception("install_from_url background error: %s", e)
            _ip.complete(install_id, success=False,
                         error=f"unexpected: {e}")

    threading.Thread(
        target=_bg_install, daemon=True,
        name=f"skill-install-{install_id}",
    ).start()
    return {"install_id": install_id, "ok": True}


@router.get("/skills/install-progress/{install_id}")
async def get_install_progress(
    install_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    """Poll endpoint for ``install-from-url-async`` progress.

    Returns the latest snapshot. Once status is success/error, the entry
    sticks around for ~5min then GC'd.
    """
    from ...skills import install_progress as _ip
    st = _ip.get(install_id)
    if st is None:
        raise HTTPException(404, f"install_id {install_id!r} not found")
    return st.to_dict()


@router.post("/skills/install-from-url")
async def install_skill_from_url_endpoint(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Direct-install a Anthropic Agent Skills (SKILL.md) skill from URL.

    Body:
        url: ClawHub URL (https://clawhub.ai/<author>/<slug>),
             ``author/slug``, ``slug``, any ``.zip`` URL, or raw SKILL.md URL
        overwrite: bool (default False) — replace existing entry with same slug

    Returns the install result (skill_id, name, version, catalog_path).
    Skips the LLM-driven semantic conversion in skill-converter — use
    this when the source skill is already in standard Anthropic format.
    """
    from ...skills.url_installer import (
        install_skill_from_url, InstallFromUrlError,
    )
    import os

    url = (body.get("url") or "").strip()
    if not url:
        raise HTTPException(400, "url is required")
    overwrite = bool(body.get("overwrite", False))

    skill_registry = getattr(hub, "skill_registry", None)
    if skill_registry is None:
        raise HTTPException(500, "skill_registry not initialized on hub")

    # Catalog dir: prefer the user-level <data_dir>/skill_catalog (matches
    # what hub setup uses as the writable catalog dir).
    data_dir = getattr(hub, "_data_dir", None) or os.path.expanduser(
        "~/.tudou_claw"
    )
    catalog_dir = os.path.join(data_dir, "skill_catalog")

    actor = getattr(user, "user_id", "") or "unknown"
    try:
        result = install_skill_from_url(
            url,
            catalog_dir=catalog_dir,
            skill_registry=skill_registry,
            skill_store=getattr(hub, "skill_store", None),
            installed_by=actor,
            overwrite=overwrite,
        )
        return result
    except InstallFromUrlError as e:
        raise HTTPException(400, str(e))
    except Exception as e:
        logger.exception("install_from_url failed: %s", e)
        raise HTTPException(500, f"unexpected error: {e}")


@router.post("/skill-pkgs/{skill_id}/uninstall")
async def uninstall_skill_package(
    skill_id: str,
    body: dict = Body(default={}),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Uninstall a skill package."""
    try:
        reg = getattr(hub, "skill_registry", None)
        if not reg:
            raise HTTPException(503, "Skill registry unavailable")
        ok = reg.uninstall(skill_id)
        return {"ok": ok}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/skill-pkgs/{skill_id}/invoke")
async def invoke_skill_package(
    skill_id: str,
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Invoke a skill package.

    Parity with legacy portal_routes_post: skill-not-found and invalid
    inputs surface as 400 (user error), not 500. Only genuinely
    unexpected exceptions become 500.
    """
    agent_id = body.get("agent_id", "")
    inputs = body.get("inputs", {}) or {}
    reg = getattr(hub, "skill_registry", None)
    if not reg:
        raise HTTPException(503, "Skill registry unavailable")
    try:
        result = reg.invoke(skill_id, agent_id, inputs)
    except HTTPException:
        raise
    except Exception as e:
        # Legacy returned 400 for any reg.invoke() error — matches the
        # semantic "client asked for something invalid" better than 500.
        raise HTTPException(status_code=400, detail=str(e))
    return {"ok": True, "result": result}


@router.get("/skill-pkgs/{skill_id}/agents")
async def get_skill_agents(
    skill_id: str,
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """List agents that have this skill package granted."""
    try:
        reg = getattr(hub, "skill_registry", None)
        if not reg:
            return {"agents": []}
        inst = reg.get(skill_id)
        if not inst:
            raise HTTPException(404, "Skill not found")
        agent_ids = getattr(inst, "granted_agents", []) or []
        agents = []
        for aid in agent_ids:
            a = hub.get_agent(aid) if hasattr(hub, "get_agent") else None
            agents.append({
                "agent_id": aid,
                "agent_name": a.name if a else aid,
            })
        return {"agents": agents, "skill_id": skill_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Skill granting and revocation
# ---------------------------------------------------------------------------

@router.post("/skill-pkgs/{skill_id}/grant")
async def grant_skill_to_agents(
    skill_id: str,
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Grant skill package to agents."""
    try:
        skill = _get_skill_or_404(hub, skill_id)
        agent_ids = body.get("agent_ids", [])

        if hasattr(skill, "grant_to_agents"):
            result = skill.grant_to_agents(agent_ids)
            return {"ok": True, "result": result}

        return {"ok": True}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/skill-pkgs/{skill_id}/revoke")
async def revoke_skill_from_agents(
    skill_id: str,
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Revoke skill package from agents."""
    try:
        skill = _get_skill_or_404(hub, skill_id)
        agent_ids = body.get("agent_ids", [])

        if hasattr(skill, "revoke_from_agents"):
            result = skill.revoke_from_agents(agent_ids)
            return {"ok": True, "result": result}

        return {"ok": True}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Prompt packs
# ---------------------------------------------------------------------------

@router.get("/prompt-packs")
async def list_prompt_packs(
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """List all prompt packs — matches legacy portal_routes_get."""
    try:
        from ...core.prompt_enhancer import get_prompt_pack_registry
        registry = get_prompt_pack_registry()
        skills = [s.to_dict() for s in registry.store.get_active()]
        return {"skills": skills, "stats": registry.store.get_stats()}
    except (ImportError, Exception) as e:
        return {"skills": [], "stats": {}}


@router.post("/prompt-packs")
async def manage_prompt_packs(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Create, update, delete, or browse catalog of prompt packs."""
    try:
        action = body.get("action", "create")

        if action == "catalog":
            # -----------------------------------------------------------
            # Browse community_skills.json catalog (matches old server)
            # -----------------------------------------------------------
            import json as _json
            from pathlib import Path as _Path
            catalog_path = _Path(__file__).resolve().parent.parent.parent / "data" / "community_skills.json"
            try:
                with open(catalog_path, "r", encoding="utf-8") as f:
                    catalog = _json.load(f)
            except FileNotFoundError:
                return {"skills": [], "categories": [], "total": 0}

            category_filter = body.get("category", "")
            search_query = body.get("search", "").lower()
            page = body.get("page", 1)
            per_page = body.get("per_page", 20)

            skills = catalog.get("skills", [])
            if category_filter:
                skills = [s for s in skills if s.get("category") == category_filter]
            if search_query:
                skills = [
                    s for s in skills
                    if search_query in s.get("name", "").lower()
                    or search_query in s.get("description", "").lower()
                ]

            total = len(skills)
            start = (page - 1) * per_page
            paginated = skills[start : start + per_page]

            result = [
                {
                    "id": s.get("id", ""),
                    "name": s.get("name", ""),
                    "description": s.get("description", ""),
                    "icon": s.get("icon", ""),
                    "category": s.get("category", ""),
                }
                for s in paginated
            ]
            return {
                "skills": result,
                "categories": catalog.get("categories", []),
                "total": total,
                "page": page,
                "per_page": per_page,
            }

        elif action == "create":
            pack = hub.create_prompt_pack(body) if hasattr(hub, "create_prompt_pack") else {}
            return {"ok": True, "prompt_pack": pack}
        elif action == "update":
            pack = hub.update_prompt_pack(body.get("pack_id"), body) if hasattr(hub, "update_prompt_pack") else {}
            return {"ok": True, "prompt_pack": pack}
        elif action == "delete":
            hub.delete_prompt_pack(body.get("pack_id")) if hasattr(hub, "delete_prompt_pack") else None
            return {"ok": True}
        else:
            raise HTTPException(400, f"Unknown action: {action}")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Skill store
# ---------------------------------------------------------------------------

@router.get("/skill-store")
async def get_skill_store_catalog(
    source: str = Query("", description="Filter by source"),
    tag: str = Query("", description="Filter by tag"),
    q: str = Query("", description="Search query"),
    all: str = Query("0", description="Include disallowed"),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Get skill store catalog — matches legacy portal_routes_get."""
    try:
        store = getattr(hub, "skill_store", None)
        if store is None:
            return {"entries": [], "installed": {}, "annotations": [], "stats": {}}

        include_all = all in ("1", "true", "yes")
        entries = store.list_catalog(
            source_filter=source or "",
            tag=tag or "",
            query=q or "",
            include_disallowed=include_all,
        )

        # Pull installed-skill metadata for UI
        installed_map: dict = {}
        reg = getattr(hub, "skill_registry", None)
        if reg is not None:
            try:
                for inst in reg.list_all():
                    installed_map[inst.id] = {
                        "id": inst.id,
                        "name": inst.manifest.name,
                        "status": inst.status,
                        "granted_to": list(inst.granted_to),
                        "runtime": inst.manifest.runtime,
                    }
            except Exception:
                pass

        return {
            "entries": [e.to_dict() for e in entries],
            "installed": installed_map,
            "annotations": store.list_annotations(),
            "stats": store.stats(),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/skill-store")
async def manage_skill_store(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Unified skill store management endpoint.

    Mirrors the old BaseHTTPRequestHandler at /api/portal/skill-store.
    Supported actions: search, browse, rescan, install, uninstall,
    grant, revoke, annotate, clear_annotation, set_allowed_sources,
    import, import_bulk, import_from_url, scan_url, cleanup_scan.
    """
    store = getattr(hub, "skill_store", None)
    if store is None:
        raise HTTPException(503, "skill store not initialized")

    action = body.get("action", "search")

    try:
        if action == "search":
            if hasattr(hub, "search_skill_store"):
                return {"results": hub.search_skill_store(body.get("query", ""))}
            return {"results": []}

        if action == "browse":
            if hasattr(hub, "browse_skill_store"):
                return {"results": hub.browse_skill_store(body.get("category", ""))}
            return {"results": []}

        if action == "rescan":
            n = store.scan()
            return {"ok": True, "count": n, "stats": store.stats()}

        if action == "install":
            entry_id = body.get("entry_id", "")
            who = body.get("installed_by", "portal")
            result = store.install_entry(entry_id, installed_by=who)
            store.scan()  # refresh installed flags
            return {"ok": True, "result": result}

        if action == "uninstall":
            entry_id = body.get("entry_id", "")
            ok = store.uninstall_entry(entry_id)
            store.scan()
            return {"ok": ok}

        if action == "disable":
            # Soft-disable: skill stays on disk, hidden from default list.
            entry_id = body.get("entry_id", "")
            disabled = bool(body.get("disabled", True))
            ok = store.disable_entry(entry_id, disabled=disabled)
            store.scan()
            return {"ok": ok, "disabled": disabled}

        if action == "delete_catalog":
            # Hard-delete: removes the catalog directory + uninstalls if
            # currently installed. NOT recoverable — caller must confirm.
            entry_id = body.get("entry_id", "")
            ok = store.delete_catalog_entry(entry_id)
            store.scan()
            return {"ok": ok}

        if action == "grant":
            installed_id = body.get("installed_id", "")
            agent_id = body.get("agent_id", "")
            agent = hub.get_agent(agent_id) if hasattr(hub, "get_agent") else None
            wdir = ""
            sync_result = None
            if agent is not None:
                wdir = getattr(agent, "working_directory", "") or ""
                if installed_id not in getattr(agent, "granted_skills", []):
                    try:
                        agent.granted_skills.append(installed_id)
                        hub._save_agents()
                    except Exception as e:
                        logger.warning("grant skill %s to agent %s: %s", installed_id, agent_id, e)
                # Sync skill package to agent workspace
                try:
                    inst = store._registry.get(installed_id) if hasattr(store, "_registry") else None
                    if inst is not None and hasattr(agent, "sync_skill_to_workspace"):
                        sync_result = agent.sync_skill_to_workspace(inst)
                        hub._save_agents()
                except Exception as e:
                    logger.debug("sync skill %s workspace %s: %s", installed_id, agent_id, e)
            ok = store.grant(installed_id, agent_id, agent_working_dir=wdir)
            resp = {"ok": ok}
            if sync_result:
                resp["sync"] = sync_result
            return resp

        if action == "revoke":
            installed_id = body.get("installed_id", "")
            agent_id = body.get("agent_id", "")
            agent = hub.get_agent(agent_id) if hasattr(hub, "get_agent") else None
            wdir = getattr(agent, "working_directory", "") if agent else ""
            if agent is not None and installed_id in getattr(agent, "granted_skills", []):
                try:
                    agent.granted_skills.remove(installed_id)
                    hub._save_agents()
                except Exception as e:
                    logger.warning("revoke skill %s from agent %s: %s", installed_id, agent_id, e)
            # Remove skill from agent workspace
            if agent is not None and hasattr(agent, "remove_skill_from_workspace"):
                try:
                    inst = store._registry.get(installed_id) if hasattr(store, "_registry") else None
                    skill_name = ""
                    if inst:
                        skill_name = getattr(getattr(inst, "manifest", None), "name", "") or getattr(inst, "id", "")
                    if skill_name:
                        agent.remove_skill_from_workspace(skill_name)
                        hub._save_agents()
                except Exception:
                    pass
            ok = store.revoke(installed_id, agent_id, agent_working_dir=wdir or "")
            return {"ok": ok}

        if action == "annotate":
            skill_id = body.get("skill_id", "")
            text = body.get("text", "")
            author = body.get("author", "portal")
            ann = store.annotate(skill_id, text, author=author)
            return {"ok": True, "annotation": ann}

        if action == "clear_annotation":
            skill_id = body.get("skill_id", "")
            ok = store.clear_annotation(skill_id)
            return {"ok": ok}

        if action == "set_allowed_sources":
            sources = body.get("sources", []) or []
            store.set_allowed_sources(sources)
            return {"ok": True, "allowed": store.allowed_sources()}

        if action == "import":
            from ...skill_store import import_agent_skill as _import_one
            src_path = (body.get("src_path") or "").strip()
            tier = body.get("tier", "community")
            auto_install = body.get("auto_install", True)
            if not src_path:
                raise HTTPException(400, "src_path is required")
            catalog_dir = store.catalog_dirs[-1] if store.catalog_dirs else ""
            if not catalog_dir:
                raise HTTPException(500, "no catalog dir configured")
            result = _import_one(src_path, catalog_dir, source_tier=tier)
            if auto_install and result.get("ok"):
                store.scan()
                eid = result.get("entry_id", "")
                if eid:
                    try:
                        store.install_entry(eid, installed_by="portal")
                        store.scan()
                    except Exception:
                        pass
            return result

        if action == "import_bulk":
            from ...skill_store import import_anthropic_skills_bulk as _import_bulk
            src_root = (body.get("src_root") or body.get("directory") or "").strip()
            auto_install = body.get("auto_install", True)
            if not src_root:
                raise HTTPException(400, "src_root is required")
            catalog_dir = store.catalog_dirs[-1] if store.catalog_dirs else ""
            if not catalog_dir:
                raise HTTPException(500, "no catalog dir configured")
            results = _import_bulk(src_root, catalog_dir)
            if auto_install:
                store.scan()
                for r in (results if isinstance(results, list) else []):
                    eid = r.get("entry_id", "")
                    if eid and r.get("ok"):
                        try:
                            store.install_entry(eid, installed_by="portal")
                        except Exception:
                            pass
                store.scan()
            return {"ok": True, "results": results}

        if action == "import_from_url":
            url = (body.get("url") or "").strip()
            auto_install = body.get("auto_install", True)
            if not url:
                raise HTTPException(400, "url is required")
            catalog_dir = store.catalog_dirs[-1] if store.catalog_dirs else ""

            # ClawHub URL / bare slug / .zip URL → fast direct-install path
            # via url_installer (recognizes ClawHub's zip download API and
            # handles SKILL.md-only skills without manifest.yaml).
            from ...skills.url_installer import (
                install_skill_from_url, InstallFromUrlError,
                resolve_download_url,
            )
            try:
                # Cheap probe: does this URL look like something url_installer
                # can handle? It raises on unrecognized formats — for those
                # we fall through to the legacy scan_remote_url path which
                # handles GitHub repo clones, tar.gz, multi-skill archives.
                resolve_download_url(url)
                use_direct = True
            except InstallFromUrlError:
                use_direct = False

            if use_direct:
                try:
                    direct = install_skill_from_url(
                        url,
                        catalog_dir=catalog_dir,
                        skill_registry=getattr(hub, "skill_registry", None),
                        skill_store=store,
                        installed_by="portal",
                        overwrite=bool(body.get("overwrite", True)),
                    )
                    # Shape result like the legacy multi-skill format so the
                    # caller UI doesn't have to branch.
                    return {
                        "ok": True,
                        "results": [{
                            "ok": True,
                            "name": direct.get("name"),
                            "entry_id": direct.get("skill_id"),
                            "path": direct.get("catalog_path"),
                            "version": direct.get("version"),
                        }],
                        "install_path": "direct",
                    }
                except InstallFromUrlError as e:
                    # Direct path failed → try legacy as fallback (e.g.
                    # ClawHub URL might just be down)
                    logger.info("install_skill_from_url failed (%s); "
                                "falling back to scan_remote_url", e)

            # Legacy path: scan_remote_url for GitHub repos, tar.gz, etc.
            from ...skill_store import scan_remote_url, import_from_scan_result
            scan_data = scan_remote_url(url)
            if not isinstance(scan_data, dict) or not scan_data.get("ok"):
                err = (scan_data or {}).get("error") if isinstance(scan_data, dict) else "scan failed"
                raise HTTPException(400, f"scan failed: {err}")
            # import_from_scan_result(temp_dir, skill_names, catalog_dir, tier="community")
            temp_dir = scan_data.get("temp_dir", "")
            skill_names = [s.get("name", "") for s in (scan_data.get("skills") or []) if s.get("name")]
            results = import_from_scan_result(temp_dir, skill_names, catalog_dir)
            if auto_install:
                store.scan()
                for r in (results if isinstance(results, list) else []):
                    eid = r.get("entry_id", "")
                    if eid and r.get("ok"):
                        try:
                            store.install_entry(eid, installed_by="portal")
                        except Exception:
                            pass
                store.scan()
            return {"ok": True, "results": results, "install_path": "legacy"}

        if action == "scan_url":
            from ...skill_store import scan_remote_url
            url = (body.get("url") or "").strip()
            if not url:
                raise HTTPException(400, "url is required")
            data = scan_remote_url(url)
            return data

        if action == "cleanup_scan":
            from ...skill_store import cleanup_scan_temp
            temp_dir = body.get("temp_dir", "")
            cleanup_scan_temp(temp_dir)
            return {"ok": True}

        raise HTTPException(400, f"Unknown skill-store action: {action}")
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("skill-store action=%s error", action)
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Pending skills (SkillForge drafts awaiting admin approval)
# ---------------------------------------------------------------------------

@router.get("/pending-skills")
async def list_pending_skills(
    status: str = Query("", description="Filter by status: draft, exported, approved, rejected"),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """List skill drafts generated by SkillForge, pending admin review."""
    try:
        from ...skills._skill_forge import get_skill_forge
        forge = get_skill_forge()
        drafts = forge.list_drafts()

        if status:
            drafts = [d for d in drafts if d.status == status]

        return {
            "drafts": [d.to_dict() for d in drafts],
            "total": len(drafts),
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pending-skills/{draft_id}")
async def get_pending_skill(
    draft_id: str,
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Get detail of a single skill draft including code files and parsed manifest."""
    try:
        from ...skills._skill_forge import get_skill_forge
        forge = get_skill_forge()
        drafts = {d.id: d for d in forge.list_drafts()}
        draft = drafts.get(draft_id)
        if not draft:
            raise HTTPException(404, f"Draft not found: {draft_id}")
        result = draft.to_dict()
        # Add parsed manifest for easier rendering
        try:
            import yaml as _yaml
            if _yaml and draft.manifest_yaml:
                result["manifest_parsed"] = _yaml.safe_load(draft.manifest_yaml)
        except Exception:
            pass
        return result
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/pending-skills/{draft_id}/approve")
async def approve_pending_skill(
    draft_id: str,
    body: dict = Body(default={}),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Approve a skill draft and import it into the skill store.

    forge.approve_draft() auto-exports the package. This endpoint
    then rescans the skill store catalog to pick it up.
    """
    try:
        from ...skills._skill_forge import get_skill_forge
        forge = get_skill_forge()
        result = forge.approve_draft(draft_id)
        if "error" in result:
            raise HTTPException(404, result["error"])

        # Rescan skill store catalog to pick up the exported package
        import_result = {}
        export_dir = result.get("export_dir", "")
        try:
            store = getattr(hub, "skill_store", None)
            if store is not None:
                store.scan()  # rescan all catalog_dirs including pending_skills
                import_result = {"imported": True, "export_dir": export_dir}
            else:
                import_result = {"imported": False, "reason": "skill_store not available"}
        except Exception as ie:
            logger.warning("Skill store rescan after approval failed: %s", ie)
            import_result = {"imported": False, "error": str(ie)}

        return {
            "ok": True,
            "draft_id": draft_id,
            "status": "approved",
            "import": import_result,
        }
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/pending-skills/import")
async def import_skill_as_draft(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Import a skill directory (created by an agent) as a SkillForge draft.

    Body: { "path": "/abs/path/to/skill_dir" }
    or:   { "agent_id": "xxx", "dir_name": "pptx_skill" }
    """
    import os
    import time as _time
    import yaml as _yaml
    from ...skills._skill_forge import get_skill_forge, SkillDraft

    forge = get_skill_forge()

    # Resolve the skill directory path
    skill_dir = body.get("path", "").strip()
    if not skill_dir:
        agent_id = body.get("agent_id", "").strip()
        dir_name = body.get("dir_name", "").strip()
        if agent_id and dir_name:
            # Look in agent workspace
            agent = hub.get_agent(agent_id) if agent_id else None
            if agent and hasattr(agent, "workspace_dir"):
                skill_dir = os.path.join(str(agent.workspace_dir), dir_name)
            else:
                # Fallback: search common workspace paths
                from ... import DEFAULT_DATA_DIR
                candidates = [
                    os.path.join(DEFAULT_DATA_DIR, "workspaces", "agents", agent_id, "workspace", dir_name),
                ]
                for c in candidates:
                    if os.path.isdir(c):
                        skill_dir = c
                        break
        if not skill_dir:
            raise HTTPException(400, "Provide 'path' or 'agent_id'+'dir_name'")

    if not os.path.isdir(skill_dir):
        raise HTTPException(404, f"Directory not found: {skill_dir}")

    # Read skill files
    manifest_yaml = ""
    skill_md = ""
    code_files = {}
    name = os.path.basename(skill_dir)
    description = ""
    runtime = "markdown"
    triggers = []

    manifest_path = os.path.join(skill_dir, "manifest.yaml")
    if os.path.isfile(manifest_path):
        manifest_yaml = open(manifest_path, "r", encoding="utf-8").read()
        try:
            m = _yaml.safe_load(manifest_yaml) or {}
            name = m.get("name", name)
            desc = m.get("description", "")
            description = desc if isinstance(desc, str) else (desc.get("zh-CN") or desc.get("en") or str(desc))
            runtime = m.get("runtime", "markdown")
            triggers = m.get("triggers", [])
        except Exception:
            pass

    skill_md_path = os.path.join(skill_dir, "SKILL.md")
    if os.path.isfile(skill_md_path):
        skill_md = open(skill_md_path, "r", encoding="utf-8").read()

    # Collect code files (*.py)
    for fn in os.listdir(skill_dir):
        fp = os.path.join(skill_dir, fn)
        if os.path.isfile(fp) and fn.endswith(".py"):
            try:
                code_files[fn] = open(fp, "r", encoding="utf-8").read()
            except Exception:
                pass

    if not manifest_yaml and not skill_md:
        raise HTTPException(400, "No manifest.yaml or SKILL.md found in directory")

    # Check for duplicate: same name + same version = reject
    import_version = ""
    if manifest_yaml:
        try:
            _mv = _yaml.safe_load(manifest_yaml) or {}
            import_version = _mv.get("version", "")
        except Exception:
            pass
    for existing in forge._drafts.values():
        if existing.name == name and existing.status in ("draft", "exported", "approved"):
            existing_version = ""
            if existing.manifest_yaml:
                try:
                    em = _yaml.safe_load(existing.manifest_yaml) or {}
                    existing_version = em.get("version", "")
                except Exception:
                    pass
            if existing_version == import_version:
                raise HTTPException(
                    409,
                    f"技能 '{name}' v{import_version} 已存在"
                    f"（ID: {existing.id}, 状态: {existing.status}）。"
                    f"请修改 version 后重新导入。"
                )

    # Create a SkillDraft
    draft_id = f"SF-{_time.strftime('%Y%m%d')}-IMP-{os.urandom(3).hex()}"
    draft = SkillDraft(
        id=draft_id,
        name=name,
        description=description,
        source_experiences=[],
        role=body.get("role", ""),
        scene_pattern="",
        triggers=triggers if isinstance(triggers, list) else [triggers],
        manifest_yaml=manifest_yaml,
        skill_md=skill_md,
        confidence=0.9,
        created_at=_time.time(),
        status="exported",
        runtime=runtime,
        code_files=code_files,
    )

    # Add to forge and save
    forge._drafts[draft_id] = draft
    forge._save_drafts()

    logger.info("Imported skill draft from directory: %s → %s", skill_dir, draft_id)
    return {
        "ok": True,
        "draft_id": draft_id,
        "name": name,
        "runtime": runtime,
        "code_files": list(code_files.keys()),
        "source_dir": skill_dir,
    }


@router.post("/pending-skills/{draft_id}/reject")
async def reject_pending_skill(
    draft_id: str,
    body: dict = Body(default={}),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Reject a skill draft."""
    try:
        from ...skills._skill_forge import get_skill_forge
        forge = get_skill_forge()
        result = forge.reject_draft(draft_id)
        if "error" in result:
            raise HTTPException(404, result["error"])
        return {"ok": True, "draft_id": draft_id, "status": "rejected"}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# PPTX theme gallery — portal UI support for the pptx-author skill's
# template system. Exposes 3 endpoints:
#   GET  /api/portal/pptx-themes                list 10 themes + previews
#   GET  /api/portal/pptx-themes/{id}/preview   serve PNG (inline)
#   POST /api/portal/pptx-themes/recommend      {query} → top-3 matches
# ---------------------------------------------------------------------------

def _load_pptx_loader():
    """Lazy-import _template_loader from the pptx-author skill dir."""
    import importlib.util
    from pathlib import Path
    skill_dir = (Path(__file__).parents[2]
                 / "skills" / "builtin" / "tudou-builtin" / "pptx-author")
    loader_path = skill_dir / "_template_loader.py"
    if not loader_path.exists():
        raise HTTPException(503, "pptx-author skill not installed")
    # Importlib with a namespaced name to avoid polluting sys.modules globally.
    spec = importlib.util.spec_from_file_location(
        "_tudou_pptx_loader", str(loader_path))
    mod = importlib.util.module_from_spec(spec)
    # pptx-author's _template_loader does `from _pptx_helpers import ...` at
    # render time, so make sure skill dir is on sys.path.
    import sys as _sys
    if str(skill_dir) not in _sys.path:
        _sys.path.insert(0, str(skill_dir))
    spec.loader.exec_module(mod)
    return mod


@router.get("/pptx-themes")
async def list_pptx_themes(
    user: CurrentUser = Depends(get_current_user),
):
    """List all PPTX themes with preview URLs."""
    try:
        mod = _load_pptx_loader()
        themes = mod.list_themes()
        for t in themes:
            t["preview_url"] = f"/api/portal/pptx-themes/{t['id']}/preview.png"
        layouts = mod.list_layouts()
        return {"themes": themes, "layouts": layouts, "count": len(themes)}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("list_pptx_themes failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pptx-themes/{theme_id}/preview.png")
async def get_pptx_theme_preview(
    theme_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    """Serve a theme's preview PNG (generates on first access if missing)."""
    from pathlib import Path
    from fastapi.responses import FileResponse
    try:
        mod = _load_pptx_loader()
        skill_dir = (Path(__file__).parents[2]
                     / "skills" / "builtin" / "tudou-builtin" / "pptx-author")
        preview_path = (skill_dir / "templates" / "themes"
                        / theme_id / "preview.png")
        if not preview_path.exists():
            # Regenerate on demand if missing.
            mod.generate_theme_preview(theme_id, force=True)
        if not preview_path.exists():
            raise HTTPException(404, f"preview for '{theme_id}' not found")
        return FileResponse(
            str(preview_path),
            media_type="image/png",
            headers={"Cache-Control": "public, max-age=3600"},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("get_pptx_theme_preview failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pptx-layouts")
async def list_pptx_layouts(
    user: CurrentUser = Depends(get_current_user),
):
    """List all PPTX layouts with showcase preview URLs.

    Each layout exposes:
      - id / name / category / summary / when_to_use
      - in_showcase: True if it's one of the 3 layouts rendered for
        per-theme preview thumbnails

    Response shape::
        {
          "layouts": [
            {"id": "T01_cover", "name": "...", ...,
             "preview_url": "/api/portal/pptx-layouts/T01_cover/preview.png?theme=corporate"},
            ...
          ],
          "showcase_layouts": ["T01_cover", "T24_kpi_metrics", "T26_process_flow"],
          "count": 15,
        }
    """
    try:
        mod = _load_pptx_loader()
        layouts = mod.list_layouts()
        showcase = set(getattr(mod, "PREVIEW_SHOWCASE_LAYOUTS", ()))
        for L in layouts:
            L["in_showcase"] = L["id"] in showcase
            # Default preview against the corporate theme; the UI can
            # swap by adding ?theme=<name> to the URL.
            L["preview_url"] = (
                f"/api/portal/pptx-layouts/{L['id']}/preview.png"
            )
        return {
            "layouts": layouts,
            "showcase_layouts": list(showcase),
            "count": len(layouts),
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("list_pptx_layouts failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pptx-layouts/{layout_id}/preview.png")
async def get_pptx_layout_preview(
    layout_id: str,
    theme: str = "corporate",
    user: CurrentUser = Depends(get_current_user),
):
    """Serve a layout × theme preview PNG (generates on first access).

    Query param ``theme`` defaults to ``corporate``. Generated PNG is
    cached at ``templates/_shared/.previews/<theme>/<layout_id>.png``.
    """
    from pathlib import Path
    from fastapi.responses import FileResponse
    try:
        mod = _load_pptx_loader()
        skill_dir = (Path(__file__).parents[2]
                     / "skills" / "builtin" / "tudou-builtin" / "pptx-author")
        preview_path = (skill_dir / "templates" / "_shared"
                        / ".previews" / theme / f"{layout_id}.png")
        if not preview_path.exists():
            mod.generate_layout_preview(layout_id, theme, force=True)
        if not preview_path.exists():
            raise HTTPException(
                404,
                f"preview for layout '{layout_id}' theme '{theme}' not found",
            )
        return FileResponse(
            str(preview_path),
            media_type="image/png",
            headers={"Cache-Control": "public, max-age=3600"},
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("get_pptx_layout_preview failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/pptx-themes/recommend")
async def recommend_pptx_theme(
    body: dict = Body(...),
    user: CurrentUser = Depends(get_current_user),
):
    """Recommend top-3 themes for a user query.

    Request:  {"query": "AI 大模型发布会", "top_k": 3}
    Response: {"results": [{"id","name","score","matched","description"}, ...]}
    """
    q = str(body.get("query") or "").strip()
    top_k = int(body.get("top_k") or 3)
    if not q:
        raise HTTPException(400, "query required")
    try:
        mod = _load_pptx_loader()
        results = mod.recommend_theme(q, top_k=top_k)
        return {"query": q, "results": results}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("recommend_pptx_theme failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pptx-themes/{theme_id}/demo")
async def list_pptx_theme_demo(
    theme_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    """List pre-rendered demo slides for a theme.

    Returns URLs for each slide_N.png and the full .pptx download path.
    """
    from pathlib import Path
    try:
        mod = _load_pptx_loader()  # validate pptx-author is installed
        skill_dir = (Path(__file__).parents[2]
                     / "skills" / "builtin" / "tudou-builtin" / "pptx-author")
        demo_dir = skill_dir / "templates" / "themes" / theme_id / "demo"
        if not demo_dir.exists():
            raise HTTPException(404, f"demo for '{theme_id}' not found")
        slides = sorted(demo_dir.glob("slide_*.png"),
                        key=lambda p: int(p.stem.split("_")[1]))
        return {
            "theme_id": theme_id,
            "slide_count": len(slides),
            "slides": [
                {
                    "index": int(p.stem.split("_")[1]),
                    "url": f"/api/portal/pptx-themes/{theme_id}/demo/{p.name}",
                }
                for p in slides
            ],
            "pptx_url": f"/api/portal/pptx-themes/{theme_id}/demo.pptx",
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("list_pptx_theme_demo failed")
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/pptx-themes/{theme_id}/demo/{filename}")
async def get_pptx_demo_slide(
    theme_id: str,
    filename: str,
    user: CurrentUser = Depends(get_current_user),
):
    """Serve an individual demo slide PNG."""
    from pathlib import Path
    from fastapi.responses import FileResponse
    # Prevent path traversal
    if "/" in filename or "\\" in filename or not filename.endswith(".png"):
        raise HTTPException(400, "bad filename")
    skill_dir = (Path(__file__).parents[2]
                 / "skills" / "builtin" / "tudou-builtin" / "pptx-author")
    fp = skill_dir / "templates" / "themes" / theme_id / "demo" / filename
    if not fp.exists():
        raise HTTPException(404, "not found")
    return FileResponse(
        str(fp), media_type="image/png",
        headers={"Cache-Control": "public, max-age=3600"})


@router.get("/pptx-themes/{theme_id}/demo.pptx")
async def download_pptx_demo(
    theme_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    """Download the full demo .pptx file (so user can open in PowerPoint)."""
    from pathlib import Path
    from fastapi.responses import FileResponse
    skill_dir = (Path(__file__).parents[2]
                 / "skills" / "builtin" / "tudou-builtin" / "pptx-author")
    fp = skill_dir / "templates" / "themes" / theme_id / "demo.pptx"
    if not fp.exists():
        raise HTTPException(404, "not found")
    return FileResponse(
        str(fp),
        media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
        filename=f"{theme_id}_demo.pptx",
    )


@router.post("/files/exists")
async def check_files_exist(
    body: dict = Body(...),
    user: CurrentUser = Depends(get_current_user),
):
    """Batch existence check for filesystem paths.

    Used by the artifact preview panel to filter out non-existent
    paths (e.g. when the LLM hallucinated a `/workspace/foo.md` write
    that never actually landed on disk) before rendering tabs.

    Request:  {"paths": ["/abs/path1", "~/path2", ...]}
    Response: {"exists": {"/abs/path1": true, "~/path2": false}}
    """
    from pathlib import Path
    paths = body.get("paths") or []
    if not isinstance(paths, list):
        raise HTTPException(400, "paths must be a list")
    out: dict[str, bool] = {}
    for p in paths:
        if not isinstance(p, str) or not p.strip():
            continue
        try:
            out[p] = Path(p).expanduser().is_file()
        except Exception:
            out[p] = False
    return {"exists": out}


@router.post("/file-preview")
async def preview_file(
    body: dict = Body(...),
    user: CurrentUser = Depends(get_current_user),
):
    """Render a preview for common document formats.

    Supported:
      - .md / .markdown / .txt → raw text (frontend renders via
                                   _renderSimpleMarkdown for .md/.markdown,
                                   or <pre> for .txt)
      - .docx                  → simplified HTML (headings, paragraphs,
                                   lists, tables) via python-docx
      - .pdf                   → recommend <iframe src=url> (no server work)
      - .pptx                  → point to /pptx-preview

    Request:  {"path": "/abs/path/to/file"}
    Response: {"kind":"markdown","text":"..."} | {"kind":"html","html":"..."}
               | {"kind":"iframe","url":"..."} | {"kind":"unsupported"}
    """
    from pathlib import Path
    path = str(body.get("path") or "").strip()
    if not path:
        raise HTTPException(400, "path required")
    # Expand ~ + resolve to absolute. Without this, paths like
    # ``~/.tudou_claw/workspaces/<id>/foo.md`` (which the artifact
    # panel sends when the agent's working_dir starts with ~) 404
    # because Path("~/...").exists() doesn't auto-expand the tilde.
    p = Path(path).expanduser()
    if not p.exists() or not p.is_file():
        raise HTTPException(404, f"file not found: {path}")

    # Prevent huge files (>10 MB) from blowing up the response.
    try:
        sz = p.stat().st_size
    except Exception:
        sz = 0
    if sz > 10 * 1024 * 1024:
        raise HTTPException(413, f"file too large: {sz} bytes (limit 10MB)")

    suffix = p.suffix.lower()

    # ── Markdown / plain text ──
    if suffix in (".md", ".markdown", ".txt", ".log"):
        try:
            text = p.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            raise HTTPException(500, f"read failed: {e}")
        kind = "markdown" if suffix in (".md", ".markdown") else "text"
        return {
            "kind": kind,
            "text": text,
            "filename": p.name,
            "size": sz,
        }

    # ── DOCX → HTML ──
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(503, "python-docx not installed")
        try:
            doc = Document(str(p))
        except Exception as e:
            raise HTTPException(500, f"docx parse failed: {e}")

        import html as _html
        parts: list[str] = []
        # Walk document body in order. python-docx's `paragraphs` / `tables`
        # give flat lists — to preserve order we iterate over the body XML.
        from docx.oxml.ns import qn
        body = doc.element.body
        p_tag = qn("w:p")
        tbl_tag = qn("w:tbl")

        # Build a map of (paragraph_xml_element → Paragraph) and
        # (table_xml_element → Table) for lookup.
        p_map = {para._element: para for para in doc.paragraphs}
        t_map = {tbl._element: tbl for tbl in doc.tables}

        def _render_para(para):
            style = (para.style.name if para.style else "") or ""
            raw_text = para.text or ""
            if not raw_text.strip() and not style.startswith("Heading"):
                return "<br>"
            # Decide tag
            tag = "p"
            lvl = 0
            if style.startswith("Heading"):
                try:
                    lvl = int(style.split()[-1])
                    lvl = max(1, min(6, lvl))
                    tag = f"h{lvl}"
                except Exception:
                    tag = "h2"
            elif style in ("List Paragraph", "List Bullet"):
                return f"<li>{_html.escape(raw_text)}</li>"
            # runs may have bold / italic
            pieces: list[str] = []
            for run in para.runs:
                rtext = _html.escape(run.text or "")
                if not rtext:
                    continue
                if run.bold:
                    rtext = f"<strong>{rtext}</strong>"
                if run.italic:
                    rtext = f"<em>{rtext}</em>"
                pieces.append(rtext)
            inner = "".join(pieces) or _html.escape(raw_text)
            return f"<{tag}>{inner}</{tag}>"

        def _render_table(tbl):
            rows_html = []
            for row in tbl.rows:
                cells_html = []
                for cell in row.cells:
                    txt = _html.escape(cell.text or "")
                    cells_html.append(f"<td>{txt}</td>")
                rows_html.append("<tr>" + "".join(cells_html) + "</tr>")
            return ("<table style='border-collapse:collapse;"
                    "border:1px solid #ccc;margin:8px 0'>"
                    + "".join(rows_html) + "</table>")

        in_list = False
        for child in body.iterchildren():
            if child.tag == p_tag and child in p_map:
                rendered = _render_para(p_map[child])
                if rendered.startswith("<li>"):
                    if not in_list:
                        parts.append("<ul>")
                        in_list = True
                    parts.append(rendered)
                else:
                    if in_list:
                        parts.append("</ul>")
                        in_list = False
                    parts.append(rendered)
            elif child.tag == tbl_tag and child in t_map:
                if in_list:
                    parts.append("</ul>")
                    in_list = False
                parts.append(_render_table(t_map[child]))
        if in_list:
            parts.append("</ul>")

        # Basic HTML wrapper with modest styling. Table cells get simple
        # border/padding for readability. No inline JS.
        html_body = "\n".join(parts)
        return {
            "kind": "html",
            "html": html_body,
            "filename": p.name,
            "size": sz,
        }

    # ── PDF → iframe (delegate to browser native viewer) ──
    if suffix == ".pdf":
        # Provide the attachment URL so frontend can iframe it. Portal has
        # an attachment endpoint that serves arbitrary paths with auth.
        import urllib.parse as _url
        attach_url = "/api/portal/attachment?path=" + _url.quote(str(p))
        return {
            "kind": "iframe",
            "url": attach_url,
            "filename": p.name,
            "size": sz,
        }

    # ── PPTX → delegate to existing /pptx-preview ──
    if suffix == ".pptx":
        return {
            "kind": "delegate",
            "delegate_to": "/api/portal/pptx-preview",
            "filename": p.name,
            "size": sz,
        }

    return {
        "kind": "unsupported",
        "suffix": suffix,
        "filename": p.name,
        "size": sz,
    }


@router.post("/pptx-preview")
async def preview_arbitrary_pptx(
    body: dict = Body(...),
    user: CurrentUser = Depends(get_current_user),
):
    """Convert any .pptx file into slide PNG thumbnails. General-purpose
    endpoint for previewing agent-produced PPT files.

    Request:  {"path": "/abs/path/to/file.pptx", "width": 640}
    Response: {"slides": [{"index":1,"data_url":"data:image/png;base64,..."},...]}
    """
    import base64
    from pathlib import Path
    import tempfile, os as _os, shutil

    path = str(body.get("path") or "").strip()
    width = int(body.get("width") or 640)
    if not path:
        raise HTTPException(400, "path required")
    p = Path(path)
    if not p.exists() or not p.is_file():
        raise HTTPException(404, f"file not found: {path}")
    if p.suffix.lower() != ".pptx":
        raise HTTPException(400, "only .pptx supported")

    tmp_dir = tempfile.mkdtemp(prefix="pptx_preview_")
    try:
        mod = _load_pptx_loader()
        pngs = mod.pptx_to_pngs(str(p), tmp_dir, width=width, force=True)
        out = []
        for i, png_path in enumerate(pngs, start=1):
            with open(png_path, "rb") as fh:
                b = base64.b64encode(fh.read()).decode("ascii")
            out.append({
                "index": i,
                "data_url": "data:image/png;base64," + b,
            })
        return {"slide_count": len(out), "slides": out,
                "source_file": str(p)}
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("preview_arbitrary_pptx failed")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        try:
            shutil.rmtree(tmp_dir, ignore_errors=True)
        except Exception:
            pass
