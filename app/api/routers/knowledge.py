"""Knowledge management and RAG router — knowledge base, search, RAG providers."""
from __future__ import annotations

import logging
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Query, Body

from ..deps.hub import get_hub
from ..deps.auth import CurrentUser, get_current_user

logger = logging.getLogger("tudouclaw.api.knowledge")

router = APIRouter(prefix="/api/portal", tags=["knowledge"])


# ---------------------------------------------------------------------------
# Knowledge base listing and management
# ---------------------------------------------------------------------------

@router.get("/knowledge")
async def list_knowledge_entries(
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """List knowledge base entries — matches legacy portal_routes_get."""
    try:
        from ... import knowledge
        entries = knowledge.list_entries()
        return {"entries": entries}
    except (ImportError, Exception) as e:
        return {"entries": []}


@router.get("/knowledge/search")
async def search_knowledge(
    q: str = Query("", description="Search query"),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Search knowledge base — matches legacy portal_routes_get."""
    try:
        if not q:
            return {"entries": []}
        from ... import knowledge
        entries = knowledge.search(q)
        return {"entries": entries}
    except (ImportError, Exception) as e:
        return {"entries": []}


@router.post("/knowledge")
async def add_knowledge_entry(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Add a knowledge entry to the shared Knowledge Wiki.

    Bug fix (Nov 2026): the old implementation imported a non-existent
    ``core.memory.get_knowledge_manager``, caught the ImportError with
    a too-broad ``except (ImportError, Exception)``, then fell back to
    a non-existent ``hub.add_knowledge_entry`` method, and returned a
    stub ``{"title": title}`` — UI reported success but nothing was
    persisted. Now routed directly through the real module-level API
    in ``app.knowledge``.
    """
    try:
        title = body.get("title", "").strip()
        content = body.get("content", "").strip()
        tags = body.get("tags", []) or []
        if not title or not content:
            raise HTTPException(400, "title and content are required")
        if not isinstance(tags, list):
            tags = []
        from ... import knowledge as _kb
        entry = _kb.add_entry(title, content, tags)
        return {"ok": True, "entry": entry}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/knowledge/{entry_id}")
async def update_knowledge_entry(
    entry_id: str,
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Update a knowledge entry by id. Same bug fix as add_knowledge_entry."""
    try:
        title = body.get("title")
        content = body.get("content")
        tags = body.get("tags")
        if tags is not None and not isinstance(tags, list):
            tags = None
        from ... import knowledge as _kb
        entry = _kb.update_entry(entry_id, title=title,
                                  content=content, tags=tags)
        if entry:
            return {"ok": True, "entry": entry}
        raise HTTPException(404, "Entry not found")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/knowledge/{entry_id}/delete")
async def delete_knowledge_entry(
    entry_id: str,
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Delete a knowledge entry by id. Returns 404 if entry not found."""
    try:
        from ... import knowledge as _kb
        ok = _kb.delete_entry(entry_id)
        if not ok:
            raise HTTPException(404, f"Entry '{entry_id}' not found")
        return {"ok": True, "deleted_id": entry_id}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# RAG provider management
# ---------------------------------------------------------------------------

@router.get("/rag/providers")
async def list_rag_providers(
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """List available RAG providers — matches legacy portal_routes_get."""
    try:
        from ...rag_provider import get_rag_registry
        reg = get_rag_registry()
        return {"providers": [p.to_dict() for p in reg.list_providers()]}
    except (ImportError, Exception):
        return {"providers": []}


# ---------------------------------------------------------------------------
# RAG provider management (register, update, delete, rebuild)
# ---------------------------------------------------------------------------

@router.post("/rag/providers")
async def register_rag_provider(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Register a new RAG provider."""
    try:
        from ...rag_provider import get_rag_registry
        reg = get_rag_registry()
        entry = reg.register(
            name=body.get("name", ""),
            kind=body.get("kind", "remote"),
            base_url=body.get("base_url", ""),
            api_key=body.get("api_key", ""),
            config=body.get("config", {}),
        )
        return entry.to_dict() if hasattr(entry, "to_dict") else {"ok": True, "entry": entry}
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except ValueError as e:
        # register() rejects ghost entries (empty remote name+url). That
        # is a client error, not a server error — surface as 400 so the
        # Portal form can show the validation message.
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/rag/providers/{provider_id}/update")
async def update_rag_provider(
    provider_id: str,
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Update an existing RAG provider."""
    try:
        from ...rag_provider import get_rag_registry
        reg = get_rag_registry()
        kwargs = {k: v for k, v in body.items() if k != "provider_id"}
        entry = reg.update(provider_id, **kwargs)
        if entry:
            return entry.to_dict() if hasattr(entry, "to_dict") else {"ok": True}
        raise HTTPException(404, "Provider not found")
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/rag/providers/{provider_id}/delete")
async def delete_rag_provider(
    provider_id: str,
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Delete a RAG provider."""
    try:
        from ...rag_provider import get_rag_registry
        reg = get_rag_registry()
        ok = reg.remove(provider_id)
        return {"ok": ok}
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/rag/providers/{provider_id}/rebuild")
async def rebuild_rag_provider(
    provider_id: str,
    body: dict = Body(default={}),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Rebuild/reindex a RAG provider's data."""
    try:
        from ...rag_provider import get_rag_registry
        reg = get_rag_registry()
        provider = reg.get(provider_id) if hasattr(reg, "get") else None
        if not provider:
            raise HTTPException(404, "Provider not found")
        if hasattr(reg, "rebuild"):
            result = reg.rebuild(provider_id)
            return {"ok": True, "result": result}
        if hasattr(reg, "reindex"):
            result = reg.reindex(provider_id)
            return {"ok": True, "result": result}
        # Fallback: re-list collections to confirm provider is reachable
        colls = reg.list_collections(provider_id) if hasattr(reg, "list_collections") else []
        return {
            "ok": True,
            "message": "provider reachable, no explicit rebuild needed",
            "collections": len(colls),
        }
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/rag/index")
async def rag_index(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Index documents into RAG (batch indexing)."""
    try:
        provider_id = body.get("provider_id", "")
        collection = body.get("collection", "knowledge")
        documents = body.get("documents", [])
        if not documents:
            raise HTTPException(400, "Missing documents")

        from ...rag_provider import get_rag_registry
        reg = get_rag_registry()
        count = reg.ingest(
            provider_id=provider_id,
            collection=collection,
            documents=documents,
        )
        return {"ok": True, "count": count}
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# RAG search and ingestion
# ---------------------------------------------------------------------------

@router.post("/rag/search")
async def search_rag(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Search using RAG.

    Bug fix (Nov 2026): old code checked ``hasattr(hub, 'search_rag')``
    — hub has no such method, so the endpoint silently returned
    ``{"results": []}`` for every query. Now routes through the real
    RAG provider registry.
    """
    try:
        query = (body.get("query") or "").strip()
        if not query:
            raise HTTPException(400, "query is required")
        provider = body.get("provider", "") or ""
        collection = body.get("collection", "") or ""
        if not collection:
            raise HTTPException(400, "collection is required")
        limit = int(body.get("limit", 10) or 10)

        from ...rag_provider import get_rag_registry
        results = get_rag_registry().search(
            provider, collection, query, top_k=limit,
        )
        return {
            "ok": True,
            "results": [
                r.to_dict() if hasattr(r, "to_dict") else r
                for r in (results or [])
            ],
            "count": len(results or []),
        }
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/rag/ingest")
async def ingest_rag_documents(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Ingest documents into RAG system.

    Bug fix (Nov 2026): same silent-success pattern as /rag/search.
    Now routes through the real registry.
    """
    try:
        provider = body.get("provider", "") or ""
        collection = body.get("collection", "") or ""
        documents = body.get("documents", []) or []
        if not documents or not isinstance(documents, list):
            raise HTTPException(400, "documents must be a non-empty list")
        if not collection:
            raise HTTPException(400, "collection is required")

        from ...rag_provider import get_rag_registry
        count = get_rag_registry().ingest(provider, collection, documents)
        return {
            "ok": True,
            "ingested": count,
            "requested": len(documents),
        }
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Domain knowledge base management
# ---------------------------------------------------------------------------

@router.post("/domain-kb/list")
async def list_domain_knowledge_bases(
    body: dict = Body(default={}),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """List domain knowledge bases."""
    try:
        # Prefer the hub convenience if present; otherwise go direct to
        # the store. The hub method was never fully wired up, and silent
        # `return []` here matched the silent `return {"ok": True}` in
        # create() below → users thought the feature was broken.
        if hasattr(hub, "list_domain_knowledge_bases"):
            kbs = hub.list_domain_knowledge_bases()
        else:
            from ...rag_provider import get_domain_kb_store
            kbs = get_domain_kb_store().list_all()
        kbs_list = [k.to_dict() if hasattr(k, "to_dict") else k for k in kbs]
        return {"knowledge_bases": kbs_list}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# Curated list of embedding models the UI proactively recommends.
# Marker fields:
#   recommended  — UI pre-selects this (and shows ⭐ badge)
#   note         — one-liner shown under the dropdown when selected
#   size_mb      — first-download cost (for the "(~XGB)" hint)
# Anything cached locally (see _scan_local_st_models) is also accepted
# at create time so admins who pre-downloaded their own model don't
# have to wait for us to add it here.
CURATED_EMBEDDING_MODELS = [
    {"id": "", "label": "默认 (服务器配置)", "dim": None, "size_mb": None,
     "note": "fallback to DEFAULT_EMBEDDING_MODEL"},
    {"id": "BAAI/bge-m3", "label": "BAAI/bge-m3", "dim": 1024, "size_mb": 2300,
     "note": "多语言 + 中文最强 + 跨语言对齐, 当前最佳通用选项",
     "recommended": True},
    {"id": "BAAI/bge-large-zh-v1.5", "label": "BAAI/bge-large-zh-v1.5", "dim": 1024, "size_mb": 1300,
     "note": "纯中文 KB + 资源紧张时用 (体积比 bge-m3 小 70%)"},
    {"id": "all-MiniLM-L6-v2", "label": "all-MiniLM-L6-v2", "dim": 384, "size_mb": 80,
     "note": "英文为主, 体积小, 速度快 (旧默认)"},
    {"id": "paraphrase-multilingual-MiniLM-L12-v2",
     "label": "paraphrase-multilingual-MiniLM-L12-v2", "dim": 384, "size_mb": 470,
     "note": "多语言 + 体积小"},
]


def _scan_local_st_models() -> list[dict]:
    """Scan the HF cache directories for sentence-transformer models the
    user has already downloaded. Lets admins use any model they've
    pre-fetched without needing it added to the curated list.

    Detection: a HuggingFace model dir is a sentence-transformer if its
    snapshot contains ``config_sentence_transformers.json``. Plain
    transformers / LLM weights skip — those won't work as embedders.
    """
    import os
    from pathlib import Path
    out: dict[str, dict] = {}
    seen_roots = set()
    for env_var in ("SENTENCE_TRANSFORMERS_HOME", "HF_HOME", "HF_HUB_CACHE"):
        root = os.environ.get(env_var)
        if not root:
            continue
        for search_root in (root, os.path.join(root, "hub")):
            if not os.path.isdir(search_root) or search_root in seen_roots:
                continue
            seen_roots.add(search_root)
            try:
                entries = os.listdir(search_root)
            except OSError:
                continue
            for entry in entries:
                if not entry.startswith("models--"):
                    continue
                model_dir = os.path.join(search_root, entry)
                snap_dir = os.path.join(model_dir, "snapshots")
                if not os.path.isdir(snap_dir):
                    continue
                is_st = False
                try:
                    for sha in os.listdir(snap_dir):
                        marker = os.path.join(snap_dir, sha, "config_sentence_transformers.json")
                        if os.path.isfile(marker):
                            is_st = True
                            break
                except OSError:
                    continue
                if not is_st:
                    continue
                # "models--BAAI--bge-m3" → "BAAI/bge-m3"
                parts = entry[len("models--"):].split("--")
                model_id = "/".join(parts) if len(parts) > 1 else parts[0]
                # Approx size
                try:
                    size_mb = sum(
                        f.stat().st_size for f in Path(model_dir).rglob("*") if f.is_file()
                    ) // (1024 * 1024)
                except OSError:
                    size_mb = None
                out.setdefault(model_id, {
                    "id": model_id,
                    "label": model_id,
                    "dim": None,        # not known without loading the model
                    "size_mb": size_mb,
                    "note": "本地已缓存",
                    "local": True,
                })
    return list(out.values())


def _allowed_model_ids() -> set[str]:
    """Anything in the curated list OR already cached locally is OK."""
    ids = {m["id"] for m in CURATED_EMBEDDING_MODELS}
    for m in _scan_local_st_models():
        ids.add(m["id"])
    return ids


@router.get("/domain-kb/embedding-models")
async def list_embedding_models(user: CurrentUser = Depends(get_current_user)):
    """Catalog of selectable embedding models for KB creation.

    Returns curated entries + any sentence-transformer model the admin
    has already pre-downloaded (so they can pick a custom model without
    needing it added to the curated list).
    """
    curated = list(CURATED_EMBEDDING_MODELS)
    curated_ids = {m["id"] for m in curated}
    extra = [m for m in _scan_local_st_models() if m["id"] not in curated_ids]
    return {"models": curated, "local_extra": extra}


# Cross-encoder rerankers — separate catalog from embeddings because the
# trade-off is different (rerankers are optional; embeddings always run).
# The flag ``recommended: True`` pre-selects in the UI.
CURATED_RERANKER_MODELS = [
    {"id": "", "label": "不使用 (默认)", "size_mb": None,
     "note": "vector search 结果直接返回, 速度最快"},
    {"id": "BAAI/bge-reranker-v2-m3", "label": "BAAI/bge-reranker-v2-m3",
     "size_mb": 568,
     "note": "多语言精排, 与 bge-m3 embedding 配套使用 (推荐)",
     "recommended": True},
    {"id": "BAAI/bge-reranker-large", "label": "BAAI/bge-reranker-large",
     "size_mb": 1100,
     "note": "英文精排, 召回侧用 bge-large-en 时配它"},
    {"id": "BAAI/bge-reranker-base", "label": "BAAI/bge-reranker-base",
     "size_mb": 280,
     "note": "轻量版精排, 体积小一半"},
]


def _allowed_reranker_ids() -> set[str]:
    """Allow-list for reranker model field. Empty string is always
    valid (== no rerank). Curated catalog ∪ anything cached locally."""
    ids = {m["id"] for m in CURATED_RERANKER_MODELS}
    # Local-cache scan: cross-encoders also have config.json with
    # text-classification head; we re-use the embedding scanner since
    # the BAAI/bge-reranker-* dirs follow the same models--owner--name
    # naming, even though they're not technically "sentence transformers".
    # The act of pre-downloading is itself the vetting.
    for m in _scan_local_st_models():
        ids.add(m["id"])
    return ids


@router.get("/domain-kb/reranker-models")
async def list_reranker_models(user: CurrentUser = Depends(get_current_user)):
    """Catalog of cross-encoder rerankers for KB creation."""
    curated = list(CURATED_RERANKER_MODELS)
    curated_ids = {m["id"] for m in curated}
    # Show locally-cached models that look reranker-ish (heuristic:
    # name contains "rerank" or "cross" — avoids polluting the dropdown
    # with the bi-encoder embedders, which would technically load via
    # CrossEncoder but produce garbage scores).
    extra = []
    for m in _scan_local_st_models():
        nm = m["id"].lower()
        if m["id"] in curated_ids:
            continue
        if "rerank" in nm or "cross" in nm or "reranker" in nm:
            extra.append({**m, "label": m["id"]})
    return {"models": curated, "local_extra": extra}


@router.post("/domain-kb/create")
async def create_domain_knowledge_base(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Create a domain knowledge base.

    Bug fix: the old implementation only created a KB when the hub
    exposed a ``create_domain_knowledge_base`` method, which it never
    did. The endpoint then silently returned ``{"ok": True}`` — UI
    reported success but nothing was persisted. Now we call the store
    directly.
    """
    try:
        name = body.get("name", "")
        if not name:
            raise HTTPException(400, "Missing name")

        description = body.get("description", "") or ""
        provider_id = body.get("provider_id", "") or ""
        embedding_model = (body.get("embedding_model") or "").strip()
        # Allow-list only applies to LOCAL sentence-transformers models.
        # When the KB routes embedding through an LLM provider
        # (``embedding_provider_id`` set), ``embedding_model`` is just
        # whatever model id that provider exposes — free text, no gate.
        _emb_provider_id_pre = (body.get("embedding_provider_id") or "").strip()
        if (embedding_model
                and not _emb_provider_id_pre
                and embedding_model not in _allowed_model_ids()):
            raise HTTPException(
                400,
                f"embedding_model not allowed: {embedding_model!r}. "
                f"GET /domain-kb/embedding-models for the catalog, "
                f"or pre-download into ~/.tudou_claw/hf_cache/. "
                f"(For LLM-provider embedding, set "
                f"``embedding_provider_id`` and use the provider's "
                f"native model id.)",
            )
        reranker_model = (body.get("reranker_model") or "").strip()
        if reranker_model and reranker_model not in _allowed_reranker_ids():
            raise HTTPException(
                400,
                f"reranker_model not allowed: {reranker_model!r}. "
                f"GET /domain-kb/reranker-models for the catalog.",
            )
        tags = body.get("tags") or []
        if not isinstance(tags, list):
            tags = []

        # Optional: route embedding through an existing LLM provider.
        # When ``embedding_provider_id`` is set, the KB looks up that
        # provider in app.llm at ingest/search time and POSTs to its
        # /v1/embeddings endpoint. ``embedding_model`` is then the
        # model id at that provider (free text — we don't gate it
        # against the local sentence-transformers allow-list).
        embedding_provider_id = (body.get("embedding_provider_id") or "").strip()
        if embedding_provider_id:
            # When using LLM-provider embedding, the model is whatever
            # that provider supports — bypass the local allow-list.
            embedding_model = (body.get("embedding_model") or "").strip()
            if not embedding_model:
                raise HTTPException(
                    400,
                    "embedding_model is required when "
                    "embedding_provider_id is set (e.g. "
                    "'text-embedding-3-small', 'nomic-embed-text', "
                    "'embedding-2' — depends on which provider).",
                )

        # Prefer hub method if a future version exposes it.
        if hasattr(hub, "create_domain_knowledge_base"):
            kb = hub.create_domain_knowledge_base(body)
            if hasattr(kb, "to_dict"):
                kb = kb.to_dict()
            return {"ok": True, "knowledge_base": kb}

        # Direct-to-store path — what the endpoint should have been doing
        # from day one.
        from ...rag_provider import get_domain_kb_store
        store = get_domain_kb_store()
        kb = store.create(
            name=name,
            description=description,
            provider_id=provider_id,
            tags=[str(t).strip() for t in tags if str(t).strip()],
            embedding_model=embedding_model,
            reranker_model=reranker_model,
            embedding_provider_id=embedding_provider_id,
        )
        return {"ok": True, "knowledge_base": kb.to_dict()}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/domain-kb/search")
async def search_domain_knowledge_base(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Search a domain knowledge base.

    Body:
      • ``kb_id``         (required) — target knowledge base
      • ``query``         (required) — natural-language query
      • ``top_k``         default 5  — max results to return
      • ``max_distance``  default 1.0 (no filter) — drop results
        whose vector distance exceeds this. With bge-m3 + cosine,
        empirical thresholds (verified 2026-05-04 on docs/ corpus):
          - relevant queries:   distance ≈ 0.30 - 0.55
          - irrelevant queries: distance ≈ 0.75 - 0.85
        So ``max_distance: 0.65`` is a reasonable cutoff to filter
        out garbage matches on noisy queries.
      • ``min_rerank_score`` default null — when a reranker is
        configured, drop results below this cross-encoder score
        (typical: keep > 0.0).

    Returns: ``{"results": [...], "filtered_count": N}`` — N is the
    number of raw results dropped by max_distance / min_rerank_score
    so callers can detect "noisy query → no useful hits".
    """
    try:
        from ...rag_provider import get_domain_kb_store, get_rag_registry
        store = get_domain_kb_store()
        kb_id = body.get("kb_id", "")
        query = body.get("query", "")
        top_k = int(body.get("top_k", 5))
        max_distance = float(body.get("max_distance", 1.0))
        min_rerank_raw = body.get("min_rerank_score")
        min_rerank = float(min_rerank_raw) if min_rerank_raw is not None else None
        if not kb_id or not query:
            raise HTTPException(400, "Missing kb_id or query")
        kb = store.get(kb_id)
        if not kb:
            raise HTTPException(404, "knowledge base not found")
        # RAG v1-B: prefer hybrid (BM25 + vector + RRF) when available.
        # Both code paths involve sync embedding (bge-m3.encode) which
        # would block the asyncio event loop — push to threadpool.
        from starlette.concurrency import run_in_threadpool
        reg = get_rag_registry()
        if hasattr(reg, "hybrid_search"):
            try:
                results = await run_in_threadpool(
                    reg.hybrid_search,
                    kb.provider_id, kb.collection, query, top_k,
                )
            except Exception:
                results = await run_in_threadpool(
                    reg.search, kb.provider_id, kb.collection, query, top_k,
                )
        else:
            results = await run_in_threadpool(
                reg.search, kb.provider_id, kb.collection, query, top_k,
            )
        # Threshold filtering — done AFTER retrieval so callers can
        # see how many got dropped (signal for "noisy query").
        before = len(results)
        if max_distance < 1.0:
            results = [r for r in results if (r.distance or 0) <= max_distance]
        if min_rerank is not None:
            def _rerank_ok(r):
                rs = (r.metadata or {}).get("rerank_score")
                return rs is None or rs >= min_rerank
            results = [r for r in results if _rerank_ok(r)]
        filtered_count = before - len(results)
        return {
            "results": [r.to_dict() for r in results],
            "filtered_count": filtered_count,
        }
    except HTTPException:
        raise
    except ImportError:
        return {"results": [], "filtered_count": 0}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.get("/domain-kb/{kb_id}/image/{filename}")
async def get_domain_kb_image(
    kb_id: str,
    filename: str,
    user: CurrentUser = Depends(get_current_user),
):
    """Serve an original image stored alongside its OCR text.

    Path is referenced from chunk metadata (``image_url``) — RAG
    retrieval surfaces this URL so chat / UI can render the image
    inline (markdown ``![](...)``) right next to the OCR'd text.

    Auth: JWT — same as other portal endpoints. We don't expose
    raw bytes via ``X-Hub-Secret`` because images may contain
    sensitive content (screenshots of dashboards, scanned docs).

    Path traversal: ``filename`` is constrained — must match the
    ``base_id.ext`` shape stamped at ingest time, and we resolve
    against the kb's directory only. Anything else 404.
    """
    import re as _re
    import os as _os
    from fastapi.responses import FileResponse

    # Sanitise — only allow our own filename pattern (alnum + dot + underscore).
    if not _re.fullmatch(r"[A-Za-z0-9._-]+", filename):
        raise HTTPException(400, "invalid filename")
    # KB id similarly safe.
    if not _re.fullmatch(r"[A-Za-z0-9_.-]+", kb_id):
        raise HTTPException(400, "invalid kb_id")

    from ...paths import data_dir as _resolve_data_dir
    img_path = _resolve_data_dir() / "kb_images" / kb_id / filename
    img_path = img_path.resolve()
    # Defense-in-depth: confirm the resolved path stays under kb_images.
    expected_root = (_resolve_data_dir() / "kb_images" / kb_id).resolve()
    try:
        img_path.relative_to(expected_root)
    except ValueError:
        raise HTTPException(404, "not found")
    if not img_path.is_file():
        raise HTTPException(404, "not found")

    # MIME hint from extension.
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    mime_map = {
        "jpg": "image/jpeg", "jpeg": "image/jpeg",
        "png": "image/png", "webp": "image/webp",
        "gif": "image/gif", "bmp": "image/bmp",
        "tif": "image/tiff", "tiff": "image/tiff",
    }
    return FileResponse(
        path=str(img_path),
        media_type=mime_map.get(ext, "application/octet-stream"),
    )


@router.post("/domain-kb/update")
async def update_domain_knowledge_base(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Update a domain knowledge base metadata.

    Accepts: ``name`` / ``description`` / ``tags`` /
    ``embedding_model`` / ``reranker_model``. Only fields present
    in the body are touched. Pass ``""`` to clear a model field
    back to the server default.

    Notes on switching models:
      • ``embedding_model``: switching invalidates the existing
        ingested chunks (different vector space). Recommend
        re-ingest after changing this.
      • ``reranker_model``: pure inference-time, no re-ingest
        required. Empty string disables reranking entirely.
    """
    try:
        from ...rag_provider import get_domain_kb_store
        store = get_domain_kb_store()
        kb_id = body.get("id", "")
        # Validate model ids against the curated lists (same gate as create)
        embedding_model = body.get("embedding_model")
        if embedding_model is not None and embedding_model:
            if embedding_model not in _allowed_model_ids():
                raise HTTPException(
                    400,
                    f"embedding_model not allowed: {embedding_model!r}. "
                    f"GET /domain-kb/embedding-models for catalog.",
                )
        reranker_model = body.get("reranker_model")
        if reranker_model is not None and reranker_model:
            if reranker_model not in _allowed_reranker_ids():
                raise HTTPException(
                    400,
                    f"reranker_model not allowed: {reranker_model!r}. "
                    f"GET /domain-kb/reranker-models for catalog.",
                )
        kb = store.update(
            kb_id,
            name=body.get("name"),
            description=body.get("description"),
            tags=body.get("tags"),
            embedding_model=embedding_model,
            reranker_model=reranker_model,
        )
        if kb:
            return kb.to_dict()
        raise HTTPException(404, "knowledge base not found")
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/domain-kb/delete")
async def delete_domain_knowledge_base(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Delete a domain knowledge base."""
    try:
        from ...rag_provider import get_domain_kb_store
        store = get_domain_kb_store()
        kb_id = body.get("id", "")
        ok = store.delete(kb_id)
        return {"ok": ok}
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── file parsing helper (shared by single-file & folder import) ──────
# Extensions we know how to parse. Unknown extensions are skipped with
# a warning in folder-import; for single-file they fall through to
# best-effort text decode.
_KNOWN_KB_EXTENSIONS: tuple[str, ...] = (
    "txt", "md", "markdown", "rst", "log",
    "csv", "tsv",
    "json", "yaml", "yml",
    "html", "htm",
    "pdf", "docx",
    "py", "js", "ts", "go", "rs", "java", "c", "cpp", "h",
    # Images — OCR'd via rapidocr-onnxruntime (lazy-loaded, ~100 MB
    # models on first use). Text-in-image only; pure visual content
    # — diagrams, charts, screenshots — gets whatever caption text
    # the OCR can pick up. For semantic image search use a multimodal
    # embedder instead (planned, not built yet).
    "jpg", "jpeg", "png", "webp", "bmp", "tif", "tiff", "gif",
)


# Lazy-init OCR engine. Module-level cache avoids re-loading the
# 100 MB ONNX model on every image. Init triggers download to
# ${TUDOU_HF_CACHE} on first run.
_ocr_engine = None
_ocr_init_lock = None


def _get_ocr_engine():
    """Return a lazily-initialised RapidOCR instance, or None when
    the dep isn't installed (returns None silently — caller should
    skip image files with a clear error)."""
    global _ocr_engine, _ocr_init_lock
    if _ocr_engine is not None:
        return _ocr_engine
    if _ocr_init_lock is None:
        import threading
        _ocr_init_lock = threading.Lock()
    with _ocr_init_lock:
        if _ocr_engine is not None:
            return _ocr_engine
        try:
            from rapidocr_onnxruntime import RapidOCR
        except ImportError:
            logger.warning(
                "RapidOCR not installed — image files cannot be ingested. "
                "Install via `pip install rapidocr-onnxruntime` to enable.",
            )
            _ocr_engine = False  # negative cache; falsy
            return None
        try:
            _ocr_engine = RapidOCR()
            logger.info("RapidOCR engine initialised")
        except Exception as e:
            logger.error("RapidOCR init failed: %s", e)
            _ocr_engine = False
            return None
        return _ocr_engine


# ── Per-format parsers (each returns markdown w/ # / ## / ### headings) ──
#
# Architecture: every parser produces normalized markdown text.
# Heading hierarchy is encoded as ``#`` prefixes regardless of the
# source format. This keeps the downstream chunker
# (``_chunk_text_for_rag`` → ``_split_by_headings``) format-agnostic
# — it splits on '#' lines and stamps ``heading_path`` metadata on
# chunks. ``mode="outline"`` and aggregate queries then work
# uniformly across docx / pdf / html / md.
#
# Adding a new format: implement ``_NEW_to_markdown(raw) -> (text, method)``,
# follow the contract:
#   - return ``("", "ext_unsupported")`` if the dep is missing
#   - emit ``# heading`` / ``## heading`` etc. on their own lines
#   - never raise; fall back to best-effort plain text on errors


def _docx_to_markdown(raw: bytes) -> tuple[str, str]:
    """Word doc → markdown. Word ``Heading N`` / ``Title`` styles
    become ``#`` prefixes; tables flatten to pipe-joined rows."""
    try:
        import docx  # type: ignore
        import io as _io
    except ImportError:
        return "", "ext_unsupported_docx"
    try:
        doc = docx.Document(_io.BytesIO(raw))
        parts: list[str] = []
        for para in doc.paragraphs:
            txt = para.text.strip()
            if not txt:
                continue
            lvl = 0
            try:
                style_name = (para.style.name or "").strip()
            except Exception:
                style_name = ""
            if style_name == "Title":
                lvl = 1
            elif style_name.startswith("Heading "):
                try:
                    lvl = max(1, min(6, int(style_name.split()[-1])))
                except (ValueError, IndexError):
                    lvl = 0
            if lvl > 0:
                parts.append(f"\n{'#' * lvl} {txt}\n")
            else:
                parts.append(txt)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n\n".join(parts), "python-docx-headings"
    except Exception as e:
        logger.debug("docx parse failed: %s", e)
        return raw.decode("utf-8", errors="replace"), "raw_decode"


def _html_to_markdown(raw: bytes) -> tuple[str, str]:
    """HTML → markdown. ``<h1>..<h6>`` tags become ``#``..``######``;
    ``<p>`` and other text-bearing tags become flat paragraphs.
    Falls back to ``BeautifulSoup.get_text`` then regex-strip on any
    error (preserves backward compat with the old behaviour)."""
    try:
        from bs4 import BeautifulSoup  # type: ignore
    except ImportError:
        return "", "ext_unsupported_html"
    try:
        soup = BeautifulSoup(raw, "html.parser")
        # Strip script / style — they pollute downstream RAG.
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        parts: list[str] = []
        # Walk body (or whole doc) in document order, materialising
        # heading tags as markdown and other text-bearing nodes as paragraphs.
        body = soup.body or soup
        for node in body.descendants:
            name = getattr(node, "name", None)
            if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                lvl = int(name[1])
                txt = node.get_text(" ", strip=True)
                if txt:
                    parts.append(f"\n{'#' * lvl} {txt}\n")
        # Now grab the remaining text content (everything that wasn't a heading)
        # — simplest: do another pass for non-heading text nodes.
        # Trade-off: paragraph order may not exactly match heading-interleaved
        # original, but for RAG retrieval this is fine; what matters is the
        # heading_path metadata for each chunk after split.
        body_text = soup.get_text("\n", strip=True)
        if body_text:
            parts.append(body_text)
        if parts:
            return "\n\n".join(parts), "beautifulsoup-headings"
        return body_text, "beautifulsoup"
    except Exception as e:
        logger.debug("html parse failed: %s", e)
        import re as _re
        return _re.sub(r"<[^>]+>", "", raw.decode("utf-8", errors="replace")), "regex_strip"


def _pdf_to_markdown(raw: bytes) -> tuple[str, str]:
    """PDF → markdown. Uses pymupdf's TOC (PDF outline / bookmarks)
    when present to inject ``#`` markers at the right page/position;
    falls back to pdfplumber → pymupdf → raw decode.
    Caveat: PDFs without a TOC stay flat — heading detection from
    visual layout (font size / boldness) is a known hard problem
    and out of scope for this iteration."""
    import io as _io

    # Path 1: pymupdf with TOC injection (best — preserves heading hierarchy)
    try:
        import fitz  # type: ignore
        doc = fitz.open(stream=raw, filetype="pdf")
        toc = doc.get_toc(simple=True) or []  # [[level, title, page1], ...]
        if toc:
            # Group entries by page so we can inject before that page's text.
            from collections import defaultdict
            entries_by_page: dict[int, list[tuple[int, str]]] = defaultdict(list)
            for lvl, title, page in toc:
                if page >= 1 and title:
                    entries_by_page[page - 1].append(  # 0-indexed
                        (max(1, min(6, lvl)), title.strip()),
                    )
            parts: list[str] = []
            for i, page in enumerate(doc):
                # Inject any TOC entries that point at this page first
                for lvl, title in entries_by_page.get(i, []):
                    parts.append(f"\n{'#' * lvl} {title}\n")
                txt = page.get_text() or ""
                if txt.strip():
                    parts.append(txt)
            return "\n\n".join(parts), "pymupdf-toc"
        # No TOC → fall through to flat text (still pymupdf, no headings)
        text = "\n\n".join(page.get_text() for page in doc)
        return text, "pymupdf"
    except Exception as e:
        logger.debug("pymupdf parse failed: %s", e)

    # Path 2: pdfplumber (better text extraction for some PDFs, no TOC support)
    try:
        import pdfplumber  # type: ignore
        with pdfplumber.open(_io.BytesIO(raw)) as pdf:
            text = "\n\n".join(p.extract_text() or "" for p in pdf.pages)
        return text, "pdfplumber"
    except Exception as e:
        logger.debug("pdfplumber parse failed: %s", e)

    # Path 3: raw decode (last resort)
    return raw.decode("utf-8", errors="replace"), "raw_decode"


def _parse_file_bytes_to_text(raw: bytes, file_name: str) -> tuple[str, str]:
    """Shared file-to-text logic.

    Dispatches to format-specific parsers each of which produces
    normalised markdown (heading hierarchy as ``#`` prefixes, body
    text as plain paragraphs). Returns ``(text, method)``.

    Best-effort — always returns SOMETHING (may be raw-decoded)
    rather than raising, so folder walks don't abort on one bad file.
    """
    ext = file_name.rsplit(".", 1)[-1].lower() if "." in file_name else "txt"
    text = ""
    method = "raw"
    if ext == "pdf":
        text, method = _pdf_to_markdown(raw)
    elif ext == "docx":
        text, method = _docx_to_markdown(raw)
    elif ext in ("html", "htm"):
        text, method = _html_to_markdown(raw)
    elif ext in ("jpg", "jpeg", "png", "webp", "bmp", "tif", "tiff", "gif"):
        # Image OCR via RapidOCR. The engine returns line-level boxes
        # with text + confidence; we join lines into a flat string so
        # the existing chunker / embedder treats it like any text doc.
        engine = _get_ocr_engine()
        if engine is None:
            logger.warning(
                "Skipping image %s: OCR engine not available", file_name,
            )
            text = ""
            method = "ocr_unavailable"
        else:
            try:
                # RapidOCR accepts bytes / numpy / path. Bytes is the
                # simplest path here (we already have raw).
                result, _elapsed = engine(raw)
                lines: list[str] = []
                if result:
                    for entry in result:
                        # entry = [box_pts, text, confidence] in v1+
                        if isinstance(entry, (list, tuple)) and len(entry) >= 2:
                            line_text = str(entry[1] or "").strip()
                            if line_text:
                                lines.append(line_text)
                text = "\n".join(lines)
                method = "rapidocr"
                if not text:
                    method = "rapidocr_empty"
            except Exception as e:
                logger.warning(
                    "OCR failed on %s: %s — skipping", file_name, e,
                )
                text = ""
                method = f"ocr_error:{type(e).__name__}"
    else:
        for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                text = raw.decode(enc)
                method = enc
                break
            except (UnicodeDecodeError, LookupError):
                continue
        if not text:
            text = raw.decode("utf-8", errors="replace")
            method = "utf8_replace"
    return text, method


# ── RAG v1-A: heading-aware recursive chunker ────────────────────────
# Splits in priority order:
#   1. Markdown headings (# / ## / ### / ...)  — strongest boundary
#   2. Double-newline paragraphs
#   3. Sentence terminators (. / 。 / ! / ？ etc.)
#   4. Raw character slice (last resort, should rarely fire)
# Adds 15% overlap between adjacent chunks to keep cross-boundary context.
# Preserves the heading path ("第一章 / 第二节") on each chunk's metadata
# so retrieval can show WHERE the chunk lives in the source.

import re as _re_rag
import hashlib as _hashlib_rag
import time

_HEADING_RE = _re_rag.compile(r"^(#{1,6})\s+(.+?)\s*$", _re_rag.MULTILINE)
# Sentence terminators — handles both English (.!?) and CJK (。！？).
# We split on the terminator but keep it attached to the left side so
# sentences round-trip correctly.
_SENTENCE_RE = _re_rag.compile(r"(?<=[.!?。！？])\s+")

_CHUNK_OVERLAP_RATIO = 0.15


def _split_by_headings(text: str) -> list[tuple[str, str]]:
    """Split text at Markdown heading boundaries.

    Returns a list of (heading_path, section_body) pairs. Text before
    the first heading gets heading_path="" (preamble). Nested headings
    build a breadcrumb path like "第一章 / 第二节".
    """
    lines = text.splitlines()
    sections: list[tuple[str, str]] = []
    stack: list[tuple[int, str]] = []        # [(level, title), ...]
    current_body: list[str] = []
    current_path = ""

    def _path_of_stack() -> str:
        return " / ".join(t for _l, t in stack)

    for line in lines:
        m = _HEADING_RE.match(line)
        if m:
            # Flush current section.
            body = "\n".join(current_body).strip()
            if body:
                sections.append((current_path, body))
            current_body = []
            level = len(m.group(1))
            title = m.group(2).strip()
            # Pop to level - 1, then push this.
            while stack and stack[-1][0] >= level:
                stack.pop()
            stack.append((level, title))
            current_path = _path_of_stack()
        else:
            current_body.append(line)
    # Flush tail.
    body = "\n".join(current_body).strip()
    if body:
        sections.append((current_path, body))
    # If no headings at all, return single section with empty path.
    if not sections:
        sections = [("", text.strip())] if text.strip() else []
    return sections


def _split_by_paragraphs(text: str) -> list[str]:
    return [p.strip() for p in text.split("\n\n") if p.strip()]


def _split_by_sentences(text: str) -> list[str]:
    parts = _SENTENCE_RE.split(text)
    return [p.strip() for p in parts if p.strip()]


def _force_char_slice(text: str, size: int) -> list[str]:
    """Last-resort slicer for a single unit that's too big."""
    return [text[i:i + size] for i in range(0, len(text), size) if text[i:i + size].strip()]


def _recursive_split(body: str, chunk_size: int) -> list[str]:
    """Hierarchical split: paragraphs → sentences → char. Returns
    chunk-body strings sized at or below ``chunk_size`` (overlap added
    later by the caller)."""
    body = (body or "").strip()
    if not body:
        return []
    if len(body) <= chunk_size:
        return [body]

    out: list[str] = []
    # Level 1: paragraphs.
    paragraphs = _split_by_paragraphs(body)
    buf = ""
    for para in paragraphs:
        if len(para) > chunk_size:
            # Flush pending buffer first.
            if buf.strip():
                out.append(buf.strip())
                buf = ""
            # Level 2: sentences inside the oversized paragraph.
            sentences = _split_by_sentences(para)
            sbuf = ""
            for sent in sentences:
                if len(sent) > chunk_size:
                    # Flush sentence buffer.
                    if sbuf.strip():
                        out.append(sbuf.strip())
                        sbuf = ""
                    # Level 3: force-slice the giant sentence.
                    for piece in _force_char_slice(sent, chunk_size):
                        out.append(piece.strip())
                elif len(sbuf) + len(sent) + 1 > chunk_size and sbuf:
                    out.append(sbuf.strip())
                    sbuf = sent
                else:
                    sbuf += (" " + sent) if sbuf else sent
            if sbuf.strip():
                out.append(sbuf.strip())
        elif len(buf) + len(para) + 2 > chunk_size and buf:
            out.append(buf.strip())
            buf = para
        else:
            buf += ("\n\n" + para) if buf else para
    if buf.strip():
        out.append(buf.strip())
    return out


def _apply_overlap(pieces: list[str], chunk_size: int,
                   ratio: float = _CHUNK_OVERLAP_RATIO) -> list[str]:
    """Prepend a tail slice of the previous chunk onto each chunk after
    the first, to preserve cross-boundary context during retrieval.
    Overlap size = ratio * chunk_size, capped by the previous chunk's
    actual length."""
    if not pieces or len(pieces) == 1:
        return list(pieces)
    overlap_chars = max(0, int(chunk_size * ratio))
    if overlap_chars == 0:
        return list(pieces)
    out: list[str] = [pieces[0]]
    for i in range(1, len(pieces)):
        prev = pieces[i - 1]
        tail = prev[-overlap_chars:] if len(prev) > overlap_chars else prev
        out.append(f"…{tail}\n\n{pieces[i]}")
    return out


def _chunk_text_for_rag(text: str, base_id: str, base_title: str,
                        tags: list, chunk_size: int,
                        source: str = "domain_import",
                        source_file: str = "") -> list[dict]:
    """Heading-aware recursive chunker with overlap.

    Emits chunk dicts ready for RAG ingest. Each chunk's metadata carries
    (v1-C):
      * content_hash  — SHA-256 of content (enables dedup)
      * heading_path  — breadcrumb of Markdown headings leading into it
      * source_file   — relative path, if caller passed it
      * chunk_index   — ordinal within this document
      * imported_at   — unix timestamp
    """
    text = (text or "").strip()
    if not text:
        return []

    now = time.time()
    sections = _split_by_headings(text)

    chunks: list[dict] = []
    chunk_idx = 0
    for heading_path, body in sections:
        if not body.strip():
            continue
        pieces = _recursive_split(body, chunk_size)
        pieces = _apply_overlap(pieces, chunk_size)
        for piece in pieces:
            if not piece.strip():
                continue
            chunk_idx += 1
            content = piece.strip()
            title_suffix = f" · {heading_path}" if heading_path else ""
            display_title = (
                base_title + title_suffix
                + (f" (Part {chunk_idx})" if chunk_idx > 1 else "")
            ) if chunk_idx > 1 else (base_title + title_suffix)
            content_hash = _hashlib_rag.sha256(
                content.encode("utf-8", errors="replace"),
            ).hexdigest()
            chunks.append({
                "id": f"{base_id}_{chunk_idx:04d}",
                "title": display_title,
                "content": content,
                "tags": list(tags or []),
                "source": source,
                # v1-C metadata fields; _ingest_local flattens these
                # into ChromaDB metadata via the `metadata` subdict.
                "content_hash": content_hash,
                "heading_path": heading_path or "",
                "source_file": source_file or "",
                "chunk_index": chunk_idx,
                "imported_at": now,
            })
    return chunks


@router.post("/domain-kb/import")
async def import_domain_knowledge(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Import content into a domain knowledge base."""
    try:
        from ...rag_provider import get_domain_kb_store, get_rag_registry
        store = get_domain_kb_store()
        kb_id = body.get("kb_id", "")
        kb = store.get(kb_id)
        if not kb:
            raise HTTPException(404, "knowledge base not found")
        raw_content = body.get("content", "")
        title = body.get("title", "Imported")
        tags = body.get("tags", [])
        chunk_size = int(body.get("chunk_size", 1000))
        if not raw_content.strip():
            raise HTTPException(400, "content is required")
        chunks = _chunk_text_for_rag(
            raw_content, base_id=f"dkb_{kb.id}",
            base_title=title, tags=tags, chunk_size=chunk_size,
        )
        if not chunks:
            return {"ok": True, "count": 0, "chunks": 0}
        count = get_rag_registry().ingest(kb.provider_id, kb.collection, chunks)
        store.increment_doc_count(kb_id, len(chunks))
        return {"ok": True, "count": count, "chunks": len(chunks)}
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/domain-kb/import-files")
async def import_files_into_domain_kb(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Browser-side folder import: accepts a LIST of base64-encoded files.

    The portal's new folder picker enumerates local files via
    ``<input webkitdirectory>``, reads each as base64, and batches them
    into this endpoint. No server-side path is required — the folder
    can live on the user's laptop.

    Body::

        {
          "kb_id":      "dkb_abc",
          "files": [
              {"name": "docs/a.md",   "data_base64": "..."},
              {"name": "docs/b.pdf",  "data_base64": "..."}
          ],
          "tags":       ["legal"],
          "chunk_size": 1000,
          "max_file_size_mb": 20
        }

    Returns the same shape as /domain-kb/import-folder so the UI can
    share a single result-renderer.
    """
    import base64 as _base64
    try:
        from ...rag_provider import get_domain_kb_store, get_rag_registry
        store = get_domain_kb_store()
        kb_id = (body.get("kb_id") or "").strip()
        if not kb_id:
            raise HTTPException(400, "kb_id is required")
        kb = store.get(kb_id)
        if not kb:
            raise HTTPException(404, "knowledge base not found")

        files = body.get("files") or []
        if not isinstance(files, list) or not files:
            raise HTTPException(400, "files must be a non-empty list")

        tags = body.get("tags") or []
        if not isinstance(tags, list):
            tags = []
        chunk_size = int(body.get("chunk_size", 1000) or 1000)
        max_file_size_mb = float(body.get("max_file_size_mb", 20) or 20)
        max_bytes = int(max_file_size_mb * 1024 * 1024)

        by_file: list[dict] = []
        skipped: list[dict] = []
        all_chunks: list[dict] = []
        files_imported = 0
        files_scanned = 0

        for f in files:
            if not isinstance(f, dict):
                skipped.append({"name": "", "reason": "not_a_dict"})
                continue
            name = (f.get("name") or "").strip()
            data_b64 = f.get("data_base64") or ""
            if not name:
                skipped.append({"name": "", "reason": "missing_name"})
                continue
            if not data_b64:
                skipped.append({"name": name, "reason": "empty_data"})
                continue
            files_scanned += 1
            try:
                raw = _base64.b64decode(data_b64)
            except Exception as e:
                skipped.append({"name": name,
                                "reason": f"base64_decode_failed: {e}"})
                continue
            if not raw:
                skipped.append({"name": name, "reason": "empty"})
                continue
            if len(raw) > max_bytes:
                skipped.append({
                    "name": name,
                    "reason": f"too_large ({len(raw) // (1024*1024)}MB > "
                              f"{max_file_size_mb}MB)",
                })
                continue
            # Parse to text (reuses folder-import helper).
            text, method = _parse_file_bytes_to_text(raw, name)
            if not text or not text.strip():
                skipped.append({"name": name, "reason": "parse_empty",
                                "method": method})
                continue
            # Title/base_id derived from the client-supplied relative name.
            import hashlib as _hl
            base_id = (f"dkb_{kb.id}_u"
                       + _hl.md5(name.encode()).hexdigest()[:10])

            # ── Image preservation (for in-chat display) ────────────
            # When the source is an image, save the original bytes to
            # disk and stamp ``image_url`` onto every chunk's metadata
            # so retrieval downstream can render the image alongside
            # the OCR text in chat / UI.
            ext = name.rsplit(".", 1)[-1].lower() if "." in name else ""
            image_url = ""
            if ext in ("jpg", "jpeg", "png", "webp", "bmp", "tif", "tiff", "gif"):
                try:
                    from ...paths import data_dir as _resolve_data_dir
                    img_dir = _resolve_data_dir() / "kb_images" / kb.id
                    img_dir.mkdir(parents=True, exist_ok=True)
                    safe_ext = ext if ext != "jpeg" else "jpg"
                    img_filename = f"{base_id}.{safe_ext}"
                    img_path = img_dir / img_filename
                    img_path.write_bytes(raw)
                    image_url = f"/api/portal/domain-kb/{kb.id}/image/{img_filename}"
                    logger.debug("Saved KB image: %s (%d bytes)", img_path, len(raw))
                except Exception as _ie:
                    logger.warning("Failed to save image bytes for %s: %s",
                                   name, _ie)

            chunks = _chunk_text_for_rag(
                text, base_id=base_id, base_title=name,
                tags=tags, chunk_size=chunk_size,
                source=f"domain_import_upload:{name}",
                source_file=name,
            )
            if not chunks:
                skipped.append({"name": name, "reason": "no_chunks"})
                continue
            # Stamp image_url onto every chunk produced from this image.
            # Stored at TOP-LEVEL of the chunk dict (not in a metadata
            # sub-dict) because the ingest layer (rag_provider._ingest_local)
            # promotes specific top-level keys into ChromaDB metadata. A
            # nested ``metadata`` would be silently dropped.
            if image_url:
                mime = f"image/{ext if ext != 'jpg' else 'jpeg'}"
                for c in chunks:
                    c["image_url"] = image_url
                    c["is_image"] = True
                    c["mime_hint"] = mime
            all_chunks.extend(chunks)
            files_imported += 1
            by_file.append({
                "name": name,
                "chunks": len(chunks),
                "method": method,
                "size_bytes": len(raw),
                "ok": True,
            })

        if not all_chunks:
            return {
                "ok": True,
                "kb_id": kb_id,
                "files_scanned": files_scanned,
                "files_imported": 0,
                "files_skipped": len(skipped),
                "chunks_total": 0,
                "ingest_count": 0,
                "by_file": by_file,
                "skipped": skipped,
                "message": "No files ingested — see skipped for reasons.",
            }

        # Dedup by content_hash within THIS request (same file uploaded
        # twice gets one copy). Cross-request dedup kicks in at the
        # ingest layer when we wire v1-C up; for now it's upsert-by-id
        # in ChromaDB.
        seen_hashes: set[str] = set()
        deduped: list[dict] = []
        for c in all_chunks:
            h = c.get("content_hash", "")
            if h and h in seen_hashes:
                continue
            if h:
                seen_hashes.add(h)
            deduped.append(c)

        # ── Run ingest in a thread pool ─────────────────────────────────
        # bge-m3.encode() is sync CPU-heavy (10s+ for 30 chunks on CPU,
        # longer on first call when the model loads to RAM). Calling it
        # directly from this async route blocks the asyncio event loop
        # — every other HTTP request stalls until ingest finishes.
        # Symptom (reproduced 2026-05-04): user uploads 56 files,
        # browser tab freezes for minutes, the portal becomes
        # unreachable, "卡死" complaints in logs.
        # ``run_in_threadpool`` runs the sync call on a worker thread so
        # the event loop keeps serving other requests.
        from starlette.concurrency import run_in_threadpool
        reg = get_rag_registry()
        ingest_count = await run_in_threadpool(
            reg.ingest, kb.provider_id, kb.collection, deduped,
        )
        store.increment_doc_count(kb_id, len(deduped))

        return {
            "ok": True,
            "kb_id": kb_id,
            "files_scanned": files_scanned,
            "files_imported": files_imported,
            "files_skipped": len(skipped),
            "chunks_total": len(deduped),
            "chunks_deduped": len(all_chunks) - len(deduped),
            "ingest_count": ingest_count,
            "by_file": by_file,
            "skipped": skipped,
        }
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/domain-kb/import-folder")
async def import_folder_into_domain_kb(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Bulk-import every file in a server-local folder into a domain KB.

    Body::

        {
          "kb_id":      "dkb_abc",
          "folder":     "/absolute/path/to/docs",
          "recursive":  true,                    # walk subfolders (default true)
          "extensions": ["pdf","docx","md","txt","html"],  # default = _KNOWN_KB_EXTENSIONS
          "tags":       ["legal", "2026"],       # per-chunk tag list
          "chunk_size": 1000,                    # chars per chunk
          "max_files":  500,                     # safety cap (default 500)
          "max_file_size_mb": 20,                # skip anything bigger (default 20MB)
        }

    Returns::

        {
          "ok": true,
          "kb_id": "dkb_abc",
          "files_scanned":  17,
          "files_imported": 15,
          "files_skipped":  2,     // details in `skipped`
          "chunks_total":   312,
          "ingest_count":   312,
          "by_file": [{"path": ..., "chunks": N, "method": "...", "ok": true}, ...],
          "skipped": [{"path": ..., "reason": "..."}, ...],
        }
    """
    import os as _os
    try:
        from ...rag_provider import get_domain_kb_store, get_rag_registry
        store = get_domain_kb_store()

        kb_id = (body.get("kb_id") or "").strip()
        folder = (body.get("folder") or "").strip()
        if not kb_id:
            raise HTTPException(400, "kb_id is required")
        if not folder:
            raise HTTPException(400, "folder is required")
        kb = store.get(kb_id)
        if not kb:
            raise HTTPException(404, "knowledge base not found")
        if not _os.path.isdir(folder):
            raise HTTPException(400, f"folder does not exist or is not a directory: {folder}")

        recursive = bool(body.get("recursive", True))
        # Normalize extensions list → lowercase, strip leading dots.
        exts_raw = body.get("extensions")
        if exts_raw is None:
            exts_set = set(_KNOWN_KB_EXTENSIONS)
        elif isinstance(exts_raw, list) and exts_raw:
            exts_set = {
                str(x).strip().lstrip(".").lower()
                for x in exts_raw if str(x).strip()
            }
        else:
            exts_set = set(_KNOWN_KB_EXTENSIONS)

        tags = body.get("tags") or []
        if not isinstance(tags, list):
            tags = []
        chunk_size = int(body.get("chunk_size", 1000) or 1000)
        max_files = int(body.get("max_files", 500) or 500)
        max_file_size_mb = float(body.get("max_file_size_mb", 20) or 20)
        max_bytes = int(max_file_size_mb * 1024 * 1024)

        # Walk folder.
        targets: list[str] = []
        if recursive:
            for dirpath, _dirs, filenames in _os.walk(folder):
                for fn in filenames:
                    targets.append(_os.path.join(dirpath, fn))
        else:
            for fn in _os.listdir(folder):
                full = _os.path.join(folder, fn)
                if _os.path.isfile(full):
                    targets.append(full)
        targets.sort()

        by_file: list[dict] = []
        skipped: list[dict] = []
        all_chunks: list[dict] = []
        files_imported = 0
        files_scanned = 0

        for path in targets:
            if len(by_file) + len(skipped) >= max_files:
                skipped.append({"path": path, "reason": "max_files_cap_reached"})
                continue
            files_scanned += 1
            ext = path.rsplit(".", 1)[-1].lower() if "." in path else ""
            if exts_set and ext not in exts_set:
                skipped.append({"path": path, "reason": f"extension_{ext}_not_in_allowlist"})
                continue
            try:
                size = _os.path.getsize(path)
            except OSError as e:
                skipped.append({"path": path, "reason": f"stat_failed: {e}"})
                continue
            if size > max_bytes:
                skipped.append({
                    "path": path,
                    "reason": f"too_large ({size // (1024*1024)}MB > {max_file_size_mb}MB)",
                })
                continue
            if size == 0:
                skipped.append({"path": path, "reason": "empty"})
                continue
            try:
                with open(path, "rb") as fh:
                    raw = fh.read()
            except OSError as e:
                skipped.append({"path": path, "reason": f"read_failed: {e}"})
                continue
            # Parse to text.
            file_name = _os.path.basename(path)
            text, method = _parse_file_bytes_to_text(raw, file_name)
            if not text or not text.strip():
                skipped.append({"path": path, "reason": "parse_empty",
                                "method": method})
                continue
            # Title: relative to folder root for readability; file_name fallback.
            try:
                rel = _os.path.relpath(path, folder)
            except ValueError:
                rel = file_name
            base_title = rel
            # base_id: path-safe, derived from the relpath hash so chunks
            # are globally unique within this KB.
            import hashlib as _hashlib
            base_id = f"dkb_{kb.id}_f{_hashlib.md5(rel.encode()).hexdigest()[:10]}"
            chunks = _chunk_text_for_rag(
                text, base_id=base_id, base_title=base_title,
                tags=tags, chunk_size=chunk_size,
                source=f"domain_import_folder:{rel}",
            )
            if not chunks:
                skipped.append({"path": path, "reason": "no_chunks_after_split"})
                continue
            all_chunks.extend(chunks)
            files_imported += 1
            by_file.append({
                "path": path,
                "relative_path": rel,
                "chunks": len(chunks),
                "method": method,
                "size_bytes": size,
                "ok": True,
            })

        if not all_chunks:
            return {
                "ok": True,
                "kb_id": kb_id,
                "files_scanned": files_scanned,
                "files_imported": 0,
                "files_skipped": len(skipped),
                "chunks_total": 0,
                "ingest_count": 0,
                "by_file": by_file,
                "skipped": skipped,
                "message": "No files ingested — see skipped for reasons.",
            }

        # Bulk ingest — one call for all chunks keeps RAG provider happy.
        ingest_count = get_rag_registry().ingest(
            kb.provider_id, kb.collection, all_chunks,
        )
        store.increment_doc_count(kb_id, len(all_chunks))

        return {
            "ok": True,
            "kb_id": kb_id,
            "files_scanned": files_scanned,
            "files_imported": files_imported,
            "files_skipped": len(skipped),
            "chunks_total": len(all_chunks),
            "ingest_count": ingest_count,
            "by_file": by_file,
            "skipped": skipped,
        }
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# RAG collections
# ---------------------------------------------------------------------------

@router.post("/rag/collections")
async def list_rag_collections(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """List collections on a RAG provider."""
    try:
        from ...rag_provider import get_rag_registry
        provider_id = body.get("provider_id", "")
        colls = get_rag_registry().list_collections(provider_id)
        return {"collections": [c.to_dict() for c in colls]}
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/rag/collection/create")
async def create_rag_collection(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Create a new RAG collection."""
    try:
        from ...rag_provider import get_rag_registry
        coll = get_rag_registry().create_collection(
            provider_id=body.get("provider_id", ""),
            collection_name=body.get("name", ""),
            description=body.get("description", ""),
        )
        return coll.to_dict()
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/rag/parse-file")
async def parse_file_for_rag(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Parse an uploaded file (base64) into text for RAG ingestion."""
    import base64
    try:
        file_data = body.get("file_data", "")
        file_name = body.get("file_name", "unknown.txt")
        if not file_data:
            raise HTTPException(400, "file_data is required")
        raw = base64.b64decode(file_data)
        # All format-specific parsing now lives in
        # ``_parse_file_bytes_to_text`` (a single dispatcher that
        # routes to per-format helpers — ``_pdf_to_markdown`` /
        # ``_docx_to_markdown`` / ``_html_to_markdown`` / etc.).
        # Each helper emits markdown so the downstream chunker
        # preserves heading hierarchy uniformly.
        text, method = _parse_file_bytes_to_text(raw, file_name)
        return {"text": text, "length": len(text),
                "method": method, "file_name": file_name}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@router.post("/rag/import")
async def import_rag_content(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Import knowledge from raw text content (split into chunks)."""
    try:
        from ...rag_provider import get_rag_registry
        raw_content = body.get("content", "")
        title = body.get("title", "Imported Document")
        collection = body.get("collection", "knowledge")
        provider_id = body.get("provider_id", "")
        tags = body.get("tags", [])
        chunk_size = int(body.get("chunk_size", 1000))
        if not raw_content.strip():
            raise HTTPException(400, "content is required")
        text = raw_content.strip()
        chunks = []
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        current_chunk = ""
        chunk_idx = 0
        for para in paragraphs:
            if len(current_chunk) + len(para) > chunk_size and current_chunk:
                chunk_idx += 1
                chunks.append({
                    "id": f"import_{hash(title) % 100000:05d}_{chunk_idx:03d}",
                    "title": f"{title} (Part {chunk_idx})",
                    "content": current_chunk.strip(),
                    "tags": tags,
                    "source": "import",
                })
                current_chunk = para
            else:
                current_chunk += ("\n\n" + para) if current_chunk else para
        if current_chunk.strip():
            chunk_idx += 1
            chunks.append({
                "id": f"import_{hash(title) % 100000:05d}_{chunk_idx:03d}",
                "title": f"{title} (Part {chunk_idx})" if chunk_idx > 1 else title,
                "content": current_chunk.strip(),
                "tags": tags,
                "source": "import",
            })
        count = get_rag_registry().ingest(provider_id, collection, chunks)
        if collection == "knowledge" and not provider_id:
            try:
                from ...core.memory import get_knowledge_manager
                knowledge = get_knowledge_manager()
                for chunk in chunks:
                    knowledge.add_entry(chunk["title"], chunk["content"], tags)
            except (ImportError, Exception):
                pass
        return {"ok": True, "count": count, "chunks": len(chunks)}
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(501, "RAG provider module not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ---------------------------------------------------------------------------
# Vector memory management
# ---------------------------------------------------------------------------

@router.post("/vector/manage")
async def manage_vector_memory(
    body: dict = Body(...),
    hub=Depends(get_hub),
    user: CurrentUser = Depends(get_current_user),
):
    """Vector memory management (migrate, stats)."""
    try:
        action = body.get("action", "")
        if action == "migrate":
            from ...core.memory import get_memory_manager
            mm = get_memory_manager()
            agent_id = body.get("agent_id", None)
            stats = mm.migrate_to_vector(agent_id)
            return stats
        elif action == "stats":
            from ...core.memory import get_memory_manager
            mm = get_memory_manager()
            stats = mm.get_vector_stats()
            return stats
        else:
            raise HTTPException(400, f"Unknown vector action: {action}")
    except HTTPException:
        raise
    except ImportError:
        raise HTTPException(501, "Memory manager not available")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
