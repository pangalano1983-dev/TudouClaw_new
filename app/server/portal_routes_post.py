"""
Portal POST and DELETE route handlers.

Extracted from portal.py to separate POST and DELETE request handling.

Domain-specific handlers have been extracted to app/server/handlers/:
  auth.py      — /api/auth/* endpoints
  config.py    — /api/portal/config, role-presets, policy, approve
  hub_sync.py  — /api/hub/* node sync & config deployment
  channels.py  — /api/portal/channels/*
  scheduler.py — /api/portal/scheduler/*
  providers.py — /api/portal/providers/*
"""
import json
import logging
import os
import re
import threading
import time
import uuid
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse, parse_qs

from ..hub import get_hub, init_hub
from ..auth import get_auth, init_auth, Role
from .. import llm, tools, knowledge
from ..agent import (Agent, AgentStatus, AgentEvent, AgentTask, TaskStatus,
                     ROLE_PRESETS, AgentProfile, MCPServerConfig,
                     ChatTask, ChatTaskStatus, get_chat_task_manager)
from ..enhancement import build_enhancer, AgentEnhancer
from ..scheduler import get_scheduler
from ..mcp.manager import get_mcp_manager
from ..template_library import get_template_library
from ..llm import get_registry
from ..channel import get_router, ChannelType

from .portal_auth import (get_client_ip, get_session_cookie, set_session_cookie,
                           require_auth, get_auth_info, get_admin_context,
                           is_super_admin, get_visible_agents)
from .handlers import PUBLIC_HANDLERS, DOMAIN_HANDLERS

logger = logging.getLogger("tudou.portal")


def is_hub_mode():
    from .portal_server import _portal_mode
    return _portal_mode == "hub"


def _get_custom_presets_path() -> Path:
    """Return path to persisted custom role presets JSON."""
    home = Path.home() / ".tudou_claw"
    home.mkdir(parents=True, exist_ok=True)
    return home / "role_presets.json"


def _save_custom_role_presets():
    """Persist current ROLE_PRESETS to disk so they survive restarts."""
    from ..agent import ROLE_PRESETS, AgentProfile
    from dataclasses import asdict
    out = {}
    for k, v in ROLE_PRESETS.items():
        entry = dict(v)
        prof = entry.get("profile")
        if prof and hasattr(prof, "__dataclass_fields__"):
            entry["profile"] = asdict(prof)
        out[k] = entry
    try:
        fp = _get_custom_presets_path()
        fp.write_text(json.dumps(out, ensure_ascii=False, indent=2), encoding="utf-8")
        logger.debug("Saved %d role presets to %s", len(out), fp)
    except Exception as e:
        logger.error("Failed to save role presets: %s", e)


def _load_custom_role_presets():
    """Load persisted role presets from disk, merging into ROLE_PRESETS."""
    from ..agent import ROLE_PRESETS, AgentProfile
    fp = _get_custom_presets_path()
    if not fp.exists():
        return
    try:
        data = json.loads(fp.read_text(encoding="utf-8"))
        for k, v in data.items():
            prof_data = v.get("profile", {})
            if isinstance(prof_data, dict):
                profile = AgentProfile(
                    personality=prof_data.get("personality", ""),
                    communication_style=prof_data.get("communication_style", ""),
                    expertise=prof_data.get("expertise", []),
                    skills=prof_data.get("skills", []),
                    allowed_tools=prof_data.get("allowed_tools") or None,
                    denied_tools=prof_data.get("denied_tools") or None,
                    auto_approve_tools=prof_data.get("auto_approve_tools") or None,
                )
            else:
                profile = prof_data
            ROLE_PRESETS[k] = {
                "name": v.get("name", k),
                "system_prompt": v.get("system_prompt", ""),
                "profile": profile,
            }
        logger.info("Loaded %d role presets from disk", len(data))
    except Exception as e:
        logger.error("Failed to load role presets: %s", e)


# Load custom presets on module import
_load_custom_role_presets()


def _bridge_standalone_to_agent(hub, st_task, st_reg) -> bool:
    """Mirror a StandaloneTask into the target agent's execution queue.

    Creates an AgentTask on the assignee agent (source='standalone') and
    writes the new agent_task_id back onto the StandaloneTask so both
    sides can be reconciled later. Returns True on successful push.
    """
    if not st_task or not getattr(st_task, "assigned_to", ""):
        return False
    agent = None
    try:
        agent = hub.agents.get(st_task.assigned_to) if hasattr(hub, "agents") else None
    except Exception:
        agent = None
    if not agent:
        logger.info(
            "standalone task %s has no resolvable agent %s — skipping bridge",
            st_task.id, st_task.assigned_to)
        return False
    try:
        _pri_map = {"low": 0, "normal": 0, "high": 1, "urgent": 2}
        pri = _pri_map.get((st_task.priority or "normal").lower(), 0)
        desc = st_task.description or ""
        if st_task.due_hint:
            desc = f"{desc}\n[Due: {st_task.due_hint}]" if desc else f"[Due: {st_task.due_hint}]"
        if st_task.source_meeting_id:
            desc = (desc + f"\n[源自会议: {st_task.source_meeting_id}]").strip()
        agent_task = agent.add_task(
            title=st_task.title or "(untitled standalone task)",
            description=desc,
            priority=pri,
            assigned_by=st_task.created_by or "user",
            source="standalone",
            tags=list(st_task.tags or []) + ["standalone"],
        )
        # Record the mirror id on the standalone task so status sync can happen later.
        try:
            st_reg.update(st_task.id, agent_task_id=agent_task.id)
        except Exception:
            try:
                st_task.agent_task_id = agent_task.id
                st_reg.save()
            except Exception:
                pass
        logger.info("bridged standalone %s → agent %s task %s",
                    st_task.id, agent.name, agent_task.id)
        return True
    except Exception as e:
        logger.warning("failed bridging standalone %s: %s", st_task.id, e)
        return False


def handle_post(handler):
    """Main POST dispatcher."""
    path = urlparse(handler.path).path
    try:
        _do_post_inner(handler, path)
    except Exception as e:
        import traceback as _tb
        logger.error("POST %s error: %s\n%s", path, e, _tb.format_exc())
        try:
            handler._json({"error": f"{type(e).__name__}: {e}"}, status=500)
        except Exception as e2:
            logger.error("Failed to send error response for POST %s: %s", path, e2)


def _do_post_inner(handler, path: str):
    """Process POST requests."""
    client_ip = get_client_ip(handler)
    logger.info("POST %s from %s", path, client_ip)
    hub = get_hub()
    body = handler._read_body()

    # ── Phase 1: public (pre-auth) handlers ──
    auth_mod = get_auth()
    for h in PUBLIC_HANDLERS:
        fn = getattr(h, "try_handle_public", None)
        if fn and fn(handler, path, hub, body, auth_mod):
            return

    # Token reset (localhost only, for recovery) — LEGACY, now in handlers/auth.py
    from ..defaults import LOCAL_ADDRESSES as _LOCAL_ADDR, MAX_CONTENT_UPLOAD, MAX_DATA_UPLOAD
    if path == "/api/auth/reset-token":
        if client_ip not in _LOCAL_ADDR:
            # Also allow local network IPs
            pass  # Allow all for now, but log it
        auth = get_auth()
        import secrets as _secrets
        raw = _secrets.token_hex(24)
        auth._create_token_obj("admin", "admin", raw)
        # Save to file
        token_file = os.path.join(auth._data_dir, ".admin_token")
        try:
            with open(token_file, "w") as f:
                f.write(raw)
            os.chmod(token_file, 0o600)
        except OSError:
            pass
        logger.warning("Token reset by %s — new admin token created", client_ip)
        handler._json({"ok": True, "token": raw})
        return

    # Auth endpoints (public)
    if path == "/api/auth/login":
        auth = get_auth()
        ip = get_client_ip(handler)

        # Check for admin login (username/password)
        username = body.get("username", "").strip()
        password = body.get("password", "").strip()
        if username and password:
            session = auth.login_admin(username, password, ip=ip)
            if session:
                handler._json({
                    "ok": True,
                    "session_id": session.session_id,
                    "role": session.role,
                    "username": session.name,
                    "admin_user_id": session.admin_user_id,
                })
                auth.audit("login", actor=session.name, role=session.role,
                           ip=ip, success=True)
            else:
                handler._json({"error": "Invalid admin credentials"}, 401)
                auth.audit("login", actor=username, role="",
                           ip=ip, success=False)
            return

        # Token login (existing behavior)
        raw_token = body.get("token", "").strip()
        token_obj = auth.validate_token(raw_token)
        if token_obj:
            session = auth.create_session(token_obj, ip=ip)
            # Return session_id in body so JS can set cookie manually
            # (Chrome blocks Set-Cookie on HTTP non-localhost sites)
            handler._json({"ok": True, "session_id": session.session_id, "role": session.role})
            auth.audit("login", actor=token_obj.name, role=token_obj.role,
                       ip=ip, success=True)
        else:
            handler._json({"error": "Invalid token"}, 401)
            auth.audit("login", actor="unknown", role="",
                       ip=ip, success=False)
        return

    if path == "/api/auth/logout":
        session_id = get_session_cookie(handler)
        auth = get_auth()
        auth.invalidate_session(session_id)
        handler.send_response(200)
        handler.send_header("Content-Type", "application/json")
        # Clear cookie with same security attributes as set_session_cookie
        import os
        is_secure = os.environ.get("TUDOU_SECURE_COOKIES", "").lower() in ("true", "1", "yes")
        secure_flag = "; Secure" if is_secure else ""
        handler.send_header("Set-Cookie", f"td_sess=; Path=/; Max-Age=0; SameSite=Lax; HttpOnly{secure_flag}")
        handler.send_header("Content-Length", "14")
        handler.end_headers()
        handler.wfile.write(b'{"ok": true}')
        return

    # Require auth for other endpoints
    if not require_auth(handler, ):
        return

    # Get auth info for logging
    auth = get_auth()
    actor_name, user_role = get_auth_info(handler)

    # ---- Role & mode enforcement ----
    # In "node" mode, block admin-only operations. These are all
    # global-write operations — node-level writes are still allowed.
    _admin_only_paths = {
        "/api/portal/config",          # Global config changes
        "/api/hub/broadcast",          # Broadcast to all agents
        "/api/hub/orchestrate",        # Multi-agent orchestration
        "/api/auth/tokens",            # Token management
        "/api/portal/providers",       # Global provider management
        "/api/portal/channels",        # Global channel management
        "/api/portal/projects",        # Global project management
        "/api/portal/templates",       # Global template library
        "/api/portal/personas",        # Global persona library
        "/api/portal/approvals",       # Global approval policies
        "/api/portal/workflows",       # Global workflows
    }
    _admin_only_prefixes = (
        "/api/hub/dispatch-config",    # Cross-node config push
        "/api/hub/batch-dispatch-config",
        "/api/auth/tokens/",           # Token delete/update
        "/api/portal/providers/",      # Provider delete/update
        "/api/portal/channels/",       # Channel delete/update (except webhooks)
        "/api/portal/templates/",      # Template delete/update
        "/api/portal/personas/",       # Persona delete/update
    )
    if not is_hub_mode():
        # Node mode: block cross-node agent creation
        if path == "/api/portal/agent/create":
            target = body.get("node_id", "local")
            if target and target != "local" and target != hub.node_id:
                handler._json({"error": "Cross-node agent creation is only available on the Hub portal"}, 403)
                return
        # Node mode: block global config changes
        if path in _admin_only_paths or any(path.startswith(p) for p in _admin_only_prefixes):
            handler._json({"error": "This operation requires Hub admin access"}, 403)
            return

    # Admin role enforcement (even on Hub, non-admins can't do these)
    role_obj = Role(user_role) if user_role in ("admin", "operator", "viewer") else Role.VIEWER
    if path == "/api/portal/config" and not role_obj.can("manage_config"):
        handler._json({"error": "Admin role required for config changes"}, 403)
        return
    if path == "/api/auth/tokens" and not role_obj.can("manage_tokens"):
        handler._json({"error": "Admin role required for token management"}, 403)
        return

    # ── Phase 2: domain-specific handlers (post-auth) ──
    for h in DOMAIN_HANDLERS:
        if h.try_handle(handler, path, hub, body, auth, actor_name, user_role):
            return

    if path == "/api/portal/agent/create":
        target_node = body.get("node_id", "local")
        logger.info("CREATE_AGENT request: name=%s role=%s target_node=%s actor=%s",
                    body.get("name"), body.get("role"), target_node, actor_name)
        if target_node and target_node != "local" and target_node != hub.node_id:
            # Create agent on remote node
            node = hub.remote_nodes.get(target_node)
            logger.info("CREATE_AGENT remote: node_id=%s found=%s url=%s",
                        target_node, node is not None, node.url if node else "N/A")
            if node and node.url:
                try:
                    headers = {"Content-Type": "application/json"}
                    if node.secret:
                        headers["X-Claw-Secret"] = node.secret
                    import requests as http_req
                    # Forward create request to remote node (as local create)
                    remote_body = dict(body)
                    remote_body["node_id"] = "local"  # Tell remote to create locally
                    target_url = f"{node.url}/api/portal/agent/create"
                    logger.info("CREATE_AGENT -> POST %s body=%s has_secret=%s",
                                target_url, {k: v for k, v in remote_body.items() if k != "system_prompt"},
                                bool(node.secret))
                    resp = http_req.post(target_url,
                        headers=headers, json=remote_body, timeout=15)
                    logger.info("CREATE_AGENT <- status=%s length=%s",
                                resp.status_code, len(resp.text))
                    if resp.status_code != 200:
                        err_msg = resp.text[:500]
                        logger.error("CREATE_AGENT REMOTE FAIL: status=%s body=%s",
                                     resp.status_code, err_msg)
                        handler._json({"error": f"Remote node returned {resp.status_code}: {err_msg}"}, resp.status_code)
                        return
                    data = resp.json()
                    logger.info("CREATE_AGENT remote OK: agent_id=%s", data.get("id", "?"))
                    hub.refresh_node(target_node)
                    auth.audit("create_remote_agent", actor=actor_name,
                               role=user_role, target=target_node,
                               ip=get_client_ip(handler))
                    handler._json(data)
                except Exception as e:
                    logger.exception("CREATE_AGENT remote EXCEPTION: %s", e)
                    handler._json({"error": f"Remote create failed: {e}"}, 500)
            else:
                logger.error("CREATE_AGENT node not found: %s available=%s",
                             target_node, list(hub.remote_nodes.keys()))
                handler._json({"error": "Node not found"}, 404)
            return

        # Local agent creation
        logger.info("CREATE_AGENT local: name=%s role=%s model=%s provider=%s priority=%s",
                    body.get("name"), body.get("role"), body.get("model"), body.get("provider"),
                    body.get("priority_level", 3))
        try:
            agent = hub.create_agent(
                name=body.get("name", ""),
                role=body.get("role", "general"),
                model=body.get("model", ""),
                provider=body.get("provider", ""),
                working_dir=body.get("working_dir", ""),
                system_prompt=body.get("system_prompt", ""),
                priority_level=int(body.get("priority_level", 3)),
                role_title=body.get("role_title", ""),
            )
        except ValueError as e:
            handler._json({"error": str(e)}, 400)
            return
        if agent and body.get("profile"):
            prof = body["profile"]
            # Preserve previously-set fields we don't get from the form
            existing = agent.profile
            agent.profile = AgentProfile(
                agent_class=prof.get("agent_class", "") or existing.agent_class,
                memory_mode=prof.get("memory_mode", "") or existing.memory_mode,
                rag_mode=prof.get("rag_mode", "") or existing.rag_mode,
                rag_provider_id=prof.get("rag_provider_id", "") or existing.rag_provider_id,
                rag_collection_ids=prof.get("rag_collection_ids", []) or list(existing.rag_collection_ids),
                personality=prof.get("personality", "") or existing.personality,
                communication_style=prof.get("communication_style", "") or existing.communication_style,
                expertise=prof.get("expertise", []) or list(existing.expertise),
                skills=prof.get("skills", []) or list(existing.skills),
                language=prof.get("language", "auto") or existing.language,
                custom_instructions=prof.get("custom_instructions", "") or existing.custom_instructions,
                max_context_messages=int(prof.get("max_context_messages", existing.max_context_messages) or existing.max_context_messages),
                temperature=float(prof.get("temperature", existing.temperature) or existing.temperature),
                exec_policy=prof.get("exec_policy", "") or existing.exec_policy,
                allowed_tools=list(existing.allowed_tools),
                denied_tools=list(existing.denied_tools),
                auto_approve_tools=list(existing.auto_approve_tools),
                mcp_servers=list(existing.mcp_servers),
            )
            hub._save_agents()  # re-save with profile
        # Set robot avatar if specified
        if agent and body.get("robot_avatar"):
            agent.robot_avatar = body["robot_avatar"]
            hub._save_agents()
        # Apply persona template if user selected one — this overrides
        # personality/communication_style/expertise/skills/temperature +
        # system_prompt with the persona's curated values.
        if agent and body.get("persona_id"):
            try:
                from ..persona import apply_persona_to_agent
                apply_persona_to_agent(agent, body["persona_id"])
                hub._save_agents()
                logger.info("CREATE_AGENT persona applied: agent=%s persona=%s",
                            agent.id, body["persona_id"])
            except Exception as e:
                logger.warning("persona apply failed: %s", e)
        logger.info("CREATE_AGENT local OK: id=%s name=%s",
                    agent.id if agent else "NONE", agent.name if agent else "")
        auth.audit("create_agent", actor=actor_name, role=user_role, target=agent.id if agent else "unknown", ip=get_client_ip(handler))
        handler._json(agent.to_dict() if agent else {})

    elif path.startswith("/api/portal/agent/") and path.endswith("/chat"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            # Try proxy to remote node
            data = hub.proxy_remote_agent_post(agent_id, "/chat", body)
            if data:
                handler._json(data)
            else:
                handler._json({"error": "Agent not found (local or remote)"}, 404)
            return
        user_msg = body.get("message", "").strip()
        attachments = body.get("attachments") or []

        # ── Handle attachments: build multimodal content + save files ──
        saved_refs: list[str] = []
        multimodal_parts: list[dict] = []  # OpenAI-style content parts
        if isinstance(attachments, list) and attachments:
            import base64 as _b64
            import os as _os
            for att in attachments[:10]:
                if not isinstance(att, dict):
                    continue
                raw_name = str(att.get("name") or "attachment.bin")
                safe_name = "".join(
                    c for c in raw_name if c.isalnum() or c in "._-"
                ) or "attachment.bin"
                data_b64 = att.get("data_base64") or ""
                mime_type = str(att.get("mime") or "application/octet-stream")
                if not data_b64:
                    continue
                try:
                    data_bytes = _b64.b64decode(data_b64)
                except Exception:
                    continue
                if len(data_bytes) > MAX_DATA_UPLOAD:
                    continue
                # Build multimodal content FIRST (independent of file save)
                if mime_type.startswith("image/"):
                    multimodal_parts.append({
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{data_b64}",
                        }
                    })
                    # Save image to disk (best-effort)
                    try:
                        base_dir = agent.working_dir or str(agent._effective_working_dir())
                        att_dir = _os.path.join(base_dir, "attachments")
                        _os.makedirs(att_dir, exist_ok=True)
                        ts = int(time.time() * 1000)
                        fname = f"{ts}_{safe_name}"
                        fpath = _os.path.join(att_dir, fname)
                        with open(fpath, "wb") as _f:
                            _f.write(data_bytes)
                        saved_refs.append(fname)
                    except Exception as _ae:
                        logger.warning("agent attachment save failed: %s", _ae)
                else:
                    # Save to disk first so we can extract text
                    fpath = None
                    try:
                        base_dir = agent.working_dir or str(agent._effective_working_dir())
                        att_dir = _os.path.join(base_dir, "attachments")
                        _os.makedirs(att_dir, exist_ok=True)
                        ts = int(time.time() * 1000)
                        fname = f"{ts}_{safe_name}"
                        fpath = _os.path.join(att_dir, fname)
                        with open(fpath, "wb") as _f:
                            _f.write(data_bytes)
                        saved_refs.append(fname)
                    except Exception as _ae:
                        logger.warning("agent attachment save failed: %s", _ae)
                    # Try to extract text content for the LLM
                    extracted = ""
                    if fpath:
                        try:
                            from app.utils.file_parser import extract_file_text
                            extracted = extract_file_text(fpath, mime_type)
                        except Exception:
                            pass
                    if extracted:
                        multimodal_parts.append({
                            "type": "text",
                            "text": f"[File: {safe_name}]\n{extracted}",
                        })
                    else:
                        multimodal_parts.append({
                            "type": "text",
                            "text": f"[Attached file: {safe_name} ({mime_type})]",
                        })
                    continue

        logger.info(
            "agent chat: msg_len=%d attachments=%d saved=%d multimodal_parts=%d",
            len(user_msg), len(attachments), len(saved_refs), len(multimodal_parts),
        )

        if not user_msg and not saved_refs:
            handler._json({"error": "Empty message"}, 400)
            return

        auth.audit("chat", actor=actor_name, role=user_role, target=agent_id, ip=get_client_ip(handler))

        # Build the message content: multimodal (list) or plain text (str)
        if multimodal_parts:
            # Multimodal message: text + image(s)
            content_parts = []
            if user_msg:
                content_parts.append({"type": "text", "text": user_msg})
            elif saved_refs:
                content_parts.append({"type": "text", "text": "请查看以下附件:"})
            content_parts.extend(multimodal_parts)
            chat_content = content_parts  # list format = multimodal
        else:
            # Plain text, optionally with file references
            chat_content = user_msg
            if saved_refs:
                suffix = "\n" + " ".join(f"📎{r}" for r in saved_refs)
                chat_content = (chat_content + suffix) if chat_content else suffix.lstrip()

        # Submit as background task — return task ID immediately
        logger.info(
            "agent chat DISPATCH: content_type=%s multimodal_parts=%d "
            "content_is_list=%s content_len=%s",
            type(chat_content).__name__, len(multimodal_parts),
            isinstance(chat_content, list),
            len(chat_content) if isinstance(chat_content, list)
            else len(str(chat_content)),
        )
        # Route through supervisor (handles isolated / in-process)
        task = hub.supervisor.chat_async(agent.id, chat_content, source="admin")
        handler._json({
            "task_id": task.id,
            "status": task.status.value,
            "attachments_saved": saved_refs,
        })

    elif "/chat-task/" in path and path.endswith("/abort"):
        # POST /api/portal/chat-task/{task_id}/abort
        parts = path.rstrip("/").split("/")
        task_id = parts[4]
        mgr = get_chat_task_manager()
        task = mgr.get_task(task_id)
        if not task:
            handler._json({"error": "Task not found"}, 404)
            return
        task.abort()
        auth.audit("abort_task", actor=actor_name, role=user_role,
                   target=task_id, ip=get_client_ip(handler))
        handler._json({"ok": True, "status": task.status.value})

    elif path.startswith("/api/portal/agent/") and path.endswith("/save-file"):
        # POST /api/portal/agent/{agent_id}/save-file
        # Save content to a file in the agent's working directory
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        filename = body.get("filename", "").strip()
        content = body.get("content", "")
        if not filename:
            handler._json({"error": "Filename is required"}, 400)
            return
        # Resolve path relative to agent's working directory
        import re as _re
        # Sanitize filename: prevent directory traversal
        filename = filename.replace("\\", "/")
        if ".." in filename or filename.startswith("/"):
            handler._json({"error": "Invalid filename"}, 400)
            return
        # Never fall back to os.getcwd() — that's the server-process CWD
        # (often the code package directory). Use the agent's private
        # workspace instead so runtime files stay out of the code tree.
        base_dir = agent.working_dir or str(agent._effective_working_dir())
        file_path = os.path.join(base_dir, filename)
        try:
            os.makedirs(os.path.dirname(file_path), exist_ok=True)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)
            auth.audit("save_file", actor=actor_name, role=user_role,
                       target=agent_id, detail=filename,
                       ip=get_client_ip(handler))
            handler._json({"ok": True, "path": file_path, "size": len(content)})
        except Exception as e:
            handler._json({"error": f"Failed to save file: {e}"}, 500)

    # --- src integration: session save/load ---
    elif path.startswith("/api/portal/agent/") and path.endswith("/save-session"):
        agent_id = path.split("/")[4]
        saved = hub.save_agent_session(agent_id)
        if saved:
            handler._json({"ok": True, "path": saved})
        else:
            handler._json({"error": "Agent not found or save failed"}, 404)

    elif path.startswith("/api/portal/agent/") and path.endswith("/load-session"):
        agent_id = path.split("/")[4]
        ok = hub.load_agent_session(agent_id)
        handler._json({"ok": ok})

    # --- src memory engine POST endpoints ---
    elif path.startswith("/api/portal/agent/") and path.endswith("/save-engine"):
        agent_id = path.split("/")[4]
        saved = hub.save_engine_session(agent_id)
        handler._json({"ok": bool(saved), "path": saved})

    elif path.startswith("/api/portal/agent/") and path.endswith("/restore-engine"):
        agent_id = path.split("/")[4]
        ok = hub.restore_engine_session(agent_id)
        handler._json({"ok": ok})

    elif path.startswith("/api/portal/agent/") and path.endswith("/compact-memory"):
        agent_id = path.split("/")[4]
        ok = hub.compact_agent_memory(agent_id)
        handler._json({"ok": ok})

    elif path.startswith("/api/portal/agent/") and path.endswith("/exec-src-tool"):
        agent_id = path.split("/")[4]
        tool_name = body.get("tool", "")
        payload = body.get("payload", "")
        handler._json(hub.execute_src_tool(agent_id, tool_name, payload))

    elif path.startswith("/api/portal/agent/") and path.endswith("/exec-src-command"):
        agent_id = path.split("/")[4]
        cmd_name = body.get("command", "")
        prompt = body.get("prompt", "")
        handler._json(hub.execute_src_command(agent_id, cmd_name, prompt))

    elif path.startswith("/api/portal/agent/") and path.endswith("/wake"):
        # POST /api/portal/agent/{agent_id}/wake — 唤醒 agent，扫描所有项目
        # 中分配给它且尚未完成的任务，逐个 spawn 后台执行。
        agent_id = path.split("/")[4]
        max_tasks = int(body.get("max_tasks", 5) or 5)
        result = hub.wake_up_agent(agent_id, max_tasks=max_tasks)
        auth.audit("wake_agent", actor=actor_name, role=user_role,
                   target=agent_id, ip=get_client_ip(handler))
        handler._json(result)

    elif path.startswith("/api/portal/agent/") and path.endswith("/clear"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if agent:
            agent.clear()
            auth.audit("clear_agent", actor=actor_name, role=user_role, target=agent_id, ip=get_client_ip(handler))
            handler._json({"ok": True})
        else:
            handler._json({"error": "Agent not found"}, 404)

    elif path.startswith("/api/portal/agent/") and path.endswith("/model"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if agent:
            agent.provider = body.get("provider", "")
            agent.model = body.get("model", "")
            hub._save_agents()
            auth.audit("switch_model", actor=actor_name, role=user_role,
                       target=agent_id,
                       detail=f"{agent.provider or 'default'}/{agent.model or 'default'}",
                       ip=get_client_ip(handler))
            handler._json({"ok": True, "provider": agent.provider, "model": agent.model})
        else:
            # Try remote agent
            node = hub.find_agent_node(agent_id)
            if node:
                hub.proxy_update_model(agent_id, node, body.get("provider", ""), body.get("model", ""))
                handler._json({"ok": True})
            else:
                handler._json({"error": "Agent not found"}, 404)

    elif path.startswith("/api/portal/agent/") and path.endswith("/learning-model"):
        # POST /api/portal/agent/{agent_id}/learning-model
        # Body: {provider, model} — sets the cheap/local LLM used for self-growth tasks.
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        agent.learning_provider = (body.get("provider", "") or "").strip()
        agent.learning_model = (body.get("model", "") or "").strip()
        hub._save_agents()
        auth.audit("set_learning_model", actor=actor_name, role=user_role,
                   target=agent_id,
                   detail=f"{agent.learning_provider or 'default'}/{agent.learning_model or 'default'}",
                   ip=get_client_ip(handler))
        handler._json({
            "ok": True,
            "learning_provider": agent.learning_provider,
            "learning_model": agent.learning_model,
        })

    elif path.startswith("/api/portal/agent/") and path.endswith("/growth-task"):
        # POST /api/portal/agent/{agent_id}/growth-task
        # Body: {learning_goal, knowledge_gap, title?}
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        try:
            task = agent.enqueue_growth_task(
                learning_goal=(body.get("learning_goal", "") or "").strip(),
                knowledge_gap=(body.get("knowledge_gap", "") or "").strip(),
                title=(body.get("title", "") or "").strip(),
            )
            hub._save_agents()
            auth.audit("enqueue_growth_task", actor=actor_name, role=user_role,
                       target=agent_id, detail=task.title,
                       ip=get_client_ip(handler))
            handler._json({"ok": True, "task": task.to_dict()})
        except Exception as e:
            handler._json({"error": str(e)}, 500)

    elif path.startswith("/api/portal/agent/") and "/profile" in path:
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if agent:
            try:
                # Update core fields if provided
                if "name" in body and body["name"].strip():
                    agent.name = body["name"].strip()
                if "role" in body:
                    agent.role = body["role"]
                if "working_dir" in body:
                    # Normalize working_dir so the agent's jail root is always
                    # an absolute, user-intended path. Relative strings like
                    # "developer" must be resolved under the agent's default
                    # private workspace — NOT against the server-process CWD
                    # (which SandboxPolicy._resolve_root would otherwise do,
                    # producing confusing double-folder artefacts like
                    # "/.../developer/developer/file.html").
                    _raw_wd = (body.get("working_dir") or "").strip()
                    if not _raw_wd:
                        # Empty → fall back to default private workspace
                        try:
                            _ws = agent._ensure_workspace_layout()
                            agent.working_dir = str(_ws)
                        except Exception:
                            agent.working_dir = ""
                    else:
                        try:
                            _p = os.path.expanduser(_raw_wd)
                            if not os.path.isabs(_p):
                                # Relative path → append under the agent's
                                # default private workspace base.
                                try:
                                    _base = agent._ensure_workspace_layout()
                                except Exception:
                                    _base = None
                                if _base is None:
                                    handler._json({
                                        "error": (
                                            "cannot resolve relative working_dir: "
                                            "default workspace unavailable"
                                        )
                                    }, 400)
                                    return
                                _resolved = (Path(str(_base)) / _p).resolve()
                                try:
                                    _resolved.mkdir(parents=True, exist_ok=True)
                                except OSError:
                                    pass
                                agent.working_dir = str(_resolved)
                            else:
                                agent.working_dir = os.path.abspath(_p)
                        except Exception as _wd_err:
                            handler._json({"error": f"invalid working_dir: {_wd_err}"}, 400)
                            return
                if "provider" in body:
                    agent.provider = body["provider"]
                if "model" in body:
                    agent.model = body["model"]
                # ── 方案甲: learning / multimodal 专用 LLM slot ──
                if "learning_provider" in body:
                    agent.learning_provider = str(body.get("learning_provider") or "")
                if "learning_model" in body:
                    agent.learning_model = str(body.get("learning_model") or "")
                if "multimodal_provider" in body:
                    agent.multimodal_provider = str(body.get("multimodal_provider") or "")
                if "multimodal_model" in body:
                    agent.multimodal_model = str(body.get("multimodal_model") or "")
                # ── 方案乙(a): extra_llms 任意 N 个 LLM slot ──
                if "extra_llms" in body:
                    raw_slots = body.get("extra_llms") or []
                    if not isinstance(raw_slots, list):
                        raw_slots = []
                    cleaned: list[dict] = []
                    for s in raw_slots:
                        if not isinstance(s, dict):
                            continue
                        label = str(s.get("label") or "").strip()
                        # label 不能为空、不能和 purpose 同时缺失
                        if not label:
                            continue
                        cleaned.append({
                            "label": label,
                            "provider": str(s.get("provider") or "").strip(),
                            "model": str(s.get("model") or "").strip(),
                            "purpose": str(s.get("purpose") or "").strip(),
                            "note": str(s.get("note") or "").strip(),
                        })
                    agent.extra_llms = cleaned
                # ── 方案乙(b): auto_route 启发式路由 ──
                if "auto_route" in body:
                    raw_ar = body.get("auto_route") or {}
                    if not isinstance(raw_ar, dict):
                        raw_ar = {}
                    try:
                        _threshold = int(raw_ar.get("complex_threshold_chars", 2000) or 2000)
                    except (TypeError, ValueError):
                        _threshold = 2000
                    agent.auto_route = {
                        "enabled": bool(raw_ar.get("enabled")),
                        "default": str(raw_ar.get("default") or "").strip(),
                        "complex": str(raw_ar.get("complex") or "").strip(),
                        "multimodal": str(raw_ar.get("multimodal") or "").strip(),
                        "complex_threshold_chars": max(1, _threshold),
                    }
                if "department" in body:
                    agent.department = (body.get("department") or "").strip()
                if "robot_avatar" in body:
                    agent.robot_avatar = body["robot_avatar"]
                agent.profile = AgentProfile(
                    agent_class=body.get("agent_class", agent.profile.agent_class),
                    memory_mode=body.get("memory_mode", agent.profile.memory_mode),
                    rag_mode=body.get("rag_mode", agent.profile.rag_mode),
                    rag_provider_id=body.get("rag_provider_id", agent.profile.rag_provider_id),
                    rag_collection_ids=body.get("rag_collection_ids", agent.profile.rag_collection_ids),
                    personality=body.get("personality", agent.profile.personality),
                    communication_style=body.get("communication_style", agent.profile.communication_style),
                    expertise=body.get("expertise", agent.profile.expertise),
                    skills=body.get("skills", agent.profile.skills),
                    language=body.get("language", agent.profile.language),
                    max_context_messages=body.get("max_context_messages", agent.profile.max_context_messages),
                    allowed_tools=body.get("allowed_tools", agent.profile.allowed_tools),
                    denied_tools=body.get("denied_tools", agent.profile.denied_tools),
                    auto_approve_tools=body.get("auto_approve_tools", agent.profile.auto_approve_tools),
                    temperature=body.get("temperature", agent.profile.temperature),
                    custom_instructions=body.get("custom_instructions", agent.profile.custom_instructions),
                    exec_policy=body.get("exec_policy", agent.profile.exec_policy),
                    exec_blacklist=body.get("exec_blacklist", agent.profile.exec_blacklist),
                    exec_whitelist=body.get("exec_whitelist", agent.profile.exec_whitelist),
                )
                hub._save_agents()
                auth.audit("update_agent_profile", actor=actor_name, role=user_role, target=agent_id, ip=get_client_ip(handler))
                handler._json({"ok": True})
            except Exception as e:
                import traceback
                traceback.print_exc()
                handler._json({"error": f"Failed to update profile: {e}"}, 500)
        else:
            handler._json({"error": "Agent not found"}, 404)

    elif path.startswith("/api/portal/mcp/source/"):
        # PUT-like save: body = {"content": "..."}. Admin-only, no worker approval.
        # Accepts builtin MCPs whose command_template points at any
        # ``python -m app.X.Y.Z`` module (new layout app.mcp.builtins.* or
        # legacy flat app.tudou_*_mcp). The resolved file is required to
        # live under ``app/`` as a path-traversal defence.
        from .portal_auth import is_super_admin as _is_super
        from .portal_routes_get import _mcp_source_resolver
        if not _is_super(handler):
            handler._json({"error": "admin only"}, 403)
            return
        mcp_id = path.split("/")[-1]
        from ..mcp.manager import MCP_CATALOG as _CAT, get_mcp_manager as _gmm
        cap = _CAT.get(mcp_id)
        if cap is None:
            handler._json({"error": "unknown mcp_id"}, 404)
            return
        app_dir, _proj_root, _resolve = _mcp_source_resolver()
        info = _resolve(getattr(cap, "command_template", "") or "")
        if info is None:
            handler._json({"error": "mcp has no editable python source"}, 400)
            return
        _dotted, fpath = info
        fpath = os.path.normpath(fpath)
        if not fpath.startswith(app_dir + os.sep):
            handler._json({"error": "forbidden path"}, 403)
            return
        content = body.get("content")
        if not isinstance(content, str):
            handler._json({"error": "content (string) required"}, 400)
            return
        if len(content) > MAX_CONTENT_UPLOAD:
            handler._json({"error": "content too large (>2MB)"}, 413)
            return
        # Syntax-check before writing — reject broken Python so we don't leave
        # the MCP server in an unstartable state.
        try:
            import ast as _ast
            _ast.parse(content)
        except SyntaxError as _se:
            handler._json({
                "error": f"SyntaxError at line {_se.lineno}: {_se.msg}",
                "lineno": _se.lineno,
            }, 400)
            return
        # Write via tempfile + rename for atomicity + keep a .bak of the old copy.
        bak_path = fpath + ".bak"
        try:
            if os.path.isfile(fpath):
                try:
                    with open(fpath, "r", encoding="utf-8") as _rf:
                        old_text = _rf.read()
                    with open(bak_path, "w", encoding="utf-8") as _bf:
                        _bf.write(old_text)
                except Exception as _bkerr:
                    logger.warning("mcp source save: backup failed: %s", _bkerr)
            import tempfile as _tf
            fd, tmp_path = _tf.mkstemp(
                prefix=".mcp_edit_", suffix=".py",
                dir=os.path.dirname(fpath))
            try:
                with os.fdopen(fd, "w", encoding="utf-8") as _wf:
                    _wf.write(content)
                os.replace(tmp_path, fpath)
            except Exception:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
                raise
        except Exception as _wr_err:
            handler._json({"error": f"write failed: {_wr_err}"}, 500)
            return
        # Try to restart the MCP server so the new code is picked up. Best-effort.
        restart_status = "not_running"
        try:
            mcp_mgr = _gmm()
            restart_fn = getattr(mcp_mgr, "restart_mcp", None)
            if callable(restart_fn):
                ok = restart_fn(mcp_id)
                restart_status = "restarted" if ok else "restart_failed"
            else:
                restart_status = "no_restart_api"
        except Exception as _rs_err:
            logger.warning("mcp source save: restart failed: %s", _rs_err)
            restart_status = f"error: {_rs_err}"
        auth.audit("mcp_source_saved", actor=actor_name, role=user_role,
                   target=mcp_id, ip=get_client_ip(handler),
                   detail=f"bytes={len(content)} restart={restart_status}")
        handler._json({
            "ok": True,
            "mcp_id": mcp_id,
            "rel_path": f"app/{mod}.py",
            "bytes": len(content),
            "backup": os.path.basename(bak_path) if os.path.isfile(bak_path) else "",
            "restart": restart_status,
        })

    elif path == "/api/portal/config":
        cfg = llm.get_config()
        for k in ("provider", "model", "ollama_url", "openai_base_url", "openai_api_key", "claude_api_key", "unsloth_base_url", "unsloth_api_key"):
            if k in body and body[k]:
                cfg[k] = body[k]
        # global_system_prompt: allow empty string so users can clear it.
        if "global_system_prompt" in body:
            val = body.get("global_system_prompt")
            if isinstance(val, str):
                cfg["global_system_prompt"] = val
        # Persist to disk so changes survive restart
        llm.save_config()
        auth.audit("update_config", actor=actor_name, role=user_role, target="config", ip=get_client_ip(handler))
        handler._json({"ok": True})

    elif path == "/api/portal/role-presets/update":
        # Create or update a role preset
        # NOTE: ROLE_PRESETS and AgentProfile already imported at module level
        key = (body.get("key") or "").strip()
        if not key:
            handler._json({"error": "key required"}, 400)
            return
        name = body.get("name", key)
        system_prompt = body.get("system_prompt", "")
        prof_data = body.get("profile", {})
        profile = AgentProfile(
            personality=prof_data.get("personality", ""),
            communication_style=prof_data.get("communication_style", ""),
            expertise=prof_data.get("expertise", []),
            skills=prof_data.get("skills", []),
            allowed_tools=prof_data.get("allowed_tools") or [],
            denied_tools=prof_data.get("denied_tools") or [],
            auto_approve_tools=prof_data.get("auto_approve_tools") or [],
        )
        ROLE_PRESETS[key] = {
            "name": name,
            "system_prompt": system_prompt,
            "profile": profile,
        }
        # Persist custom presets to disk
        _save_custom_role_presets()
        auth.audit("update_role_preset", actor=actor_name, role=user_role,
                   target=key, ip=get_client_ip(handler))
        handler._json({"ok": True})

    elif path == "/api/portal/role-presets/delete":
        # ROLE_PRESETS already imported at module level
        key = (body.get("key") or "").strip()
        if not key:
            handler._json({"error": "key required"}, 400)
            return
        if key in ROLE_PRESETS:
            del ROLE_PRESETS[key]
            _save_custom_role_presets()
        auth.audit("delete_role_preset", actor=actor_name, role=user_role,
                   target=key, ip=get_client_ip(handler))
        handler._json({"ok": True})

    elif path == "/api/portal/policy":
        auth.tool_policy.update_policy_config(body)
        auth.audit("update_policy", actor=actor_name, role=user_role,
                   target="tool_policy", ip=get_client_ip(handler))
        handler._json({"ok": True})

    elif path == "/api/portal/approve":
        approval_id = body.get("approval_id", "")
        action = body.get("action", "")
        if not approval_id:
            handler._json({"ok": False, "error": "approval_id required"}, 400)
            return
        ok = False
        scope = body.get("scope", "once")
        if action == "approve":
            ok = auth.tool_policy.approve(approval_id,
                                          decided_by=actor_name,
                                          scope=scope)
        elif action == "deny":
            ok = auth.tool_policy.deny(approval_id, decided_by=actor_name)
        else:
            handler._json({"ok": False, "error": f"unknown action: {action}"}, 400)
            return
        auth.audit("approval_" + action, actor=actor_name,
                   role=user_role, target=approval_id,
                   ip=get_client_ip(handler), success=ok)
        if not ok:
            handler._json({"ok": False,
                        "error": "approval not found or already decided"}, 404)
            return
        handler._json({"ok": True})

    elif path == "/api/auth/tokens":
        token_name = body.get("name", "").strip()
        token_role = body.get("role", "viewer")
        token_admin_uid = body.get("admin_user_id", "")
        if not token_name:
            handler._json({"error": "Token name required"}, 400)
            return
        auth = get_auth()
        token_obj = auth.create_token(token_name, token_role,
                                      admin_user_id=token_admin_uid)
        auth.audit("create_token", actor=actor_name, role=user_role, target=token_name, ip=get_client_ip(handler))
        handler._json({"token": token_obj.to_dict(), "raw_token": token_obj._raw_token, "name": token_name, "role": token_role})

    elif path == "/api/hub/register":
        nid = body.get("node_id") or uuid.uuid4().hex[:8]
        hub.register_node(
            node_id=nid,
            name=body.get("name", "remote"),
            url=body.get("url", ""),
            agents=body.get("agents", []),
        )
        auth.audit("register_node", actor=actor_name, role=user_role, target=nid, ip=get_client_ip(handler))
        handler._json({"ok": True, "node_id": nid})

    elif path == "/api/hub/sync":
        nid = body.get("node_id", "")
        if nid:
            hub.update_node_agents(nid, body.get("agents", []))
        handler._json({"ok": True})

    elif path == "/api/hub/refresh":
        nid = body.get("node_id", "")
        ok = hub.refresh_node(nid) if nid else False
        handler._json({"ok": ok})

    elif path == "/api/hub/message":
        msg = hub.send_message(
            from_agent=body.get("from_agent", "hub"),
            to_agent=body.get("to_agent", ""),
            content=body.get("content", ""),
            msg_type=body.get("msg_type", "task"),
        )
        auth.audit("send_message", get_client_ip(handler), role=user_role, target=body.get("to_agent", ""), success=True)
        handler._json(msg.to_dict())

    elif path == "/api/hub/deliver":
        to_id = body.get("to_agent", "")
        agent = hub.get_agent(to_id)
        if agent:
            content = body.get("content", "")
            from_id = body.get("from_agent", "remote")
            threading.Thread(
                target=agent.delegate, args=(content, from_id), daemon=True
            ).start()
            handler._json({"ok": True})
        else:
            handler._json({"error": "Agent not found"}, 404)

    # ---- Config deployment endpoints ----

    elif path == "/api/hub/dispatch-config":
        # Push config to a specific agent on a specific node
        from ..hub import AgentConfigPayload
        node_id = body.get("node_id", "local")
        agent_id = body.get("agent_id", "")
        config = AgentConfigPayload.from_dict(body.get("config", {}))
        dep = hub.dispatch_config(node_id, agent_id, config)
        auth.audit("dispatch_config", actor=actor_name, role=user_role,
                   target=f"{node_id}/{agent_id}", ip=get_client_ip(handler))
        handler._json(dep.to_dict())

    elif path == "/api/hub/batch-dispatch-config":
        # Push configs to multiple agents across nodes
        configs = body.get("configs", [])
        deps = hub.batch_dispatch_config(configs)
        auth.audit("batch_dispatch_config", actor=actor_name, role=user_role,
                   target=f"{len(deps)} deployments", ip=get_client_ip(handler))
        handler._json({"deployments": [d.to_dict() for d in deps]})

    elif path == "/api/hub/apply-config":
        # Receive config from hub (called BY remote nodes, or by self for local)
        from ..hub import AgentConfigPayload
        deploy_id = body.get("deploy_id", "")
        agent_id = body.get("agent_id", "")
        config = AgentConfigPayload.from_dict(body.get("config", {}))
        ok = hub.apply_config_to_local_agent(agent_id, config)
        # If we have a deploy_id, confirm back to the caller
        handler._json({"ok": ok, "applied": ok,
                     "deploy_id": deploy_id,
                     "agent_id": agent_id})

    elif path == "/api/hub/confirm-config":
        # Remote node confirms config was applied
        deploy_id = body.get("deploy_id", "")
        success = body.get("success", True)
        error = body.get("error", "")
        ok = hub.confirm_config_applied(deploy_id, success, error)
        handler._json({"ok": ok})

    # --- Node-scoped config management ---
    elif re.match(r"^/api/portal/node/([^/]+)/config/sync$", path):
        # POST: Push config to remote node
        nid = re.match(r"^/api/portal/node/([^/]+)/config/sync$", path).group(1)
        if not is_hub_mode():
            handler._json({"error": "Only Hub can sync config to nodes"}, 403)
            return
        if not Role(user_role).can("manage_config"):
            handler._json({"error": "Admin role required"}, 403)
            return
        result = hub.sync_node_config(nid)
        auth.audit("sync_node_config", actor=actor_name, role=user_role,
                   target=nid, ip=get_client_ip(handler))
        handler._json(result)

    elif re.match(r"^/api/portal/node/([^/]+)/config$", path):
        # POST: Set/update or delete a config item for a node
        nid = re.match(r"^/api/portal/node/([^/]+)/config$", path).group(1)
        # Permission check: admin can edit any; node can only edit own
        if not is_hub_mode() and nid != "local":
            handler._json({"error": "Node mode: can only modify own config"}, 403)
            return
        if is_hub_mode() and not Role(user_role).can("manage_config"):
            handler._json({"error": "Admin role required for config management"}, 403)
            return
        action = body.get("action", "set")
        if action == "delete":
            key = body.get("key", "")
            ok = hub.delete_node_config_item(nid, key)
            auth.audit("delete_node_config", actor=actor_name, role=user_role,
                       target=f"{nid}/{key}", ip=get_client_ip(handler))
            handler._json({"ok": ok, "deleted": key})
        else:
            item = hub.set_node_config_item(
                node_id=nid,
                key=body.get("key", ""),
                value=body.get("value", ""),
                description=body.get("description", ""),
                category=body.get("category", "general"),
                is_secret=body.get("is_secret", False),
                created_by=actor_name,
            )
            auth.audit("set_node_config", actor=actor_name, role=user_role,
                       target=f"{nid}/{item.key}", ip=get_client_ip(handler))
            handler._json(item.to_dict(mask=True))

    elif path == "/api/hub/apply-node-config":
        # Receive config push from Hub (called on remote nodes)
        result = hub.apply_received_node_config(body)
        handler._json(result)

    # --- Enhancement module management ---
    elif re.match(r"^/api/portal/agent/([^/]+)/enhancement$", path):
        aid = re.match(r"^/api/portal/agent/([^/]+)/enhancement$", path).group(1)
        agent = hub.get_agent(aid)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        action = body.get("action", "enable")
        if action == "enable":
            # Accept either single domain or list of up to 8 domains
            domains = body.get("domains")
            if domains is None:
                domains = body.get("domain", "general")
            if isinstance(domains, list):
                domains = domains[:8]
            stats = agent.enable_enhancement(domains)
            hub._save_agents()
            label = "+".join(domains) if isinstance(domains, list) else str(domains)
            auth.audit("enable_enhancement", actor=actor_name, role=user_role,
                       target=f"{aid}/{label}", ip=get_client_ip(handler))
            handler._json({"ok": True, "stats": stats})
        elif action == "disable":
            agent.disable_enhancement()
            hub._save_agents()
            auth.audit("disable_enhancement", actor=actor_name, role=user_role,
                       target=aid, ip=get_client_ip(handler))
            handler._json({"ok": True})
        elif action == "add_knowledge":
            if not agent.enhancer:
                handler._json({"error": "Enhancement not enabled"}, 400)
                return
            entry = agent.enhancer.knowledge.add(
                title=body.get("title", ""),
                content=body.get("content", ""),
                category=body.get("category", "general"),
                tags=body.get("tags", []),
                priority=body.get("priority", 0),
                source=actor_name,
            )
            hub._save_agents()
            handler._json(entry.to_dict())
        elif action == "remove_knowledge":
            if not agent.enhancer:
                handler._json({"error": "Enhancement not enabled"}, 400)
                return
            ok = agent.enhancer.knowledge.remove(body.get("entry_id", ""))
            hub._save_agents()
            handler._json({"ok": ok})
        elif action == "add_reasoning_pattern":
            if not agent.enhancer:
                handler._json({"error": "Enhancement not enabled"}, 400)
                return
            pattern = agent.enhancer.reasoning.add_pattern(
                name=body.get("name", ""),
                description=body.get("description", ""),
                trigger_keywords=body.get("trigger_keywords", []),
                steps=body.get("steps", []),
                reflection_prompt=body.get("reflection_prompt", ""),
            )
            hub._save_agents()
            handler._json(pattern.to_dict())
        elif action == "add_memory":
            if not agent.enhancer:
                handler._json({"error": "Enhancement not enabled"}, 400)
                return
            node = agent.enhancer.memory.add(
                title=body.get("title", ""),
                content=body.get("content", ""),
                kind=body.get("kind", "observation"),
                tags=body.get("tags", []),
                importance=body.get("importance", 0.5),
            )
            hub._save_agents()
            handler._json(node.to_dict())
        elif action == "feedback":
            # Learn from user feedback
            if not agent.enhancer:
                handler._json({"error": "Enhancement not enabled"}, 400)
                return
            node = agent.enhancer.learn_from_interaction(
                user_message=body.get("user_message", ""),
                agent_response=body.get("agent_response", ""),
                outcome=body.get("outcome", "success"),
                feedback=body.get("feedback", ""),
            )
            hub._save_agents()
            handler._json({"ok": True, "learned": node.to_dict() if node else None})
        elif action == "remove_reasoning_pattern":
            if not agent.enhancer:
                handler._json({"error": "Enhancement not enabled"}, 400)
                return
            pid = body.get("pattern_id", "")
            ok = pid in agent.enhancer.reasoning.patterns and agent.enhancer.reasoning.patterns.pop(pid, None) is not None
            hub._save_agents()
            handler._json({"ok": ok})
        elif action == "remove_memory":
            if not agent.enhancer:
                handler._json({"error": "Enhancement not enabled"}, 400)
                return
            nid = body.get("node_id", "")
            ok = nid in agent.enhancer.memory.nodes and agent.enhancer.memory.nodes.pop(nid, None) is not None
            hub._save_agents()
            handler._json({"ok": ok})
        else:
            handler._json({"error": f"Unknown action: {action}"}, 400)

    elif path.startswith("/api/portal/agent/") and path.endswith("/tasks"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            # Try proxy to remote node
            data = hub.proxy_remote_agent_post(agent_id, "/tasks", body)
            if data:
                handler._json(data)
            else:
                handler._json({"error": "Agent not found (local or remote)"}, 404)
            return
        action = body.get("action", "create")
        if action == "create":
            task = agent.add_task(
                title=body.get("title", ""),
                description=body.get("description", ""),
                priority=body.get("priority", 0),
                parent_id=body.get("parent_id", ""),
                assigned_by=actor_name,
                source=body.get("source", "admin"),
                source_agent_id=body.get("source_agent_id", ""),
                deadline=body.get("deadline", 0.0),
                tags=body.get("tags", []),
            )
            logger.info("TASK API create: agent=%s task=%s source=%s deadline=%s",
                        agent_id, task.id, task.source, task.deadline_str or "none")
            handler._json(task.to_dict())
        elif action == "update":
            task_id = body.get("task_id", "")
            updates = {}
            for k in ("title", "description", "status", "priority", "result", "tags", "deadline"):
                if k in body:
                    updates[k] = body[k]
            task = agent.update_task(task_id, **updates)
            if task:
                handler._json(task.to_dict())
            else:
                handler._json({"error": "Task not found"}, 404)
        elif action == "delete":
            ok = agent.remove_task(body.get("task_id", ""))
            handler._json({"ok": ok})
        else:
            handler._json({"error": f"Unknown action: {action}"}, 400)

    # --- Skill System POST ---
    elif path.startswith("/api/portal/agent/") and path.endswith("/prompt-packs"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        action = body.get("action", "")
        from ..core.prompt_enhancer import get_prompt_pack_registry, PromptPack
        registry = get_prompt_pack_registry()
        if action == "bind":
            skill_id = body.get("skill_id", "")
            if skill_id and skill_id not in agent.bound_prompt_packs:
                agent.bound_prompt_packs.append(skill_id)
                hub._save_agents()
            handler._json({"ok": True, "bound_prompt_packs": agent.bound_prompt_packs})
        elif action == "unbind":
            skill_id = body.get("skill_id", "")
            if skill_id in agent.bound_prompt_packs:
                agent.bound_prompt_packs.remove(skill_id)
                hub._save_agents()
            handler._json({"ok": True, "bound_prompt_packs": agent.bound_prompt_packs})
        elif action == "discover":
            import os
            scan_dirs = body.get("scan_dirs", [])
            # Auto-add agent working dir skill paths
            if agent.working_dir:
                for sub in [".claw/skills", ".claude/skills", "skills"]:
                    d = os.path.join(agent.working_dir, sub)
                    if os.path.isdir(d) and d not in scan_dirs:
                        scan_dirs.append(d)
            # Also add global default paths
            home = os.path.expanduser("~")
            for d in [os.path.join(home, ".tudou_claw", "skills"),
                      os.path.join(os.getcwd(), "skills"),
                      os.path.join(os.getcwd(), ".claw", "skills")]:
                if os.path.isdir(d) and d not in scan_dirs:
                    scan_dirs.append(d)
            new_count = registry.discover(scan_dirs if scan_dirs else None)
            handler._json({"ok": True, "new_skills": new_count,
                        "total": len(registry.store.get_active()),
                        "scan_dirs": registry.store._scan_dirs})
        elif action == "import_from_catalog":
            import json
            from pathlib import Path
            skill_ids = body.get("skill_ids", [])
            catalog_path = Path(__file__).resolve().parent.parent / "data" / "community_skills.json"
            try:
                with open(catalog_path, 'r', encoding='utf-8') as f:
                    catalog = json.load(f)
                imported_count = 0
                for skill_id in skill_ids:
                    skill_entry = None
                    for skill in catalog.get("skills", []):
                        if skill.get("id") == skill_id:
                            skill_entry = skill
                            break
                    if skill_entry:
                        record = PromptPack(
                            skill_id=skill_entry.get("id", ""),
                            name=skill_entry.get("name", ""),
                            description=skill_entry.get("description", ""),
                            category=skill_entry.get("category", "general"),
                            tags=skill_entry.get("tags", []),
                            content="",
                            origin="catalog"
                        )
                        registry.store.add_skill(record)
                        imported_count += 1
                        if skill_id not in agent.bound_prompt_packs:
                            agent.bound_prompt_packs.append(skill_id)
                hub._save_agents()
                handler._json({"ok": True, "imported": imported_count, "bound_prompt_packs": agent.bound_prompt_packs})
            except Exception as e:
                handler._json({"error": str(e)}, 500)
        elif action == "import_local":
            import os
            path = body.get("path", "")
            if not os.path.isdir(path):
                handler._json({"error": "Invalid path or directory not found"}, 400)
                return
            if path not in registry.store._scan_dirs:
                registry.store._scan_dirs.append(path)
            new_count = registry.discover([path])
            handler._json({"ok": True, "new_skills": new_count, "scan_path": path})
        else:
            handler._json({"error": f"Unknown action: {action}"}, 400)

    elif path == "/api/portal/prompt-packs":
        action = body.get("action", "")
        from ..core.prompt_enhancer import get_prompt_pack_registry, PromptPack
        registry = get_prompt_pack_registry()
        if action == "discover":
            scan_dirs = body.get("scan_dirs", [])
            new_count = registry.discover(scan_dirs)
            skills = [s.to_dict() for s in registry.store.get_active()]
            handler._json({"ok": True, "new_skills": new_count, "skills": skills})
        elif action == "create":
            record = PromptPack(
                name=body.get("name", ""),
                description=body.get("description", ""),
                category=body.get("category", "general"),
                tags=body.get("tags", []),
                content=body.get("content", ""),
            )
            sid = registry.store.add_skill(record)
            handler._json({"ok": True, "skill_id": sid, "skill": record.to_dict()})
        elif action == "catalog":
            import json
            from pathlib import Path
            catalog_path = Path(__file__).resolve().parent.parent / "data" / "community_skills.json"
            try:
                with open(catalog_path, 'r', encoding='utf-8') as f:
                    catalog = json.load(f)
                category_filter = body.get("category", "")
                search_query = body.get("search", "").lower()
                page = body.get("page", 1)
                per_page = body.get("per_page", 20)

                # Filter and search
                skills = catalog.get("skills", [])
                if category_filter:
                    skills = [s for s in skills if s.get("category") == category_filter]
                if search_query:
                    skills = [s for s in skills if search_query in s.get("name", "").lower() or search_query in s.get("description", "").lower()]

                # Pagination
                total = len(skills)
                start = (page - 1) * per_page
                end = start + per_page
                paginated = skills[start:end]

                # Return lightweight skill info
                result = []
                for s in paginated:
                    result.append({
                        "id": s.get("id", ""),
                        "name": s.get("name", ""),
                        "description": s.get("description", ""),
                        "icon": s.get("icon", ""),
                        "category": s.get("category", ""),
                    })

                handler._json({
                    "ok": True,
                    "skills": result,
                    "total": total,
                    "page": page,
                    "per_page": per_page,
                    "categories": catalog.get("categories", [])
                })
            except Exception as e:
                handler._json({"error": str(e)}, 500)
        else:
            handler._json({"error": f"Unknown action: {action}"}, 400)

    # --- Role Growth Path POST ---
    elif path.startswith("/api/portal/agent/") and path.endswith("/growth"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        action = body.get("action", "")
        if action == "init":
            gp = agent.ensure_growth_path()
            hub._save_agents()
            handler._json({"ok": True,
                        "growth_path": gp.to_dict() if gp else None})
        elif action == "complete_objective":
            objective_id = body.get("objective_id", "")
            gp = agent.ensure_growth_path()
            if gp and objective_id:
                ok = gp.mark_objective_completed(objective_id)
                advanced = gp.try_advance()
                hub._save_agents()
                handler._json({"ok": ok, "advanced": advanced,
                            "summary": gp.get_summary()})
            else:
                handler._json({"error": "No growth path or missing objective_id"}, 400)
        elif action == "trigger_learning":
            gp = agent.ensure_growth_path()
            if gp:
                obj = gp.get_next_objectives(limit=1)
                if obj:
                    from ..core.role_growth_path import build_learning_task_prompt
                    prompt = build_learning_task_prompt(obj[0], gp.role_name)
                    handler._json({"ok": True,
                                "objective": obj[0].to_dict(),
                                "learning_prompt": prompt})
                else:
                    handler._json({"ok": False,
                                "message": "All objectives in current stage completed"})
            else:
                handler._json({"error": "No growth path for this role"}, 400)
        elif action == "advance":
            gp = agent.ensure_growth_path()
            if gp:
                advanced = gp.try_advance()
                hub._save_agents()
                handler._json({"ok": advanced, "summary": gp.get_summary()})
            else:
                handler._json({"error": "No growth path"}, 400)
        else:
            handler._json({"error": f"Unknown action: {action}"}, 400)

    elif path == "/api/portal/skill-store":
        # Hub-level Skill Store — catalog / install / grant / annotate
        action = body.get("action", "")
        store = getattr(hub, "skill_store", None)
        if store is None:
            handler._json({"error": "skill store not initialized"}, 503)
            return
        try:
            if action == "rescan":
                n = store.scan()
                handler._json({"ok": True, "count": n,
                               "stats": store.stats()})
            elif action == "install":
                entry_id = body.get("entry_id", "")
                who = body.get("installed_by", "portal")
                result = store.install_entry(entry_id, installed_by=who)
                store.scan()  # refresh installed flags
                handler._json({"ok": True, "result": result})
            elif action == "uninstall":
                entry_id = body.get("entry_id", "")
                ok = store.uninstall_entry(entry_id)
                store.scan()
                handler._json({"ok": ok})
            elif action == "grant":
                installed_id = body.get("installed_id", "")
                agent_id = body.get("agent_id", "")
                agent = hub.get_agent(agent_id) if hasattr(hub, "get_agent") else None
                wdir = ""
                sync_result = None
                if agent is not None:
                    wdir = getattr(agent, "working_directory", "") or ""
                    if installed_id not in getattr(agent, "granted_skills", []):
                        try:
                            agent.granted_skills.append(installed_id)
                            hub._save_agents()
                        except Exception as e:
                            logger.warning("Failed to grant skill %s to agent %s: %s", installed_id, agent_id, e)
                    # Sync skill package to agent workspace
                    try:
                        inst = store._registry.get(installed_id) if hasattr(store, "_registry") else None
                        if inst is not None and hasattr(agent, "sync_skill_to_workspace"):
                            sync_result = agent.sync_skill_to_workspace(inst)
                            hub._save_agents()  # persist capability changes
                    except Exception as e:
                        logger.debug("Failed to sync skill %s to workspace for agent %s: %s", installed_id, agent_id, e)
                ok = store.grant(installed_id, agent_id, agent_working_dir=wdir)
                resp = {"ok": ok}
                if sync_result:
                    resp["sync"] = sync_result
                handler._json(resp)
            elif action == "revoke":
                installed_id = body.get("installed_id", "")
                agent_id = body.get("agent_id", "")
                agent = hub.get_agent(agent_id) if hasattr(hub, "get_agent") else None
                wdir = getattr(agent, "working_directory", "") if agent else ""
                if agent is not None and installed_id in getattr(agent, "granted_skills", []):
                    try:
                        agent.granted_skills.remove(installed_id)
                        hub._save_agents()
                    except Exception as e:
                        logger.warning("Failed to revoke skill %s from agent %s: %s", installed_id, agent_id, e)
                # Remove skill package from agent workspace
                if agent is not None and hasattr(agent, "remove_skill_from_workspace"):
                    try:
                        inst = store._registry.get(installed_id) if hasattr(store, "_registry") else None
                        skill_name = ""
                        if inst:
                            skill_name = getattr(getattr(inst, "manifest", None), "name", "") or getattr(inst, "id", "")
                        if skill_name:
                            agent.remove_skill_from_workspace(skill_name)
                            hub._save_agents()  # persist capability removal
                    except Exception:
                        pass
                ok = store.revoke(installed_id, agent_id, agent_working_dir=wdir or "")
                handler._json({"ok": ok})
            elif action == "annotate":
                skill_id = body.get("skill_id", "")
                text = body.get("text", "")
                author = body.get("author", "portal")
                ann = store.annotate(skill_id, text, author=author)
                handler._json({"ok": True, "annotation": ann})
            elif action == "clear_annotation":
                skill_id = body.get("skill_id", "")
                ok = store.clear_annotation(skill_id)
                handler._json({"ok": ok})
            elif action == "set_allowed_sources":
                sources = body.get("sources", []) or []
                store.set_allowed_sources(sources)
                handler._json({"ok": True, "allowed": store.allowed_sources()})

            # ── Universal skill import (local path or URL) ──
            elif action == "import":
                from ..skill_store import import_agent_skill as _import_one
                src_path = (body.get("src_path") or "").strip()
                tier = body.get("tier", "community")
                auto_install = body.get("auto_install", True)
                if not src_path:
                    handler._json({"error": "src_path is required"}, 400)
                    return
                # Resolve catalog dir from the store itself
                catalog_dir = store.catalog_dirs[-1] if store.catalog_dirs else ""
                if not catalog_dir:
                    handler._json({"error": "no catalog dir configured"}, 500)
                    return
                result = _import_one(src_path, catalog_dir, tier=tier)
                if result.get("ok") and auto_install:
                    store.scan()
                    entry = store.get_entry(f"imported/{result['name']}")
                    if entry is None:
                        # fallback: try community/<name>
                        entry = store.get_entry(f"community/{result['name']}")
                    if entry is not None and not entry.installed:
                        try:
                            inst_result = store.install_entry(entry.id,
                                                              installed_by="portal")
                            result["install"] = inst_result
                        except Exception as _ie:
                            result["install_error"] = str(_ie)
                    store.scan()  # refresh
                handler._json(result)

            elif action == "import_bulk":
                from ..skill_store import import_anthropic_skills_bulk as _bulk
                src_root = (body.get("src_root") or "").strip()
                include = body.get("include")  # list or None
                tier = body.get("tier", "community")
                auto_install = body.get("auto_install", True)
                if not src_root:
                    handler._json({"error": "src_root is required"}, 400)
                    return
                catalog_dir = store.catalog_dirs[-1] if store.catalog_dirs else ""
                if not catalog_dir:
                    handler._json({"error": "no catalog dir configured"}, 500)
                    return
                results = _bulk(src_root, catalog_dir,
                                include=include, tier=tier)
                if auto_install:
                    store.scan()
                    for r_ in results:
                        if not r_.get("ok"):
                            continue
                        eid = f"imported/{r_['name']}"
                        entry = store.get_entry(eid)
                        if entry is None:
                            entry = store.get_entry(f"community/{r_['name']}")
                        if entry is not None and not entry.installed:
                            try:
                                inst_r = store.install_entry(entry.id,
                                                              installed_by="portal")
                                r_["install"] = inst_r
                            except Exception as _ie:
                                r_["install_error"] = str(_ie)
                    store.scan()
                handler._json({"ok": True, "results": results})

            elif action == "import_from_url":
                # Download a skill archive (tar.gz / zip) from URL,
                # extract to a temp dir, then import.
                import tempfile, urllib.request, zipfile, tarfile
                url = (body.get("url") or "").strip()
                tier = body.get("tier", "community")
                if not url:
                    handler._json({"error": "url is required"}, 400)
                    return
                catalog_dir = store.catalog_dirs[-1] if store.catalog_dirs else ""
                if not catalog_dir:
                    handler._json({"error": "no catalog dir configured"}, 500)
                    return
                tmpdir = tempfile.mkdtemp(prefix="skill_import_")
                try:
                    archive_path = os.path.join(tmpdir, "skill_archive")
                    urllib.request.urlretrieve(url, archive_path)
                    # Auto-detect format
                    extract_dir = os.path.join(tmpdir, "extracted")
                    os.makedirs(extract_dir, exist_ok=True)
                    if zipfile.is_zipfile(archive_path):
                        with zipfile.ZipFile(archive_path) as zf:
                            zf.extractall(extract_dir)
                    elif tarfile.is_tarfile(archive_path):
                        with tarfile.open(archive_path) as tf:
                            tf.extractall(extract_dir)
                    else:
                        handler._json({"error": "unsupported archive format (need .zip or .tar.gz)"}, 400)
                        return
                    # Find SKILL.md inside extracted dir
                    from ..skill_store import import_agent_skill as _import_one
                    skill_dir = None
                    for _root, _dirs, _files in os.walk(extract_dir):
                        low = {f.lower() for f in _files}
                        if "skill.md" in low or "manifest.yaml" in low:
                            skill_dir = _root
                            break
                    if not skill_dir:
                        handler._json({"error": "no SKILL.md or manifest.yaml in archive"}, 400)
                        return
                    result = _import_one(skill_dir, catalog_dir, tier=tier)
                    if result.get("ok"):
                        store.scan()
                        eid = f"imported/{result['name']}"
                        entry = store.get_entry(eid)
                        if entry and not entry.installed:
                            try:
                                result["install"] = store.install_entry(
                                    entry.id, installed_by="portal")
                            except Exception as _ie:
                                result["install_error"] = str(_ie)
                        store.scan()
                    handler._json(result)
                except Exception as _ue:
                    handler._json({"error": f"URL import failed: {_ue}"}, 500)
                finally:
                    import shutil as _sh
                    _sh.rmtree(tmpdir, ignore_errors=True)

            # ── Remote URL scanning (two-step: scan → import) ──
            elif action == "scan_url":
                from ..skill_store import scan_remote_url as _scan_remote
                url = (body.get("url") or "").strip()
                if not url:
                    handler._json({"error": "url is required"}, 400)
                    return
                result = _scan_remote(url)
                handler._json(result)

            elif action == "import_scanned":
                from ..skill_store import import_from_scan_result as _import_scanned
                temp_dir = (body.get("temp_dir") or "").strip()
                skill_names = body.get("skill_names") or []
                tier = body.get("tier", "community")
                auto_install = body.get("auto_install", True)
                if not temp_dir or not skill_names:
                    handler._json({"error": "temp_dir and skill_names required"}, 400)
                    return
                catalog_dir = store.catalog_dirs[-1] if store.catalog_dirs else ""
                if not catalog_dir:
                    handler._json({"error": "no catalog dir configured"}, 500)
                    return
                results = _import_scanned(temp_dir, skill_names,
                                          catalog_dir, tier=tier)
                if auto_install:
                    store.scan()
                    for r_ in results:
                        if not r_.get("ok"):
                            continue
                        eid = f"imported/{r_['name']}"
                        entry = store.get_entry(eid)
                        if entry is None:
                            entry = store.get_entry(f"community/{r_['name']}")
                        if entry is not None and not entry.installed:
                            try:
                                inst_r = store.install_entry(entry.id,
                                                              installed_by="portal")
                                r_["install"] = inst_r
                            except Exception as _ie:
                                r_["install_error"] = str(_ie)
                    store.scan()
                handler._json({"ok": True, "results": results})

            elif action == "cleanup_scan":
                from ..skill_store import cleanup_scan_temp as _cleanup
                temp_dir = (body.get("temp_dir") or "").strip()
                _cleanup(temp_dir)
                handler._json({"ok": True})

            # ── Skill Creator: validate / preview / edit ──
            elif action == "validate":
                # Validate a SKILL.md body (raw text) for correctness.
                content = body.get("content", "")
                if not content.strip():
                    handler._json({"error": "content is empty"}, 400)
                    return
                from ..skill_store import _parse_frontmatter
                meta, md_body = _parse_frontmatter(content)
                issues: list[str] = []
                if not meta:
                    issues.append("YAML frontmatter not found (need --- ... --- block)")
                else:
                    if not meta.get("name"):
                        issues.append("frontmatter missing required field: name")
                    if not meta.get("description"):
                        issues.append("frontmatter missing required field: description")
                if not md_body.strip():
                    issues.append("body is empty (no instructions after frontmatter)")
                handler._json({
                    "ok": len(issues) == 0,
                    "issues": issues,
                    "parsed_meta": meta if meta else {},
                    "body_length": len(md_body),
                })

            elif action == "preview":
                # Read back an installed skill's SKILL.md content for editing.
                entry_id = body.get("entry_id", "")
                entry = store.get_entry(entry_id)
                if entry is None:
                    handler._json({"error": f"entry not found: {entry_id}"}, 404)
                    return
                import pathlib
                skill_md = pathlib.Path(entry.catalog_path) / "SKILL.md"
                if not skill_md.exists():
                    skill_md = pathlib.Path(entry.catalog_path) / "skill.md"
                content = ""
                files: list[dict] = []
                if skill_md.exists():
                    try:
                        content = skill_md.read_text(encoding="utf-8")
                    except Exception as e:
                        logger.debug("Failed to read skill markdown: %s", e)
                # List ancillary files
                cp = pathlib.Path(entry.catalog_path)
                for fp in sorted(cp.rglob("*")):
                    if fp.is_file():
                        try:
                            rel = str(fp.relative_to(cp))
                        except ValueError:
                            rel = fp.name
                        files.append({
                            "name": rel,
                            "size": fp.stat().st_size,
                        })
                handler._json({
                    "ok": True,
                    "entry": entry.to_dict(),
                    "content": content,
                    "files": files,
                })

            elif action == "save_edit":
                # Save edited SKILL.md content back to the catalog copy.
                entry_id = body.get("entry_id", "")
                content = body.get("content", "")
                entry = store.get_entry(entry_id)
                if entry is None:
                    handler._json({"error": f"entry not found: {entry_id}"}, 404)
                    return
                if not content.strip():
                    handler._json({"error": "content is empty"}, 400)
                    return
                # Validate first
                from ..skill_store import _parse_frontmatter
                meta, _md = _parse_frontmatter(content)
                if not meta or not meta.get("name"):
                    handler._json({"error": "invalid SKILL.md: frontmatter missing name"}, 400)
                    return
                import pathlib
                skill_md = pathlib.Path(entry.catalog_path) / "SKILL.md"
                try:
                    skill_md.chmod(0o644)
                except Exception:
                    pass
                skill_md.write_text(content, encoding="utf-8")
                # If installed, also update the installed copy
                if entry.installed:
                    reg = getattr(hub, "skill_registry", None)
                    if reg:
                        inst = reg.get(entry.installed_id)
                        if inst:
                            import pathlib as _pl
                            installed_md = _pl.Path(inst.install_dir) / "SKILL.md"
                            try:
                                installed_md.chmod(0o644)
                            except Exception:
                                pass
                            installed_md.write_text(content, encoding="utf-8")
                store.scan()
                handler._json({"ok": True, "name": meta.get("name", "")})

            elif action == "create_new":
                # Create a brand-new skill from scratch in the catalog.
                content = body.get("content", "")
                tier = body.get("tier", "local")
                if not content.strip():
                    handler._json({"error": "content is empty"}, 400)
                    return
                from ..skill_store import _parse_frontmatter
                meta, _md = _parse_frontmatter(content)
                if not meta or not meta.get("name"):
                    handler._json({"error": "invalid SKILL.md: frontmatter missing name"}, 400)
                    return
                name = meta["name"]
                catalog_dir = store.catalog_dirs[-1] if store.catalog_dirs else ""
                if not catalog_dir:
                    handler._json({"error": "no catalog dir"}, 500)
                    return
                import pathlib
                target = pathlib.Path(catalog_dir) / "local" / name
                target.mkdir(parents=True, exist_ok=True)
                (target / "SKILL.md").write_text(content, encoding="utf-8")
                store.scan()
                entry = store.get_entry(f"local/{name}") or store.get_entry(f"community/{name}")
                result: dict[str, Any] = {"ok": True, "name": name, "path": str(target)}
                if entry:
                    result["entry"] = entry.to_dict()
                    # Auto-install
                    try:
                        inst_r = store.install_entry(entry.id, installed_by="portal")
                        result["install"] = inst_r
                    except Exception as _ie:
                        result["install_error"] = str(_ie)
                    store.scan()
                handler._json(result)

            else:
                handler._json({"error": f"Unknown action: {action}"}, 400)
        except KeyError as _ke:
            handler._json({"error": str(_ke)}, 404)
        except Exception as _ex:
            handler._json({"error": str(_ex)}, 500)

    elif path == "/api/portal/providers":
        # Add new provider
        reg = get_registry()
        p = reg.add(
            name=body.get("name", ""),
            kind=body.get("kind", "openai"),
            base_url=body.get("base_url", ""),
            api_key=body.get("api_key", ""),
            enabled=body.get("enabled", True),
            manual_models=body.get("manual_models"),
            scope=body.get("scope", "local"),
            max_concurrent=max(1, int(body.get("max_concurrent", 1))),
            schedule_strategy=body.get("schedule_strategy", "serial"),
            rate_limit_rpm=max(0, int(body.get("rate_limit_rpm", 0))),
        )
        # Set models_cache from manual_models
        if body.get("manual_models"):
            p.models_cache = list(body.get("manual_models", []))
            reg._save()
        auth.audit("add_provider", actor=actor_name, role=user_role,
                   target=p.id, ip=get_client_ip(handler))
        handler._json(p.to_dict(mask_key=True))

    elif path.startswith("/api/portal/providers/") and path.endswith("/update"):
        provider_id = path.split("/")[4]
        reg = get_registry()
        kwargs = {}
        for k in ("name", "kind", "base_url", "api_key", "enabled", "manual_models"):
            if k in body:
                # Don't overwrite api_key with mask
                if k == "api_key" and body[k] == "********":
                    continue
                kwargs[k] = body[k]
        p = reg.update(provider_id, **kwargs)
        if p:
            # ── Concurrency & scheduling fields ──
            changed = False
            if "max_concurrent" in body:
                p.max_concurrent = max(1, int(body["max_concurrent"]))
                changed = True
            if "model_concurrency" in body:
                mc = body["model_concurrency"]
                p.model_concurrency = {k: int(v) for k, v in (mc or {}).items() if int(v) > 0}
                changed = True
            if "schedule_strategy" in body:
                p.schedule_strategy = body["schedule_strategy"]
                changed = True
            if "rate_limit_rpm" in body:
                p.rate_limit_rpm = max(0, int(body["rate_limit_rpm"]))
                changed = True
            if "scope" in body:
                p.scope = body["scope"]
                changed = True
            if "priority" in body:
                p.priority = int(body["priority"])
                changed = True
            if "cost_per_1k_tokens" in body:
                p.cost_per_1k_tokens = float(body["cost_per_1k_tokens"])
                changed = True
            if "context_length" in body:
                p.context_length = max(0, int(body["context_length"]))
                changed = True
            # The edit UI shows union of manual_models + models_cache as
            # editable tags. So what the user sends back IS the full model
            # list they want.  Sync models_cache accordingly.
            if "manual_models" in body:
                wanted = list(body.get("manual_models") or [])
                p.manual_models = wanted
                p.models_cache = list(set(wanted))
                changed = True
            if changed:
                reg._save()
            auth.audit("update_provider", actor=actor_name, role=user_role,
                       target=provider_id, ip=get_client_ip(handler))
            handler._json(p.to_dict(mask_key=True))
        else:
            handler._json({"error": "Provider not found"}, 404)

    elif path.startswith("/api/portal/providers/") and path.endswith("/detect"):
        provider_id = path.split("/")[4]
        reg = get_registry()
        models = reg.detect_models(provider_id)
        handler._json({"provider_id": provider_id, "models": models})

    elif path == "/api/portal/providers/detect-all":
        reg = get_registry()
        all_models = reg.detect_all_models()
        handler._json({"models": all_models})

    elif path.startswith("/api/portal/agent/") and path.endswith("/persona"):
        agent_id = path.split("/")[4]
        persona_id = body.get("persona_id", "")
        if not persona_id:
            handler._json({"error": "persona_id required"}, 400)
        else:
            ok = hub.apply_persona(agent_id, persona_id)
            if ok:
                auth.audit("apply_persona", actor=actor_name,
                           role=user_role, target=agent_id,
                           detail=persona_id, ip=get_client_ip(handler))
                handler._json({"ok": True})
            else:
                handler._json({"error": "Agent or persona not found"}, 404)

    elif path.startswith("/api/portal/agent/") and path.endswith("/soul"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        soul_md = body.get("soul_md", "")
        robot_avatar = body.get("robot_avatar", "")
        if soul_md is not None:
            agent.soul_md = soul_md
            # Also update the system_prompt from SOUL.md content
            agent.system_prompt = soul_md
        if robot_avatar is not None:
            agent.robot_avatar = robot_avatar
        # Rebuild system prompt immediately
        if agent.messages and agent.messages[0].get("role") == "system":
            agent.messages[0]["content"] = agent._build_system_prompt()
        hub._save_agents()
        handler._json({"ok": True, "agent_id": agent_id})

    elif path.startswith("/api/portal/agent/") and path.endswith("/thinking/enable"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        config = body or {}
        stats = agent.enable_active_thinking(**config)
        hub._save_agents()
        handler._json({"ok": True, "stats": stats})

    elif path.startswith("/api/portal/agent/") and path.endswith("/thinking/disable"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        agent.disable_active_thinking()
        hub._save_agents()
        handler._json({"ok": True})

    elif path.startswith("/api/portal/agent/") and path.endswith("/thinking/trigger"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        trigger = body.get("trigger", "manual")
        context = body.get("context", "")
        result = agent.trigger_thinking(trigger=trigger, context=context)
        hub._save_agents()
        handler._json({"ok": True, "result": result})

    elif path.startswith("/api/portal/agent/") and path.endswith("/thinking/history"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        history = []
        if agent.active_thinking:
            history = [r.to_dict() for r in agent.active_thinking.history[-20:]]
        stats = agent.active_thinking.get_stats() if agent.active_thinking else None
        handler._json({"history": history, "stats": stats})

    elif path == "/api/portal/agent/workspace/authorize":
        agent_id = body.get("agent_id", "")
        target_agent_id = body.get("target_agent_id", "")  # whose workspace to authorize
        if not agent_id or not target_agent_id:
            handler._json({"error": "agent_id and target_agent_id required"}, 400)
            return
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": f"Agent {agent_id} not found"}, 404)
            return
        if target_agent_id not in agent.authorized_workspaces:
            agent.authorized_workspaces.append(target_agent_id)
            hub._save_agents()
            auth.audit("authorize_workspace", actor=actor_name,
                       role=user_role, target=agent_id,
                       detail=f"authorized:{target_agent_id}", ip=get_client_ip(handler))
        handler._json({"ok": True, "authorized_workspaces": agent.authorized_workspaces})

    elif path == "/api/portal/agent/workspace/revoke":
        agent_id = body.get("agent_id", "")
        target_agent_id = body.get("target_agent_id", "")
        if not agent_id or not target_agent_id:
            handler._json({"error": "agent_id and target_agent_id required"}, 400)
            return
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": f"Agent {agent_id} not found"}, 404)
            return
        if target_agent_id in agent.authorized_workspaces:
            agent.authorized_workspaces.remove(target_agent_id)
            hub._save_agents()
            auth.audit("revoke_workspace", actor=actor_name,
                       role=user_role, target=agent_id,
                       detail=f"revoked:{target_agent_id}", ip=get_client_ip(handler))
        handler._json({"ok": True, "authorized_workspaces": agent.authorized_workspaces})

    elif path == "/api/portal/agent/workspace/list":
        agent_id = body.get("agent_id", "")
        if not agent_id:
            handler._json({"error": "agent_id required"}, 400)
            return
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": f"Agent {agent_id} not found"}, 404)
            return
        handler._json({
            "agent_id": agent_id,
            "own_workspace": agent.working_dir,
            "shared_workspace": agent.shared_workspace,
            "authorized_workspaces": agent.authorized_workspaces,
        })

    elif path == "/api/portal/workflows":
        action = body.get("action", "create")
        if action == "create":
            template_id = body.get("template_id", "")
            input_data = body.get("input_data", "")
            if template_id:
                wf = hub.create_workflow_from_template(template_id, input_data)
            else:
                wf = hub.create_custom_workflow(
                    name=body.get("name", "Custom Workflow"),
                    description=body.get("description", ""),
                    steps=body.get("steps", []),
                    input_data=input_data,
                )
            if wf:
                auth.audit("create_workflow", actor=actor_name,
                           role=user_role, target=wf.id,
                           ip=get_client_ip(handler))
                handler._json(wf.to_dict())
            else:
                handler._json({"error": "Failed to create workflow"}, 400)
        elif action == "create_from_catalog":
            catalog_id = body.get("catalog_id", "")
            custom_name = body.get("name", "")
            if not catalog_id:
                handler._json({"error": "Missing catalog_id"}, 400)
                return
            tmpl = hub.workflow_engine.create_from_catalog(catalog_id, custom_name)
            if tmpl:
                auth.audit("create_workflow_from_catalog", actor=actor_name,
                           role=user_role, target=tmpl.id,
                           detail=f"catalog={catalog_id}",
                           ip=get_client_ip(handler))
                handler._json(tmpl.to_dict())
            else:
                handler._json({"error": f"Catalog template '{catalog_id}' not found"}, 404)
        elif action == "delete":
            wf_id = body.get("workflow_id", "")
            engine = hub.workflow_engine
            # Remove from all dicts
            removed = False
            with engine._lock:
                if wf_id in engine._workflows:
                    engine._workflows.pop(wf_id)
                    removed = True
                if wf_id in engine._templates:
                    engine._templates.pop(wf_id)
                    removed = True
                if wf_id in engine._instances:
                    engine._instances.pop(wf_id)
                    removed = True
            if removed:
                engine.save()
                # Also remove from database
                db = engine._get_db()
                if db:
                    try:
                        db.delete("workflow_templates", "template_id", wf_id)
                        db.delete("workflow_instances", "instance_id", wf_id)
                    except Exception:
                        pass
                auth.audit("delete_workflow", actor=actor_name,
                           role=user_role, target=wf_id,
                           ip=get_client_ip(handler))
                handler._json({"ok": True})
            else:
                handler._json({"error": "Workflow not found"}, 404)
        elif action == "update":
            wf_id = body.get("workflow_id", "")
            engine = hub.workflow_engine
            # Update template
            tmpl = engine._templates.get(wf_id)
            wf = engine._workflows.get(wf_id)
            if not tmpl and not wf:
                handler._json({"error": "Workflow not found"}, 404)
                return
            if tmpl:
                if body.get("name"):
                    tmpl.name = body["name"]
                if body.get("description") is not None:
                    tmpl.description = body["description"]
                if body.get("steps"):
                    from app.workflow import StepTemplate
                    tmpl.steps = [StepTemplate.from_dict(s) for s in body["steps"]]
                engine.save()
                handler._json(tmpl.to_dict())
            elif wf:
                if body.get("name"):
                    wf.name = body["name"]
                if body.get("description") is not None:
                    wf.description = body["description"]
                if body.get("steps"):
                    from app.workflow import StepInstance
                    wf.steps[:] = [StepInstance.from_dict(s) for s in body["steps"]]
                engine.save()
                handler._json(wf.to_dict())
        else:
            handler._json({"error": f"Unknown action: {action}"}, 400)

    elif path.startswith("/api/portal/workflows/") and path.endswith("/start"):
        wf_id = path.split("/")[4]
        ok = hub.start_workflow(wf_id)
        if ok:
            auth.audit("start_workflow", actor=actor_name,
                       role=user_role, target=wf_id,
                       ip=get_client_ip(handler))
            handler._json({"ok": True})
        else:
            handler._json({"error": "Workflow not found or already running"}, 400)

    elif path.startswith("/api/portal/workflows/") and path.endswith("/abort"):
        wf_id = path.split("/")[4]
        ok = hub.abort_workflow(wf_id)
        if ok:
            handler._json({"ok": True})
        else:
            handler._json({"error": "Workflow not found"}, 404)

    # ---- Self-Improvement / Experience Library endpoints ----
    elif path == "/api/portal/experience/retrospective":
        agent_id = body.get("agent_id", "")
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        task_summary = body.get("task_summary", "")
        context = body.get("context", "")
        try:
            result = agent.trigger_retrospective(
                task_summary=task_summary, context=context)
            # Persist experience changes
            hub._save_agents()
            auth.audit("trigger_retrospective", actor=actor_name,
                       role=user_role, target=agent_id,
                       ip=get_client_ip(handler))
            handler._json(result)
        except Exception as e:
            logger.error(f"Retrospective error for {agent_id}: {e}")
            handler._json({"error": f"复盘失败: {e}"}, 500)

    elif path == "/api/portal/experience/learning":
        agent_id = body.get("agent_id", "")
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        learning_goal = body.get("learning_goal", "")
        knowledge_gap = body.get("knowledge_gap", "")
        try:
            result = agent.trigger_active_learning(
                learning_goal=learning_goal, knowledge_gap=knowledge_gap)
            # Persist experience changes
            hub._save_agents()
            auth.audit("trigger_learning", actor=actor_name,
                       role=user_role, target=agent_id,
                       ip=get_client_ip(handler))
            handler._json(result)
        except Exception as e:
            logger.error(f"Active learning error for {agent_id}: {e}")
            handler._json({"error": f"主动学习失败: {e}"}, 500)

    elif path.startswith("/api/portal/agent/") and path.endswith("/self-improvement/enable"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        import_exp = body.get("import_experience", True)
        import_limit = body.get("import_limit", 50)
        result = agent.enable_self_improvement(
            import_experience=import_exp, import_limit=import_limit)
        hub._save_agents()
        handler._json({"ok": True, **result})

    elif path.startswith("/api/portal/agent/") and path.endswith("/self-improvement/disable"):
        agent_id = path.split("/")[4]
        agent = hub.get_agent(agent_id)
        if not agent:
            handler._json({"error": "Agent not found"}, 404)
            return
        agent.disable_self_improvement()
        hub._save_agents()
        handler._json({"ok": True})

    elif path == "/api/portal/projects":
        action = body.get("action", "create")
        if action == "create":
            proj = hub.create_project(
                name=body.get("name", "New Project"),
                description=body.get("description", ""),
                member_configs=body.get("members", []),
                working_directory=body.get("working_directory", ""),
                node_id=body.get("node_id", "local"),
                workflow_id=body.get("workflow_id", ""),
                step_assignments=body.get("step_assignments", []),
            )
            auth.audit("create_project", actor=actor_name,
                       role=user_role, target=proj.id,
                       ip=get_client_ip(handler))
            handler._json(proj.to_dict())
        elif action == "update":
            proj = hub.get_project(body.get("project_id", ""))
            if not proj:
                handler._json({"error": "Project not found"}, 404)
                return
            if body.get("name"):
                proj.name = body["name"]
            if body.get("description") is not None:
                proj.description = body["description"]
            if body.get("status"):
                old_status = str(proj.status)
                # Route through set_status for proper enum conversion +
                # paused flag synchronization.
                ok_s, _msg_s = proj.set_status(
                    body["status"], by=actor_name or "admin", reason="")
                if not ok_s:
                    handler._json({"error": _msg_s}, 400)
                    return
                # Agent 学习闭环: 项目变为 completed 时，触发成员经验沉淀
                new_status = body["status"]
                if (str(new_status) in ("completed", "ProjectStatus.COMPLETED")
                        and "completed" not in old_status.lower()):
                    for m in proj.members:
                        agent = hub.get_agent(m.agent_id)
                        if not agent:
                            continue
                        try:
                            consolidator = agent._get_memory_consolidator()
                            if consolidator:
                                consolidator.consolidate(
                                    agent_id=m.agent_id, force=True)
                                agent.history_log.add(
                                    "project_complete_learning",
                                    f"[Learning] 项目 {proj.name} 完成，经验已沉淀")
                        except Exception as _e:
                            logger.debug("Project-complete consolidate failed for %s: %s",
                                         m.agent_id, _e)
            # If working_directory is updated, re-sync all member agents
            if body.get("working_directory"):
                proj.working_directory = body["working_directory"]
                for member in proj.members:
                    hub._sync_agent_to_project_dir(
                        member.agent_id, proj.working_directory,
                        project_id=proj.id, project_name=proj.name)
                hub._save_agents()
            # Workflow 绑定/更换
            if body.get("workflow_id") is not None:
                wf_id = body["workflow_id"]
                step_asgn = body.get("step_assignments", [])
                if wf_id:
                    # 清除旧的 WF 任务后重新绑定
                    proj.tasks = [t for t in proj.tasks
                                  if not t.title.startswith("[WF Step")]
                    hub._bind_workflow_to_project(
                        proj, wf_id, step_asgn)
                    if proj.working_directory:
                        hub._save_agents()
                else:
                    # 解绑：清除 binding 和 WF 任务
                    from ..project import WorkflowBinding
                    proj.workflow_binding = WorkflowBinding()
                    proj.tasks = [t for t in proj.tasks
                                  if not t.title.startswith("[WF Step")]
            hub._save_projects()
            auth.audit("update_project", actor=actor_name,
                       role=user_role, target=proj.id,
                       ip=get_client_ip(handler))
            handler._json(proj.to_dict())
        elif action == "delete":
            pid = body.get("project_id", "")
            ok = hub.remove_project(pid)
            if ok:
                auth.audit("delete_project", actor=actor_name,
                           role=user_role, target=pid,
                           ip=get_client_ip(handler))
                handler._json({"ok": True})
            else:
                handler._json({"error": "Project not found"}, 404)
        else:
            handler._json({"error": f"Unknown action: {action}"}, 400)

    elif path.startswith("/api/portal/projects/") and path.endswith("/members"):
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            agent_id = body.get("agent_id", "")
            responsibility = body.get("responsibility", "")
            if body.get("action") == "remove":
                proj.remove_member(agent_id)
            else:
                proj.add_member(agent_id, responsibility)
                # Sync project working_directory to new member agent
                if proj.working_directory and agent_id:
                    hub._sync_agent_to_project_dir(
                        agent_id, proj.working_directory,
                        project_id=proj.id, project_name=proj.name)
                    hub._save_agents()
            hub._save_projects()
            handler._json({"ok": True, "members": [m.to_dict() for m in proj.members]})

    elif path.startswith("/api/portal/projects/") and path.endswith("/chat"):
        proj_id = path.split("/")[4]
        content = body.get("content", "").strip()
        target_agents = body.get("target_agents")
        attachments = body.get("attachments") or []
        # Persist any attachments to the project working dir and append refs
        saved_refs: list[str] = []
        if isinstance(attachments, list) and attachments:
            proj = hub.get_project(proj_id)
            if proj is None:
                handler._json({"error": "Project not found"}, 404)
                return
            import base64 as _b64
            import os as _os
            try:
                base_dir = proj.working_directory or _os.path.join(
                    _os.environ.get("TUDOU_CLAW_DATA_DIR", "."),
                    "projects", proj_id,
                )
                att_dir = _os.path.join(base_dir, "attachments")
                _os.makedirs(att_dir, exist_ok=True)
                for att in attachments[:10]:  # cap at 10 per message
                    if not isinstance(att, dict):
                        continue
                    raw_name = str(att.get("name") or "attachment.bin")
                    safe_name = "".join(
                        c for c in raw_name if c.isalnum() or c in "._-"
                    ) or "attachment.bin"
                    data_b64 = att.get("data_base64") or ""
                    if not data_b64:
                        continue
                    try:
                        data_bytes = _b64.b64decode(data_b64)
                    except Exception:
                        continue
                    if len(data_bytes) > MAX_DATA_UPLOAD:  # hard cap
                        continue
                    ts = int(time.time() * 1000)
                    fname = f"{ts}_{safe_name}"
                    fpath = _os.path.join(att_dir, fname)
                    try:
                        with open(fpath, "wb") as _f:
                            _f.write(data_bytes)
                    except Exception:
                        continue
                    saved_refs.append(fname)
            except Exception as _ae:
                import logging as _lg
                _lg.getLogger("tudou.portal").warning(
                    "attachment save failed: %s", _ae,
                )
        # Merge attachment refs into content (rendered with 📎 inline)
        if saved_refs:
            suffix = "\n" + " ".join(f"📎{r}" for r in saved_refs)
            content = (content + suffix) if content else suffix.lstrip()
        if not content:
            handler._json({"error": "Empty message"}, 400)
        else:
            respondents = hub.project_chat(proj_id, content, target_agents)
            handler._json({
                "ok": True,
                "respondents": respondents,
                "attachments_saved": saved_refs,
            })

    elif path.startswith("/api/portal/projects/") and path.endswith("/tasks"):
        proj_id = path.split("/")[4]
        task = hub.project_assign_task(
            proj_id,
            title=body.get("title", ""),
            description=body.get("description", ""),
            assigned_to=body.get("assigned_to", ""),
            priority=body.get("priority", 0),
        )
        if task:
            handler._json(task.to_dict())
        else:
            handler._json({"error": "Project not found"}, 404)

    elif path.startswith("/api/portal/projects/") and path.endswith("/task-update"):
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            task = proj.update_task(
                body.get("task_id", ""),
                **{k: v for k, v in body.items() if k != "task_id"}
            )
            if task:
                hub._save_projects()
                # WF Step 手动标记完成 → 自动推进下一步
                new_status = body.get("status", "")
                if (new_status == "done"
                        and task.title.startswith("[WF Step")
                        and proj.workflow_binding.workflow_id):
                    try:
                        hub.project_chat_engine._auto_progress_next_step(
                            proj, task)
                    except Exception as e:
                        import logging
                        logging.getLogger("tudou.portal").warning(
                            "WF auto-progress after manual toggle failed: %s", e)
                handler._json(task.to_dict())
            else:
                handler._json({"error": "Task not found"}, 404)

    elif path.startswith("/api/portal/projects/") and path.endswith("/task-steps"):
        # Define / replace step list on a task. Body: {task_id, steps:[name,...]}
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
            return
        task_id = body.get("task_id", "")
        items = body.get("steps") or []
        task = next((t for t in proj.tasks if t.id == task_id), None)
        if not task:
            handler._json({"error": "Task not found"}, 404)
            return
        # Accept both legacy list[str] and new list[{"name", "manual_review"}]
        if not isinstance(items, list):
            handler._json({"error": "steps must be a list"}, 400)
            return
        normalized = []
        for it in items:
            if isinstance(it, str):
                normalized.append({"name": it, "manual_review": False})
            elif isinstance(it, dict) and "name" in it:
                normalized.append({
                    "name": str(it.get("name", "")),
                    "manual_review": bool(it.get("manual_review", False)),
                })
            else:
                handler._json({"error": "each step must be str or {name, manual_review}"}, 400)
                return
        # 重置 steps（保留已完成的同名 step 的状态，但同步它的 manual_review 标记）
        from ..project import TaskStep
        prev_done = {s.name: s for s in (task.steps or []) if s.status in ("done", "skipped")}
        new_steps = []
        for it in normalized:
            n = it["name"]
            if n in prev_done:
                s = prev_done[n]
                s.manual_review = it["manual_review"]
                new_steps.append(s)
            else:
                new_steps.append(TaskStep(name=n, manual_review=it["manual_review"]))
        task.steps = new_steps
        task.current_step_index = 0
        task.last_checkpoint_at = time.time()
        task.updated_at = task.last_checkpoint_at
        hub._save_projects()
        handler._json({"ok": True, "task": task.to_dict()})

    elif path.startswith("/api/portal/projects/") and path.endswith("/task-step-review"):
        # Human review of an awaiting_review step.
        # Body: {task_id, step_id, action: "approve"|"reject", result?, reason?}
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
            return
        task_id = body.get("task_id", "")
        step_id = body.get("step_id", "")
        action = body.get("action", "approve")
        task = next((t for t in proj.tasks if t.id == task_id), None)
        if not task:
            handler._json({"error": "Task not found"}, 404)
            return
        step = next((s for s in (task.steps or []) if s.id == step_id), None)
        if not step:
            handler._json({"error": "Step not found"}, 404)
            return
        if not step.manual_review:
            handler._json({"error": "step is not flagged for manual review"}, 400)
            return
        if step.status != "awaiting_review":
            handler._json({
                "error": f"step is not awaiting review (current status: {step.status})"
            }, 400)
            return
        reviewer = getattr(handler, "_admin_user", None) or "user"
        if isinstance(reviewer, dict):
            reviewer = reviewer.get("username") or reviewer.get("id") or "user"
        if action == "approve":
            ok = task.approve_step(step, reviewer_id=str(reviewer),
                                   override_result=body.get("result", ""))
            if not ok:
                handler._json({"error": "approve failed"}, 400)
                return
            # Re-trigger the runner so remaining steps can proceed.
            # If the task was waiting at this step, handle_task_assignment will
            # pick up from next_pending_step and continue.
            try:
                hub._save_projects()
            except Exception as e:
                logger.warning("Failed to save projects after task approval: %s", e)
            try:
                hub.project_chat_engine.handle_task_assignment(proj, task)
            except Exception as _e:
                # Resume failure shouldn't fail the approval itself; the next
                # restart will pick it up via _resume_interrupted_tasks.
                logger.debug("Failed to handle task assignment on approval: %s", _e)
            handler._json({"ok": True, "task": task.to_dict()})
        elif action == "reject":
            ok = task.reject_step(step, reviewer_id=str(reviewer),
                                  reason=body.get("reason", ""))
            if not ok:
                handler._json({"error": "reject failed"}, 400)
                return
            try:
                hub._save_projects()
            except Exception as e:
                logger.warning("Failed to save projects after task rejection: %s", e)
            # Re-trigger so the agent can re-run the rejected step
            try:
                hub.project_chat_engine.handle_task_assignment(proj, task)
            except Exception as e:
                logger.debug("Failed to handle task assignment on rejection: %s", e)
            handler._json({"ok": True, "task": task.to_dict()})
        else:
            handler._json({"error": "action must be approve|reject"}, 400)

    elif path.startswith("/api/portal/projects/") and path.endswith("/task-checkpoint"):
        # Manual checkpoint: {task_id, step_id, status: done|failed|skipped, result?}
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
            return
        task_id = body.get("task_id", "")
        step_id = body.get("step_id", "")
        new_status = body.get("status", "done")
        task = next((t for t in proj.tasks if t.id == task_id), None)
        if not task:
            handler._json({"error": "Task not found"}, 404)
            return
        step = next((s for s in (task.steps or []) if s.id == step_id), None)
        if not step:
            handler._json({"error": "Step not found"}, 404)
            return
        if new_status == "failed":
            task.complete_step(step, error=body.get("error", "manual fail"))
        elif new_status == "skipped":
            step.status = "skipped"
            step.completed_at = time.time()
            task.last_checkpoint_at = step.completed_at
        else:
            task.complete_step(step, result=body.get("result", ""))
        hub._save_projects()
        handler._json({"ok": True, "task": task.to_dict()})

    elif path.startswith("/api/portal/projects/") and path.endswith("/milestones"):
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            milestone = proj.add_milestone(
                name=body.get("name", ""),
                responsible_agent_id=body.get("responsible_agent_id", ""),
                due_date=body.get("due_date", ""),
            )
            hub._save_projects()
            handler._json(milestone.to_dict())

    elif path.startswith("/api/portal/projects/") and "/milestones/" in path and path.endswith("/update"):
        parts = path.split("/")
        proj_id = parts[4]
        milestone_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            milestone = proj.update_milestone(
                milestone_id,
                **{k: v for k, v in body.items() if k != "milestone_id"}
            )
            if milestone:
                hub._save_projects()
                handler._json(milestone.to_dict())
            else:
                handler._json({"error": "Milestone not found"}, 404)

    elif path.startswith("/api/portal/projects/") and "/milestones/" in path and path.endswith("/confirm"):
        parts = path.split("/")
        proj_id = parts[4]
        milestone_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            ms = proj.confirm_milestone(milestone_id, by=actor_name or "admin")
            if ms:
                proj.post_message(sender="system", sender_name="System",
                    content=f"✅ 里程碑「{ms.name}」已被 {actor_name or 'admin'} 确认通过。",
                    msg_type="system")
                hub._save_projects()
                auth.audit("confirm_milestone", actor=actor_name, role=user_role,
                           target=f"{proj_id}/{milestone_id}", ip=get_client_ip(handler))
                handler._json({"ok": True, "milestone": ms.to_dict()})
            else:
                handler._json({"error": "Milestone not found"}, 404)

    elif path.startswith("/api/portal/projects/") and "/milestones/" in path and path.endswith("/reject"):
        parts = path.split("/")
        proj_id = parts[4]
        milestone_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            reason = body.get("reason", "")
            ms = proj.reject_milestone(milestone_id, reason=reason,
                                        by=actor_name or "admin")
            if ms:
                proj.post_message(sender="system", sender_name="System",
                    content=f"❌ 里程碑「{ms.name}」被 {actor_name or 'admin'} 驳回。原因：{reason or '未说明'}",
                    msg_type="system")
                # 通知责任 agent: 触发其重新处理
                if ms.responsible_agent_id:
                    try:
                        trigger = (f"【里程碑驳回】里程碑「{ms.name}」被 admin 驳回。\n"
                                   f"原因：{reason or '未说明'}\n请修正后重新提交。")
                        threading.Thread(
                            target=hub.project_chat_engine._agent_respond,
                            args=(proj, ms.responsible_agent_id, trigger),
                            daemon=True
                        ).start()
                    except Exception as e:
                        logger.debug("Failed to start agent respond thread: %s", e)
                hub._save_projects()
                auth.audit("reject_milestone", actor=actor_name, role=user_role,
                           target=f"{proj_id}/{milestone_id}", ip=get_client_ip(handler))
                handler._json({"ok": True, "milestone": ms.to_dict()})
            else:
                handler._json({"error": "Milestone not found"}, 404)

    # ── Project Goals ──
    elif path.startswith("/api/portal/projects/") and path.endswith("/goals"):
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            g = proj.add_goal(
                name=body.get("name", ""),
                description=body.get("description", ""),
                owner_agent_id=body.get("owner_agent_id", ""),
                metric=body.get("metric", "count"),
                target_value=float(body.get("target_value", 0) or 0),
                target_text=body.get("target_text", ""),
            )
            hub._save_projects()
            handler._json(g.to_dict())

    elif path.startswith("/api/portal/projects/") and "/goals/" in path and path.endswith("/update"):
        parts = path.split("/"); proj_id = parts[4]; goal_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            allowed = {k: v for k, v in body.items()
                       if k in ("name", "description", "owner_agent_id", "metric",
                                "target_value", "current_value", "target_text", "done",
                                "linked_milestone_ids", "linked_deliverable_ids")}
            g = proj.update_goal(goal_id, **allowed)
            if g:
                hub._save_projects()
                handler._json(g.to_dict())
            else:
                handler._json({"error": "Goal not found"}, 404)

    elif path.startswith("/api/portal/projects/") and "/goals/" in path and path.endswith("/progress"):
        parts = path.split("/"); proj_id = parts[4]; goal_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            g = proj.update_goal_progress(
                goal_id,
                current_value=(float(body["current_value"]) if "current_value" in body else None),
                done=(bool(body["done"]) if "done" in body else None),
            )
            if g:
                hub._save_projects()
                handler._json(g.to_dict())
            else:
                handler._json({"error": "Goal not found"}, 404)

    elif path.startswith("/api/portal/projects/") and "/goals/" in path and path.endswith("/delete"):
        parts = path.split("/"); proj_id = parts[4]; goal_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            ok = proj.remove_goal(goal_id)
            if ok:
                hub._save_projects()
            handler._json({"ok": ok})

    # ── Project Deliverables ──
    elif path.startswith("/api/portal/projects/") and path.endswith("/deliverables"):
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            dv = proj.add_deliverable(
                title=body.get("title", ""),
                kind=body.get("kind", "document"),
                author_agent_id=body.get("author_agent_id", ""),
                task_id=body.get("task_id", ""),
                milestone_id=body.get("milestone_id", ""),
                content_text=body.get("content_text", ""),
                file_path=body.get("file_path", ""),
                url=body.get("url", ""),
            )
            hub._save_projects()
            handler._json(dv.to_dict())

    elif path.startswith("/api/portal/projects/") and "/deliverables/" in path and path.endswith("/update"):
        parts = path.split("/"); proj_id = parts[4]; dv_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            allowed = {k: v for k, v in body.items()
                       if k in ("title", "kind", "content_text", "file_path", "url",
                                "task_id", "milestone_id")}
            dv = proj.update_deliverable(dv_id, **allowed)
            if dv:
                hub._save_projects()
                handler._json(dv.to_dict())
            else:
                handler._json({"error": "Deliverable not found"}, 404)

    elif path.startswith("/api/portal/projects/") and "/deliverables/" in path and path.endswith("/submit"):
        parts = path.split("/"); proj_id = parts[4]; dv_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            dv = proj.submit_deliverable(dv_id)
            if dv:
                hub._save_projects()
                handler._json(dv.to_dict())
            else:
                handler._json({"error": "Deliverable not found"}, 404)

    elif path.startswith("/api/portal/projects/") and "/deliverables/" in path and path.endswith("/review"):
        parts = path.split("/"); proj_id = parts[4]; dv_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            approved = bool(body.get("approved", False))
            comment = body.get("comment", "") or ""
            dv = proj.review_deliverable(dv_id, approved=approved,
                                           reviewer=actor_name or "admin",
                                           comment=comment)
            if dv:
                # Notify author agent on rejection so it knows to revise.
                if (not approved) and dv.author_agent_id:
                    try:
                        trigger = (f"【交付物被驳回】「{dv.title}」\n"
                                   f"审阅意见：{comment or '未说明'}\n请修正后重新提交。")
                        threading.Thread(
                            target=hub.project_chat_engine._agent_respond,
                            args=(proj, dv.author_agent_id, trigger),
                            daemon=True,
                        ).start()
                    except Exception as e:
                        logger.debug("Failed to start deliverable review thread: %s", e)
                hub._save_projects()
                auth.audit("review_deliverable", actor=actor_name, role=user_role,
                           target=f"{proj_id}/{dv_id}/{'approve' if approved else 'reject'}",
                           ip=get_client_ip(handler))
                handler._json({"ok": True, "deliverable": dv.to_dict()})
            else:
                handler._json({"error": "Deliverable not found"}, 404)

    elif path.startswith("/api/portal/projects/") and "/deliverables/" in path and path.endswith("/delete"):
        parts = path.split("/"); proj_id = parts[4]; dv_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            ok = proj.remove_deliverable(dv_id)
            if ok:
                hub._save_projects()
            handler._json({"ok": ok})

    # ── Project Issues ──
    elif path.startswith("/api/portal/projects/") and path.endswith("/issues"):
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            iss = proj.add_issue(
                title=body.get("title", ""),
                description=body.get("description", ""),
                severity=body.get("severity", "medium"),
                reporter=body.get("reporter", "") or actor_name or "user",
                assigned_to=body.get("assigned_to", ""),
                related_task_id=body.get("related_task_id", ""),
                related_milestone_id=body.get("related_milestone_id", ""),
            )
            hub._save_projects()
            handler._json(iss.to_dict())

    elif path.startswith("/api/portal/projects/") and "/issues/" in path and path.endswith("/update"):
        parts = path.split("/"); proj_id = parts[4]; iss_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            allowed = {k: v for k, v in body.items()
                       if k in ("title", "description", "severity", "status",
                                "assigned_to", "related_task_id", "related_milestone_id")}
            iss = proj.update_issue(iss_id, **allowed)
            if iss:
                hub._save_projects()
                handler._json(iss.to_dict())
            else:
                handler._json({"error": "Issue not found"}, 404)

    elif path.startswith("/api/portal/projects/") and "/issues/" in path and path.endswith("/resolve"):
        parts = path.split("/"); proj_id = parts[4]; iss_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            iss = proj.resolve_issue(
                iss_id,
                resolution=body.get("resolution", ""),
                status=body.get("status", "resolved"),
            )
            if iss:
                hub._save_projects()
                handler._json(iss.to_dict())
            else:
                handler._json({"error": "Issue not found"}, 404)

    elif path.startswith("/api/portal/projects/") and "/issues/" in path and path.endswith("/delete"):
        parts = path.split("/"); proj_id = parts[4]; iss_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            ok = proj.remove_issue(iss_id)
            if ok:
                hub._save_projects()
            handler._json({"ok": ok})

    # ── Meetings ──
    elif path == "/api/portal/meetings":
        reg = getattr(hub, "meeting_registry", None)
        if reg is None:
            handler._json({"error": "meeting registry not initialized"}, 503)
        else:
            m = reg.create(
                title=body.get("title", "") or "Meeting",
                host=body.get("host", "") or actor_name or "user",
                participants=body.get("participants", []) or [],
                agenda=body.get("agenda", ""),
                project_id=body.get("project_id", ""),
            )
            handler._json(m.to_dict())

    elif path.startswith("/api/portal/meetings/"):
        reg = getattr(hub, "meeting_registry", None)
        if reg is None:
            handler._json({"error": "meeting registry not initialized"}, 503)
            return
        parts = path.split("/")
        mid = parts[4] if len(parts) >= 5 else ""
        action = parts[5] if len(parts) >= 6 else ""
        m = reg.get(mid)
        if not m:
            handler._json({"error": "Meeting not found"}, 404)
        elif action == "start":
            m.start(); reg.save(); handler._json(m.to_dict())
        elif action == "close":
            m.close(summary=body.get("summary", ""))
            reg.save(); handler._json(m.to_dict())
        elif action == "cancel":
            m.cancel(); reg.save(); handler._json(m.to_dict())
        elif action == "participants":
            op = body.get("op", "add")
            aid = body.get("agent_id", "")
            ok = (m.add_participant(aid) if op == "add" else m.remove_participant(aid))
            reg.save()
            handler._json({"ok": ok, "participants": list(m.participants)})
        elif action == "messages":
            user_content = body.get("content", "")
            attachments = body.get("attachments") or []

            # ── Handle attachments: save files + build multimodal content ──
            saved_refs: list[str] = []
            multimodal_parts: list[dict] = []
            if isinstance(attachments, list) and attachments:
                import base64 as _b64
                import os as _os
                try:
                    att_dir = _os.path.join(hub._data_dir, "meeting_attachments", m.id)
                    _os.makedirs(att_dir, exist_ok=True)
                    for att in attachments[:10]:
                        if not isinstance(att, dict):
                            continue
                        raw_name = str(att.get("name") or "attachment.bin")
                        safe_name = "".join(
                            c for c in raw_name if c.isalnum() or c in "._-"
                        ) or "attachment.bin"
                        data_b64 = att.get("data_base64") or ""
                        mime_type = str(att.get("mime") or "application/octet-stream")
                        if not data_b64:
                            continue
                        try:
                            data_bytes = _b64.b64decode(data_b64)
                        except Exception:
                            continue
                        if len(data_bytes) > MAX_DATA_UPLOAD:
                            continue
                        ts = int(time.time() * 1000)
                        fname = f"{ts}_{safe_name}"
                        fpath = _os.path.join(att_dir, fname)
                        try:
                            with open(fpath, "wb") as _f:
                                _f.write(data_bytes)
                        except Exception:
                            continue
                        saved_refs.append(fname)
                        if mime_type.startswith("image/"):
                            multimodal_parts.append({
                                "type": "image_url",
                                "image_url": {"url": f"data:{mime_type};base64,{data_b64}"},
                            })
                        else:
                            multimodal_parts.append({
                                "type": "text",
                                "text": f"[Attached file: {safe_name} ({mime_type})]",
                            })
                except Exception as _ae:
                    logger.warning("meeting attachment save failed: %s", _ae)

            msg = m.add_message(
                sender=body.get("sender", "") or actor_name or "user",
                sender_name=body.get("sender_name", "") or actor_name or "user",
                role=body.get("role", "user"),
                content=user_content,
                attachments=saved_refs or (body.get("attachments", []) or []),
            )
            reg.save()
            # ── Agent auto-reply loop ──
            # If a human (role=user) posts a message to an ACTIVE meeting,
            # fire each participant agent to reply in-meeting. Mirrors the
            # ProjectChatEngine._agent_respond flow but for meeting context.
            try:
                from .. import meeting as _meeting_mod
                if (msg.role == "user"
                    and m.status == _meeting_mod.MeetingStatus.ACTIVE):
                    pce = getattr(hub, "project_chat_engine", None)
                    if pce is not None and m.participants:
                        _meeting_mod.spawn_meeting_reply(
                            meeting=m,
                            registry=reg,
                            agent_chat_fn=pce._chat,
                            agent_lookup_fn=pce._lookup,
                            user_msg=msg.content,
                            target_agent_ids=(body.get("target_agents") or None),
                            multimodal_parts=multimodal_parts or None,
                        )
            except Exception as _e:
                logger.warning("meeting agent reply spawn failed: %s", _e)
            handler._json(msg.to_dict())
        elif action == "assignments":
            sub = parts[6] if len(parts) >= 7 else ""
            if sub and parts[6] != "":
                # /meetings/{id}/assignments/{aid}/update
                aid = parts[6]
                if len(parts) >= 8 and parts[7] == "update":
                    updated = m.update_assignment(aid, **{
                        k: v for k, v in body.items()
                        if k in ("title", "description", "assignee_agent_id",
                                 "due_hint", "status", "result")
                    })
                    if updated:
                        reg.save()
                        handler._json(updated.to_dict())
                    else:
                        handler._json({"error": "Assignment not found"}, 404)
                    return
            a = m.add_assignment(
                title=body.get("title", ""),
                description=body.get("description", ""),
                assignee_agent_id=body.get("assignee_agent_id", ""),
                due_hint=body.get("due_hint", ""),
                project_id=body.get("project_id", ""),
            )
            # Optionally materialize: if assign_to_project → create ProjectTask;
            # otherwise → create StandaloneTask linked back via source_meeting_id.
            if body.get("materialize", True):
                if a.project_id:
                    proj = hub.get_project(a.project_id)
                    if proj:
                        try:
                            pt = proj.add_task(
                                title=a.title,
                                description=a.description,
                                assigned_to=a.assignee_agent_id,
                            )
                            a.project_task_id = pt.id
                            hub._save_projects()
                        except Exception:
                            pass
                else:
                    st_reg = getattr(hub, "standalone_task_registry", None)
                    if st_reg is not None:
                        try:
                            st = st_reg.create(
                                title=a.title,
                                description=a.description,
                                assigned_to=a.assignee_agent_id,
                                created_by=m.host or "user",
                                due_hint=a.due_hint,
                                source_meeting_id=m.id,
                            )
                            a.standalone_task_id = st.id
                            # Push into target agent's execution queue
                            try:
                                _bridge_standalone_to_agent(hub, st, st_reg)
                            except Exception as _be:
                                logger.warning(
                                    "meeting assignment bridge failed: %s", _be)
                        except Exception:
                            pass
            reg.save()
            handler._json(a.to_dict())
        elif action == "delete":
            ok = reg.delete(mid)
            handler._json({"ok": ok})
        else:
            handler._json({"error": "unknown meeting action"}, 400)

    # ── Standalone Tasks (non-project) ──
    elif path == "/api/portal/standalone-tasks":
        reg = getattr(hub, "standalone_task_registry", None)
        if reg is None:
            handler._json({"error": "standalone task registry not initialized"}, 503)
        else:
            t = reg.create(
                title=body.get("title", ""),
                description=body.get("description", ""),
                assigned_to=body.get("assigned_to", ""),
                created_by=body.get("created_by", "") or actor_name or "user",
                priority=body.get("priority", "normal"),
                due_hint=body.get("due_hint", ""),
                tags=body.get("tags", []) or [],
                source_meeting_id=body.get("source_meeting_id", ""),
            )
            # ── Bridge to target agent's execution queue ──
            # Without this, standalone tasks only exist in the registry and
            # the assignee agent never gets notified. Push a mirror AgentTask
            # into the agent's queue so it actually runs.
            try:
                _bridge_standalone_to_agent(hub, t, reg)
            except Exception as _e:
                logger.warning("bridge standalone→agent failed: %s", _e)
            handler._json(t.to_dict())

    elif path.startswith("/api/portal/standalone-tasks/"):
        reg = getattr(hub, "standalone_task_registry", None)
        if reg is None:
            handler._json({"error": "standalone task registry not initialized"}, 503)
            return
        parts = path.split("/")
        tid = parts[4] if len(parts) >= 5 else ""
        action = parts[5] if len(parts) >= 6 else "update"
        if action == "delete":
            ok = reg.delete(tid)
            handler._json({"ok": ok})
        else:
            allowed = {k: v for k, v in body.items()
                       if k in ("title", "description", "assigned_to", "status",
                                "priority", "due_hint", "tags", "result")}
            t = reg.update(tid, **allowed)
            if t:
                handler._json(t.to_dict())
            else:
                handler._json({"error": "Standalone task not found"}, 404)

    elif path.startswith("/api/portal/projects/") and "/tasks/" in path and path.endswith("/approve-step"):
        # Approve a workflow step that was paused pending human review.
        parts = path.split("/")
        proj_id = parts[4]
        task_id = parts[6]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            task = next((t for t in proj.tasks if t.id == task_id), None)
            if not task:
                handler._json({"error": "Task not found"}, 404)
            else:
                md = getattr(task, "metadata", None) or {}
                if not md.get("pending_approval"):
                    handler._json({"error": "Task is not awaiting approval"}, 400)
                else:
                    try:
                        task.metadata["pending_approval"] = False
                        task.metadata["approved_by"] = actor_name or "admin"
                        task.metadata["approved_at"] = time.time()
                    except Exception:
                        pass
                    # Mark as in-progress and fire the agent
                    try:
                        from ..project import ProjectTaskStatus as _PTS
                        task.status = _PTS.IN_PROGRESS
                    except Exception:
                        task.status = "in_progress"
                    task.updated_at = time.time()
                    step_name = re.sub(r'^\[WF Step \d+\]\s*', '', task.title)
                    proj.post_message(
                        sender="system", sender_name="System",
                        content=f"✅ 步骤「{step_name}」已获 {actor_name or 'admin'} 批准，开始执行。",
                        msg_type="system",
                    )
                    if task.assigned_to:
                        try:
                            trigger = (
                                f"【人工已批准】步骤「{step_name}」已通过人工审核，"
                                f"请按原定职责立即开始执行该步骤。"
                                f"完成后请在回复中包含 ✅ 和 '已完成' 来标记步骤完成。"
                            )
                            threading.Thread(
                                target=hub.project_chat_engine._agent_respond,
                                args=(proj, task.assigned_to, trigger),
                                daemon=True,
                            ).start()
                        except Exception:
                            pass
                    hub._save_projects()
                    auth.audit("approve_step", actor=actor_name, role=user_role,
                               target=f"{proj_id}/{task_id}", ip=get_client_ip(handler))
                    handler._json({"ok": True, "task": task.to_dict()})

    elif path.startswith("/api/portal/projects/") and path.endswith("/status"):
        # Lifecycle transition: planning | active | suspended | cancelled | completed | archived
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            new_status = (body.get("status") or "").strip().lower()
            reason = body.get("reason", "") or ""
            ok, msg = proj.set_status(new_status, by=actor_name or "admin", reason=reason)
            if not ok:
                handler._json({"error": msg}, 400)
            else:
                try:
                    label_map = {
                        "planning": "未开始",
                        "active": "进行中",
                        "suspended": "挂起",
                        "cancelled": "停止",
                        "completed": "结束",
                        "archived": "归档",
                    }
                    label = label_map.get(new_status, new_status)
                    proj.post_message(
                        sender="system", sender_name="System",
                        content=f"📌 项目状态变更：{msg}（{label}）" + (f" — 原因：{reason}" if reason else ""),
                        msg_type="system",
                    )
                except Exception:
                    pass
                hub._save_projects()
                auth.audit("project_status", actor=actor_name, role=user_role,
                           target=f"{proj_id}:{new_status}", ip=get_client_ip(handler))
                handler._json({"ok": True, "status": new_status, "transition": msg})

    elif path.startswith("/api/portal/projects/") and path.endswith("/pause"):
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            reason = body.get("reason", "")
            proj.pause(by=actor_name or "admin", reason=reason)
            proj.post_message(sender="system", sender_name="System",
                content=f"⏸️ 项目已被 {actor_name or 'admin'} 暂停。原因：{reason or '未说明'}",
                msg_type="system")
            hub._save_projects()
            auth.audit("pause_project", actor=actor_name, role=user_role,
                       target=proj_id, ip=get_client_ip(handler))
            handler._json({"ok": True, "paused": True})

    elif path.startswith("/api/portal/projects/") and path.endswith("/resume"):
        proj_id = path.split("/")[4]
        proj = hub.get_project(proj_id)
        if not proj:
            handler._json({"error": "Project not found"}, 404)
        else:
            proj.resume(by=actor_name or "admin")
            proj.post_message(sender="system", sender_name="System",
                content=f"▶️ 项目已被 {actor_name or 'admin'} 恢复运行。",
                msg_type="system")
            # 回放暂停期间的消息
            queued = proj.drain_paused_queue()
            for q in queued:
                try:
                    hub.project_chat_engine.handle_user_message(
                        proj, q.get("content", ""),
                        target_agents=q.get("target_agents"))
                except Exception:
                    pass
            # 触发自动唤醒
            try:
                hub.project_chat_engine._resume_auto_wake(proj)
            except Exception:
                pass
            hub._save_projects()
            auth.audit("resume_project", actor=actor_name, role=user_role,
                       target=proj_id, ip=get_client_ip(handler))
            handler._json({"ok": True, "paused": False, "replayed": len(queued)})

    elif path == "/api/hub/broadcast":
        msgs = hub.broadcast(body.get("content", ""))
        handler._json({"sent": len(msgs)})

    elif path == "/api/hub/orchestrate":
        results = hub.orchestrate(
            body.get("task", ""),
            body.get("agent_ids"),
        )
        handler._json({"results": {k: v[:2000] for k, v in results.items()}})

    # ---- Channel endpoints ----
    elif path == "/api/portal/channels":
        router = get_router()
        ch = router.add_channel(
            name=body.get("name", ""),
            channel_type=ChannelType(body.get("channel_type", "webhook")),
            agent_id=body.get("agent_id", ""),
            bot_token=body.get("bot_token", ""),
            signing_secret=body.get("signing_secret", ""),
            webhook_url=body.get("webhook_url", ""),
            app_id=body.get("app_id", ""),
            app_secret=body.get("app_secret", ""),
        )
        auth.audit("add_channel", actor=actor_name, role=user_role,
                   target=ch.id, ip=get_client_ip(handler))
        handler._json(ch.to_dict(mask_secrets=True))

    elif path.startswith("/api/portal/channels/") and path.endswith("/update"):
        channel_id = path.split("/")[4]
        router = get_router()
        kwargs = {}
        for k in ("name", "channel_type", "agent_id", "bot_token",
                   "signing_secret", "webhook_url", "app_id",
                   "app_secret", "enabled"):
            if k in body:
                # Don't overwrite secrets with mask
                if k in ("bot_token", "signing_secret", "app_secret") and body[k] == "********":
                    continue
                kwargs[k] = body[k]
        ch = router.update_channel(channel_id, **kwargs)
        if ch:
            auth.audit("update_channel", actor=actor_name, role=user_role,
                       target=channel_id, ip=get_client_ip(handler))
            handler._json(ch.to_dict(mask_secrets=True))
        else:
            handler._json({"error": "Channel not found"}, 404)

    elif path.startswith("/api/portal/channels/") and path.endswith("/webhook"):
        # Inbound webhook receiver — public endpoint (no auth required for external platforms)
        channel_id = path.split("/")[4]
        router = get_router()
        headers_dict = {k: handler.headers.get(k, "") for k in handler.headers}
        result = router.handle_inbound(channel_id, body, headers_dict)
        handler._json(result)

    elif path.startswith("/api/portal/channels/") and path.endswith("/test"):
        channel_id = path.split("/")[4]
        router = get_router()
        ok = router.send_to_channel(channel_id,
            "🥔 Tudou Claws test message — channel is connected!",
            {"channel_id": body.get("channel_id", "test")})
        handler._json({"ok": ok})

    # --- Scheduler Job Management ---
    elif path == "/api/portal/scheduler/jobs":
        scheduler = get_scheduler()
        action = body.get("action", "create")

        if action == "create":
            agent_id = body.get("agent_id", "")
            target_type = body.get("target_type", "chat")
            workflow_id = body.get("workflow_id", "")
            # workflow jobs may omit agent_id if step_assignments cover all steps;
            # otherwise agent_id is the default assignee.
            if not agent_id and target_type != "workflow" and not workflow_id:
                handler._json({"error": "agent_id required"}, 400)
                return
            # Support creating from preset
            preset_id = body.get("preset_id", "")
            if preset_id and preset_id in PRESET_JOBS:
                preset = dict(PRESET_JOBS[preset_id])
                preset.update({k: v for k, v in body.items()
                               if k not in ("action", "preset_id") and v})
                job = scheduler.add_job(agent_id=agent_id, **preset)
            else:
                job = scheduler.add_job(
                    agent_id=agent_id,
                    name=body.get("name", ""),
                    description=body.get("description", ""),
                    job_type=body.get("job_type", "one_time"),
                    cron_expr=body.get("cron_expr", ""),
                    prompt_template=body.get("prompt_template", ""),
                    template_ids=body.get("template_ids", []),
                    notify_channels=body.get("notify_channels", []),
                    notify_on=body.get("notify_on", "always"),
                    tags=body.get("tags", []),
                    timeout=body.get("timeout", 300),
                    max_runs=body.get("max_runs", 0),
                    target_type=target_type,
                    workflow_id=workflow_id,
                    workflow_step_assignments=body.get("workflow_step_assignments", []) or [],
                    workflow_input=body.get("workflow_input", ""),
                )
            auth.audit("create_scheduled_job", actor=actor_name,
                       role=user_role, target=job.id, ip=get_client_ip(handler))
            handler._json(job.to_dict())

        elif action == "update":
            job_id = body.get("job_id", "")
            updates = {k: v for k, v in body.items()
                       if k not in ("action", "job_id")}
            job = scheduler.update_job(job_id, **updates)
            if job:
                auth.audit("update_scheduled_job", actor=actor_name,
                           role=user_role, target=job_id, ip=get_client_ip(handler))
                handler._json(job.to_dict())
            else:
                handler._json({"error": "Job not found"}, 404)

        elif action == "delete":
            job_id = body.get("job_id", "")
            ok = scheduler.remove_job(job_id)
            if ok:
                auth.audit("delete_scheduled_job", actor=actor_name,
                           role=user_role, target=job_id, ip=get_client_ip(handler))
            handler._json({"ok": ok})

        elif action == "trigger":
            job_id = body.get("job_id", "")
            ok = scheduler.trigger_now(job_id)
            auth.audit("trigger_scheduled_job", actor=actor_name,
                       role=user_role, target=job_id, ip=get_client_ip(handler))
            handler._json({"ok": ok})

        elif action == "toggle":
            job_id = body.get("job_id", "")
            enabled = body.get("enabled", True)
            job = scheduler.update_job(job_id, enabled=enabled)
            if job:
                handler._json(job.to_dict())
            else:
                handler._json({"error": "Job not found"}, 404)

        else:
            handler._json({"error": f"Unknown action: {action}"}, 400)

    # --- MCP Node Management ---
    elif path == "/api/portal/mcp/manage":
        mcp_mgr = get_mcp_manager()
        action = body.get("action", "")
        node_id = body.get("node_id", hub.node_id)

        if action == "add_mcp":
            from ..agent import MCPServerConfig
            config = MCPServerConfig.from_dict(body.get("config", {}))
            result = mcp_mgr.add_mcp_to_node(node_id, config)
            auth.audit("add_mcp_to_node", actor=actor_name,
                       role=user_role, target=f"{node_id}/{result.id}",
                       ip=get_client_ip(handler))
            handler._json(result.to_dict())

        elif action == "remove_mcp":
            mcp_id = body.get("mcp_id", "")
            ok = mcp_mgr.remove_mcp_from_node(node_id, mcp_id)
            auth.audit("remove_mcp_from_node", actor=actor_name,
                       role=user_role, target=f"{node_id}/{mcp_id}",
                       ip=get_client_ip(handler))
            handler._json({"ok": ok})

        elif action == "bind_agent":
            agent_id = body.get("agent_id", "")
            mcp_id = body.get("mcp_id", "")
            ok = mcp_mgr.bind_mcp_to_agent(node_id, agent_id, mcp_id)
            # Sync to agent profile
            agent = hub.get_agent(agent_id)
            if agent and ok:
                mcp_mgr.sync_agent_mcps(agent)
                hub._save_agents()
            handler._json({"ok": ok})

        elif action == "unbind_agent":
            agent_id = body.get("agent_id", "")
            mcp_id = body.get("mcp_id", "")
            ok = mcp_mgr.unbind_mcp_from_agent(node_id, agent_id, mcp_id)
            agent = hub.get_agent(agent_id)
            if agent and ok:
                mcp_mgr.sync_agent_mcps(agent)
                hub._save_agents()
            handler._json({"ok": ok})

        elif action == "generate_from_catalog":
            capability_id = body.get("capability_id", "")
            env_values = body.get("env_values", {})
            config = mcp_mgr.generate_mcp_config(capability_id, env_values)
            if config:
                result = mcp_mgr.add_mcp_to_node(node_id, config)
                handler._json(result.to_dict())
            else:
                handler._json({"error": "Capability not found in catalog"}, 404)

        elif action == "validate":
            from ..agent import MCPServerConfig
            config = MCPServerConfig.from_dict(body.get("config", {}))
            valid, msg = mcp_mgr.validate_mcp_config(config)
            handler._json({"valid": valid, "message": msg})

        elif action == "add_global_mcp":
            # Add a Global MCP. It is eagerly copied into every known
            # node's available_mcps (Invariant D.1 — Global visibility).
            from ..agent import MCPServerConfig
            raw = body.get("config", {}) or {}
            cap_id = body.get("capability_id", "")
            if cap_id and not raw:
                # Also accept catalog-based creation for symmetry with
                # generate_from_catalog.
                env_values = body.get("env_values", {}) or {}
                cfg = mcp_mgr.generate_mcp_config(cap_id, env_values)
                if cfg is None:
                    handler._json({"ok": False, "error": "capability not found"}, 404)
                    return
            else:
                cfg = MCPServerConfig.from_dict(raw)
            result = mcp_mgr.add_global_mcp(cfg)
            auth.audit("add_global_mcp", actor=actor_name,
                       role=user_role, target=result.id,
                       ip=get_client_ip(handler))
            handler._json({"ok": True, "mcp": result.to_dict()})

        elif action == "remove_global_mcp":
            mcp_id = body.get("mcp_id", "")
            ok = mcp_mgr.remove_global_mcp(mcp_id)
            auth.audit("remove_global_mcp", actor=actor_name,
                       role=user_role, target=mcp_id,
                       ip=get_client_ip(handler))
            handler._json({"ok": ok})

        elif action == "change_scope":
            # Admin re-scopes an existing MCP between Node / Global /
            # multi_node. This is the single place that mutates scope —
            # no other route should bypass change_mcp_scope because it
            # owns the propagation + cleanup semantics.
            mcp_id = body.get("mcp_id", "")
            new_scope = body.get("scope", "")
            target_nodes = body.get("target_nodes", None)
            result = mcp_mgr.change_mcp_scope(
                mcp_id, new_scope, target_nodes=target_nodes
            )
            auth.audit("change_mcp_scope", actor=actor_name,
                       role=user_role,
                       target=f"{mcp_id}→{new_scope}",
                       ip=get_client_ip(handler))
            status = 200 if result.get("ok") else 400
            handler._json(result, status)

        elif action == "test_connection":
            from ..agent import MCPServerConfig
            # Build config either from catalog+env or from raw config
            cap_id = body.get("capability_id", "")
            env_values = body.get("env_values", {}) or {}
            if cap_id:
                cfg = mcp_mgr.generate_mcp_config(cap_id, env_values)
                if cfg is None:
                    handler._json({"ok": False, "message": "目录中未找到该能力"}, 404)
                    return
            else:
                cfg = MCPServerConfig.from_dict(body.get("config", {}))
            test_result = mcp_mgr.test_mcp_connection(cfg, timeout=15.0)
            handler._json(test_result)

        elif action == "set_env":
            mcp_id = body.get("mcp_id", "")
            agent_id = body.get("agent_id", "")  # optional: per-agent override
            # delete_keys: list of env var names to remove entirely.
            # For node-level we purge both the override layer and the base
            # MCPServerConfig.env so the key truly disappears (otherwise a
            # base-layer key would silently reappear on next load).
            delete_keys = list(body.get("delete_keys", []) or [])
            if agent_id:
                # Per-agent MCP env override (e.g. different email account per agent).
                # Compose the new dict by starting from current, dropping deletes,
                # then applying upserts, then writing once.
                cur = dict(mcp_mgr.get_agent_mcp_env(node_id, agent_id, mcp_id) or {})
                for k in delete_keys:
                    cur.pop(k, None)
                for k, v in (body.get("env", {}) or {}).items():
                    cur[k] = v
                mcp_mgr.set_agent_mcp_env(node_id, agent_id, mcp_id, cur)
            else:
                # Node-level MCP env override (shared across agents)
                node_cfg = mcp_mgr.get_node_mcp_config(node_id)
                # 1) Deletes: purge from both override dict and the base env
                if delete_keys:
                    ov = node_cfg.env_overrides.get(mcp_id, {})
                    for k in delete_keys:
                        ov.pop(k, None)
                    base = node_cfg.available_mcps.get(mcp_id)
                    if base is not None and isinstance(base.env, dict):
                        for k in delete_keys:
                            base.env.pop(k, None)
                # 2) Upserts
                for k, v in body.get("env", {}).items():
                    node_cfg.set_env_override(mcp_id, k, v)
                mcp_mgr._save()
            handler._json({"ok": True})

        elif action == "get_agent_env":
            # Get agent-specific env overrides for a MCP
            mcp_id = body.get("mcp_id", "")
            agent_id = body.get("agent_id", "")
            if not agent_id:
                handler._json({"ok": False, "message": "agent_id required"}, 400)
                return
            env = mcp_mgr.get_agent_mcp_env(node_id, agent_id, mcp_id)
            handler._json({"ok": True, "env": env})

        elif action == "install":
            # ── 一键安装 MCP ──
            capability_id = body.get("capability_id", "")
            env_values = body.get("env_values", {})
            result = mcp_mgr.install_mcp(node_id, capability_id, env_values)
            auth.audit("install_mcp", actor=actor_name,
                       role=user_role, target=f"{node_id}/{capability_id}",
                       ip=get_client_ip(handler))
            handler._json(result)

        elif action == "install_status":
            # 查询安装任务状态
            task_id = body.get("task_id", "")
            if task_id:
                task = mcp_mgr.get_install_task(task_id)
                handler._json(task or {"error": "Task not found"})
            else:
                tasks = mcp_mgr.get_install_tasks(node_id)
                handler._json({"tasks": tasks})

        elif action == "retry_install":
            mcp_id = body.get("mcp_id", "")
            result = mcp_mgr.retry_install(node_id, mcp_id)
            handler._json(result)

        elif action == "prerequisites":
            prereqs = mcp_mgr.check_prerequisites()
            handler._json({"prerequisites": prereqs})

        elif action == "sync_global_to_node":
            # 同步 Global MCP 到指定 Node
            mcp_ids = body.get("mcp_ids", None)
            result = mcp_mgr.sync_global_to_node(node_id, mcp_ids)
            handler._json(result)

        elif action == "sync_global_to_all":
            # 同步 Global MCP 到所有 Node
            mcp_ids = body.get("mcp_ids", None)
            results = mcp_mgr.sync_global_to_all_nodes(mcp_ids)
            handler._json({"results": results})

        elif action == "remove_global":
            mcp_id = body.get("mcp_id", "")
            ok = mcp_mgr.remove_global_mcp(mcp_id)
            handler._json({"ok": ok})

        else:
            handler._json({"error": f"Unknown action: {action}"}, 400)

    # --- Vector Memory Management ---
    elif path == "/api/portal/vector/manage":
        action = body.get("action", "")
        if action == "migrate":
            # One-time sync of existing FTS5 data → ChromaDB
            from ..core.memory import get_memory_manager
            mm = get_memory_manager()
            agent_id = body.get("agent_id", None)  # Optional
            stats = mm.migrate_to_vector(agent_id)
            handler._json(stats)
        elif action == "stats":
            # Get vector store statistics
            from ..core.memory import get_memory_manager
            mm = get_memory_manager()
            stats = mm.get_vector_stats()
            handler._json(stats)
        else:
            handler._json({"error": f"Unknown vector action: {action}"}, 400)

    # --- Template Library CRUD ---
    elif path == "/api/portal/templates":
        from .template_library import get_template_library
        tpl_lib = get_template_library()
        action = body.get("action", "create")

        if action == "create":
            name = body.get("name", "").strip()
            content = body.get("content", "").strip()
            if not name or not content:
                handler._json({"error": "Name and content are required"}, 400)
                return
            tpl = tpl_lib.add_template(
                name=name,
                content=content,
                description=body.get("description", ""),
                roles=body.get("roles", []),
                tags=body.get("tags", []),
                category=body.get("category", "general"),
                created_by=actor_name,
            )
            auth.audit("create_template", actor=actor_name, role=user_role,
                       target=tpl.id, ip=get_client_ip(handler))
            handler._json(tpl.to_dict(include_content=True))

        elif action == "update":
            template_id = body.get("template_id", "")
            if not template_id:
                handler._json({"error": "template_id required"}, 400)
                return
            tpl = tpl_lib.update_template(
                template_id,
                content=body.get("content", ""),
                name=body.get("name", ""),
                description=body.get("description", ""),
                roles=body.get("roles"),
                tags=body.get("tags"),
            )
            if tpl:
                auth.audit("update_template", actor=actor_name, role=user_role,
                           target=template_id, ip=get_client_ip(handler))
                handler._json(tpl.to_dict(include_content=True))
            else:
                handler._json({"error": "Template not found"}, 404)

        elif action == "delete":
            template_id = body.get("template_id", "")
            if not template_id:
                handler._json({"error": "template_id required"}, 400)
                return
            ok = tpl_lib.remove_template(template_id)
            if ok:
                auth.audit("delete_template", actor=actor_name, role=user_role,
                           target=template_id, ip=get_client_ip(handler))
            handler._json({"ok": ok})

        else:
            handler._json({"error": f"Unknown action: {action}"}, 400)

    elif path == "/api/portal/admins/list":
        # List all admins (superAdmin only)
        if not is_super_admin(handler):
            handler._json({"error": "SuperAdmin role required"}, 403)
            auth.audit("admin_list", actor=actor_name, role=user_role,
                       ip=get_client_ip(handler), success=False)
            return
        admins = auth.admin_mgr.list_admins()
        auth.audit("admin_list", actor=actor_name, role=user_role,
                   ip=get_client_ip(handler))
        handler._json({"admins": admins})

    elif path == "/api/portal/admins/create":
        # Create new admin (superAdmin only)
        if not is_super_admin(handler):
            handler._json({"error": "SuperAdmin role required"}, 403)
            auth.audit("admin_create", actor=actor_name, role=user_role,
                       ip=get_client_ip(handler), success=False)
            return
        try:
            username = body.get("username", "").strip()
            password = body.get("password", "").strip()
            display_name = body.get("display_name", "").strip()
            agent_ids = body.get("agent_ids", [])
            if not username or not password:
                handler._json({"error": "username and password required"}, 400)
                return
            admin = auth.admin_mgr.create_admin(
                username=username,
                password=password,
                display_name=display_name or username,
                agent_ids=agent_ids,
            )
            auth.audit("admin_create", actor=actor_name, role=user_role,
                       target=admin.user_id, ip=get_client_ip(handler))
            handler._json({"ok": True, "admin": admin.to_dict(include_secrets=False)})
        except ValueError as e:
            handler._json({"error": str(e)}, 400)

    elif path == "/api/portal/admins/update":
        # Update admin (superAdmin only)
        if not is_super_admin(handler):
            handler._json({"error": "SuperAdmin role required"}, 403)
            auth.audit("admin_update", actor=actor_name, role=user_role,
                       ip=get_client_ip(handler), success=False)
            return
        user_id = body.get("user_id", "")
        if not user_id:
            handler._json({"error": "user_id required"}, 400)
            return
        kwargs = {}
        if "password" in body:
            kwargs["password"] = body["password"]
        if "display_name" in body:
            kwargs["display_name"] = body["display_name"]
        if "agent_ids" in body:
            kwargs["agent_ids"] = body["agent_ids"]
        if "active" in body:
            kwargs["active"] = body["active"]
        admin = auth.admin_mgr.update_admin(user_id, **kwargs)
        if not admin:
            handler._json({"error": "Admin not found"}, 404)
            return
        auth.audit("admin_update", actor=actor_name, role=user_role,
                   target=user_id, ip=get_client_ip(handler))
        handler._json({"ok": True, "admin": admin.to_dict(include_secrets=False)})

    elif path == "/api/portal/admins/delete":
        # Delete admin (superAdmin only)
        if not is_super_admin(handler):
            handler._json({"error": "SuperAdmin role required"}, 403)
            auth.audit("admin_delete", actor=actor_name, role=user_role,
                       ip=get_client_ip(handler), success=False)
            return
        user_id = body.get("user_id", "")
        if not user_id:
            handler._json({"error": "user_id required"}, 400)
            return
        try:
            ok = auth.admin_mgr.delete_admin(user_id)
            auth.audit("admin_delete", actor=actor_name, role=user_role,
                       target=user_id, ip=get_client_ip(handler), success=ok)
            handler._json({"ok": ok})
        except ValueError as e:
            handler._json({"error": str(e)}, 400)

    elif path == "/api/portal/admins/bind":
        # Bind agents to admin (superAdmin only)
        if not is_super_admin(handler):
            handler._json({"error": "SuperAdmin role required"}, 403)
            auth.audit("admin_bind", actor=actor_name, role=user_role,
                       ip=get_client_ip(handler), success=False)
            return
        user_id = body.get("user_id", "")
        agent_ids = body.get("agent_ids", [])
        if not user_id:
            handler._json({"error": "user_id required"}, 400)
            return
        admin = auth.admin_mgr.bind_agents(user_id, agent_ids)
        if not admin:
            handler._json({"error": "Admin not found"}, 404)
            return
        auth.audit("admin_bind", actor=actor_name, role=user_role,
                   target=user_id, ip=get_client_ip(handler))
        handler._json({"ok": True, "admin": admin.to_dict(include_secrets=False)})

    elif path == "/api/portal/admins/change-password":
        uid = body.get("user_id", "")
        old_password = body.get("old_password", "")
        new_password = body.get("new_password", "")
        if not uid or not old_password or not new_password:
            handler._json({"error": "Missing required fields"}, 400)
            return
        if len(new_password) < 6:
            handler._json({"error": "New password must be at least 6 characters"}, 400)
            return
        # Verify the caller is the same user or superAdmin
        caller_admin_id = get_admin_context(handler)
        if caller_admin_id != uid and not is_super_admin(handler):
            handler._json({"error": "Permission denied"}, 403)
            return
        # Verify old password
        admin_user = auth.admin_mgr.authenticate(
            auth.admin_mgr.get_admin(uid).username if auth.admin_mgr.get_admin(uid) else "",
            old_password)
        if not admin_user:
            handler._json({"error": "当前密码错误"}, 400)
            return
        # Update password
        try:
            auth.admin_mgr.update_admin(uid, password=new_password)
            auth.audit("change_password", actor=actor_name, role=user_role,
                       target=uid, ip=get_client_ip(handler))
            handler._json({"ok": True})
        except Exception as e:
            handler._json({"error": str(e)}, 500)

    elif path == "/api/portal/knowledge":
        # Add a new knowledge entry
        title = body.get("title", "").strip()
        content = body.get("content", "").strip()
        tags = body.get("tags", [])
        if not title or not content:
            handler._json({"error": "title and content are required"}, 400)
            return
        entry = knowledge.add_entry(title, content, tags)
        auth.audit("add_knowledge", actor=actor_name, role=user_role,
                   target=entry.get("id", ""), ip=get_client_ip(handler))
        handler._json(entry)

    elif path.startswith("/api/portal/knowledge/") and path.endswith("/delete"):
        # Delete a knowledge entry
        entry_id = path.split("/")[4]
        ok = knowledge.delete_entry(entry_id)
        auth.audit("delete_knowledge", actor=actor_name, role=user_role,
                   target=entry_id, ip=get_client_ip(handler), success=ok)
        handler._json({"ok": ok})

    elif path.startswith("/api/portal/knowledge/"):
        # Update a knowledge entry (POST to /api/portal/knowledge/{id})
        entry_id = path.split("/")[4]
        title = body.get("title")
        content = body.get("content")
        tags = body.get("tags")
        entry = knowledge.update_entry(entry_id, title=title, content=content, tags=tags)
        if entry:
            auth.audit("update_knowledge", actor=actor_name, role=user_role,
                       target=entry_id, ip=get_client_ip(handler))
            handler._json(entry)
        else:
            handler._json({"error": "Entry not found"}, 404)

    # --- RAG Provider Registry API ---
    elif path == "/api/portal/rag/providers":
        # List or register RAG providers
        from ..rag_provider import get_rag_registry
        reg = get_rag_registry()
        if handler.command == "GET":
            handler._json({"providers": [p.to_dict() for p in reg.list_providers()]})
        else:
            # Register new provider
            entry = reg.register(
                name=body.get("name", ""),
                kind=body.get("kind", "remote"),
                base_url=body.get("base_url", ""),
                api_key=body.get("api_key", ""),
                config=body.get("config", {}),
            )
            auth.audit("register_rag_provider", actor=actor_name,
                       role=user_role, target=entry.id, ip=get_client_ip(handler))
            handler._json(entry.to_dict())

    elif path.startswith("/api/portal/rag/provider/") and path.endswith("/delete"):
        from ..rag_provider import get_rag_registry
        provider_id = path.split("/")[4]
        ok = get_rag_registry().remove(provider_id)
        handler._json({"ok": ok})

    elif path == "/api/portal/rag/collections":
        # List collections on a provider
        from ..rag_provider import get_rag_registry
        provider_id = body.get("provider_id", "")
        colls = get_rag_registry().list_collections(provider_id)
        handler._json({"collections": [c.to_dict() for c in colls]})

    elif path == "/api/portal/rag/collection/create":
        from ..rag_provider import get_rag_registry
        coll = get_rag_registry().create_collection(
            provider_id=body.get("provider_id", ""),
            collection_name=body.get("name", ""),
            description=body.get("description", ""),
        )
        handler._json(coll.to_dict())

    elif path == "/api/portal/rag/search":
        # Search endpoint (also serves as remote node API)
        from ..rag_provider import get_rag_registry
        results = get_rag_registry().search(
            provider_id=body.get("provider_id", ""),
            collection=body.get("collection", "knowledge"),
            query=body.get("query", ""),
            top_k=int(body.get("top_k", 5)),
        )
        handler._json({"results": [r.to_dict() for r in results]})

    elif path == "/api/portal/rag/ingest":
        # Ingest documents into a collection
        from ..rag_provider import get_rag_registry
        count = get_rag_registry().ingest(
            provider_id=body.get("provider_id", ""),
            collection=body.get("collection", "knowledge"),
            documents=body.get("documents", []),
        )
        auth.audit("rag_ingest", actor=actor_name, role=user_role,
                   target=body.get("collection", ""), ip=get_client_ip(handler))
        handler._json({"ok": True, "count": count})

    elif path == "/api/portal/rag/parse-file":
        # Parse an uploaded file (base64-encoded) into plain text
        # Supports: PDF, DOCX, HTML, TXT, MD, CSV
        import base64, tempfile, os as _os
        file_b64 = body.get("file_data", "")          # base64 encoded
        file_name = body.get("file_name", "unknown")
        if not file_b64:
            handler._json({"error": "file_data is required"}, 400)
            return

        raw_bytes = base64.b64decode(file_b64)
        ext = _os.path.splitext(file_name)[1].lower()
        text = ""
        parse_method = "raw"

        try:
            if ext == ".pdf":
                # Try pdfplumber first, fallback to PyMuPDF, then raw
                try:
                    import pdfplumber
                    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
                        tmp.write(raw_bytes); tmp_path = tmp.name
                    try:
                        with pdfplumber.open(tmp_path) as pdf:
                            pages = []
                            for page in pdf.pages:
                                t = page.extract_text()
                                if t:
                                    pages.append(t)
                            text = "\n\n".join(pages)
                        parse_method = "pdfplumber"
                    finally:
                        _os.unlink(tmp_path)
                except ImportError:
                    try:
                        import fitz  # PyMuPDF
                        doc = fitz.open(stream=raw_bytes, filetype="pdf")
                        pages = [doc[i].get_text() for i in range(len(doc))]
                        text = "\n\n".join(p for p in pages if p.strip())
                        doc.close()
                        parse_method = "pymupdf"
                    except ImportError:
                        text = raw_bytes.decode("utf-8", errors="ignore")
                        parse_method = "raw_fallback"

            elif ext in (".docx",):
                try:
                    from docx import Document as DocxDocument
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                        tmp.write(raw_bytes); tmp_path = tmp.name
                    try:
                        doc = DocxDocument(tmp_path)
                        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                        # Also extract tables
                        for table in doc.tables:
                            for row in table.rows:
                                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                                if cells:
                                    paragraphs.append(" | ".join(cells))
                        text = "\n\n".join(paragraphs)
                        parse_method = "python-docx"
                    finally:
                        _os.unlink(tmp_path)
                except ImportError:
                    # Fallback: extract text from docx XML
                    import zipfile, io, re as _re
                    zf = zipfile.ZipFile(io.BytesIO(raw_bytes))
                    xml_content = zf.read("word/document.xml").decode("utf-8")
                    # Strip XML tags to get text
                    text = _re.sub(r"<[^>]+>", " ", xml_content)
                    text = _re.sub(r"\s+", " ", text).strip()
                    parse_method = "docx_xml_fallback"

            elif ext in (".html", ".htm"):
                raw_html = raw_bytes.decode("utf-8", errors="ignore")
                try:
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(raw_html, "html.parser")
                    # Remove script/style
                    for tag in soup(["script", "style", "nav", "footer", "header"]):
                        tag.decompose()
                    text = soup.get_text(separator="\n\n", strip=True)
                    parse_method = "beautifulsoup"
                except ImportError:
                    import re as _re
                    text = _re.sub(r"<script[^>]*>.*?</script>", "", raw_html, flags=_re.S|_re.I)
                    text = _re.sub(r"<style[^>]*>.*?</style>", "", text, flags=_re.S|_re.I)
                    text = _re.sub(r"<[^>]+>", "\n", text)
                    text = _re.sub(r"\n{3,}", "\n\n", text).strip()
                    parse_method = "regex_fallback"

            elif ext in (".md", ".txt", ".csv", ".tsv", ".json", ".log"):
                # Plain text formats — try common encodings
                for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
                    try:
                        text = raw_bytes.decode(enc)
                        parse_method = f"text_{enc}"
                        break
                    except (UnicodeDecodeError, LookupError):
                        continue
                if not text:
                    text = raw_bytes.decode("utf-8", errors="ignore")
                    parse_method = "text_lossy"

            else:
                # Unknown format — try as text
                text = raw_bytes.decode("utf-8", errors="ignore")
                parse_method = "unknown_as_text"

        except Exception as exc:
            handler._json({"error": f"File parsing failed: {exc}"}, 500)
            return

        handler._json({
            "text": text,
            "length": len(text),
            "method": parse_method,
            "file_name": file_name,
        })

    elif path == "/api/portal/rag/import":
        # Import knowledge from raw text content (split into chunks)
        from ..rag_provider import get_rag_registry
        raw_content = body.get("content", "")
        title = body.get("title", "Imported Document")
        collection = body.get("collection", "knowledge")
        provider_id = body.get("provider_id", "")
        tags = body.get("tags", [])
        chunk_size = int(body.get("chunk_size", 1000))

        if not raw_content.strip():
            handler._json({"error": "content is required"}, 400)
            return

        # Split content into chunks for better retrieval
        text = raw_content.strip()
        chunks = []
        # Split by double newlines (paragraphs) first
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        current_chunk = ""
        chunk_idx = 0
        for para in paragraphs:
            if len(current_chunk) + len(para) > chunk_size and current_chunk:
                chunk_idx += 1
                chunks.append({
                    "id": f"import_{hash(title) % 100000:05d}_{chunk_idx:03d}",
                    "title": f"{title} (Part {chunk_idx})",
                    "content": current_chunk.strip(),
                    "tags": tags,
                    "source": "import",
                })
                current_chunk = para
            else:
                current_chunk += ("\n\n" + para) if current_chunk else para
        if current_chunk.strip():
            chunk_idx += 1
            chunks.append({
                "id": f"import_{hash(title) % 100000:05d}_{chunk_idx:03d}",
                "title": f"{title} (Part {chunk_idx})" if chunk_idx > 1 else title,
                "content": current_chunk.strip(),
                "tags": tags,
                "source": "import",
            })

        count = get_rag_registry().ingest(provider_id, collection, chunks)
        # Also add to shared knowledge.py if targeting shared collection
        if collection == "knowledge" and not provider_id:
            for chunk in chunks:
                knowledge.add_entry(chunk["title"], chunk["content"], tags)
        auth.audit("rag_import", actor=actor_name, role=user_role,
                   target=f"{collection}:{title}", ip=get_client_ip(handler))
        handler._json({"ok": True, "count": count, "chunks": len(chunks)})

    # --- Domain Knowledge Base CRUD ---
    elif path == "/api/portal/domain-kb/list":
        from ..rag_provider import get_domain_kb_store
        store = get_domain_kb_store()
        handler._json({"knowledge_bases": [kb.to_dict() for kb in store.list_all()]})

    elif path == "/api/portal/domain-kb/create":
        from ..rag_provider import get_domain_kb_store
        store = get_domain_kb_store()
        name = body.get("name", "").strip()
        if not name:
            handler._json({"error": "name is required"}, 400)
            return
        kb = store.create(
            name=name,
            description=body.get("description", ""),
            provider_id=body.get("provider_id", ""),
            tags=[t.strip() for t in body.get("tags", []) if t.strip()],
        )
        auth.audit("create_domain_kb", actor=actor_name, role=user_role,
                   target=kb.id, ip=get_client_ip(handler))
        handler._json(kb.to_dict())

    elif path == "/api/portal/domain-kb/update":
        from ..rag_provider import get_domain_kb_store
        store = get_domain_kb_store()
        kb_id = body.get("id", "")
        kb = store.update(
            kb_id,
            name=body.get("name"),
            description=body.get("description"),
            tags=body.get("tags"),
        )
        if kb:
            handler._json(kb.to_dict())
        else:
            handler._json({"error": "knowledge base not found"}, 404)

    elif path == "/api/portal/domain-kb/delete":
        from ..rag_provider import get_domain_kb_store
        store = get_domain_kb_store()
        kb_id = body.get("id", "")
        ok = store.delete(kb_id)
        auth.audit("delete_domain_kb", actor=actor_name, role=user_role,
                   target=kb_id, ip=get_client_ip(handler))
        handler._json({"ok": ok})

    elif path == "/api/portal/domain-kb/import":
        # Import content into a domain knowledge base
        from ..rag_provider import get_domain_kb_store, get_rag_registry
        store = get_domain_kb_store()
        kb_id = body.get("kb_id", "")
        kb = store.get(kb_id)
        if not kb:
            handler._json({"error": "knowledge base not found"}, 404)
            return
        raw_content = body.get("content", "")
        title = body.get("title", "Imported")
        tags = body.get("tags", [])
        chunk_size = int(body.get("chunk_size", 1000))
        if not raw_content.strip():
            handler._json({"error": "content is required"}, 400)
            return
        # Chunk the content
        text = raw_content.strip()
        chunks = []
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]
        current_chunk = ""
        chunk_idx = 0
        for para in paragraphs:
            if len(current_chunk) + len(para) > chunk_size and current_chunk:
                chunk_idx += 1
                chunks.append({
                    "id": f"dkb_{kb.id}_{chunk_idx:04d}",
                    "title": f"{title} (Part {chunk_idx})",
                    "content": current_chunk.strip(),
                    "tags": tags,
                    "source": "domain_import",
                })
                current_chunk = para
            else:
                current_chunk += ("\n\n" + para) if current_chunk else para
        if current_chunk.strip():
            chunk_idx += 1
            chunks.append({
                "id": f"dkb_{kb.id}_{chunk_idx:04d}",
                "title": f"{title} (Part {chunk_idx})" if chunk_idx > 1 else title,
                "content": current_chunk.strip(),
                "tags": tags,
                "source": "domain_import",
            })
        count = get_rag_registry().ingest(kb.provider_id, kb.collection, chunks)
        store.increment_doc_count(kb_id, len(chunks))
        auth.audit("domain_kb_import", actor=actor_name, role=user_role,
                   target=f"{kb_id}:{title}", ip=get_client_ip(handler))
        handler._json({"ok": True, "count": count, "chunks": len(chunks)})

    elif path == "/api/portal/domain-kb/search":
        # Search within a domain knowledge base
        from ..rag_provider import get_domain_kb_store, get_rag_registry
        store = get_domain_kb_store()
        kb_id = body.get("kb_id", "")
        kb = store.get(kb_id)
        if not kb:
            handler._json({"error": "knowledge base not found"}, 404)
            return
        query = body.get("query", "")
        top_k = int(body.get("top_k", 5))
        results = get_rag_registry().search(kb.provider_id, kb.collection, query, top_k)
        handler._json({"results": [r.to_dict() for r in results]})

    # --- Cross-node heartbeat (called by downstream nodes) ---
    elif path == "/api/hub/heartbeat":
        body = handler._read_body() or {}
        provided = handler.headers.get("X-Hub-Secret", "") if hasattr(handler, "headers") else ""
        if not auth.verify_secret(provided):
            handler._json({"error": "invalid secret"}, 401)
            return
        node_id = body.get("node_id", "")
        if not node_id:
            handler._json({"error": "missing node_id"}, 400)
            return
        try:
            from ..infra.node_manager import get_node_manager
            nm = get_node_manager()
            if nm is not None and hasattr(nm, "update_heartbeat"):
                try:
                    nm.update_heartbeat(node_id=node_id)
                except TypeError:
                    # signature variant — try positional
                    nm.update_heartbeat(node_id)
        except Exception as e:
            logger.debug("heartbeat update failed: %s", e)
        # Also bump hub.remote_nodes
        try:
            rn = hub.remote_nodes.get(node_id)
            if rn is not None:
                rn.last_seen = time.time()
        except Exception:
            pass
        handler._json({"ok": True, "ts": time.time()})

    # --- Cross-node audit ingest (called by remote nodes) ---
    elif path == "/api/hub/audit/ingest":
        body = handler._read_body() or {}
        # Verify shared secret
        provided = handler.headers.get("X-Hub-Secret", "") if hasattr(handler, "headers") else ""
        if not auth.verify_secret(provided):
            handler._json({"error": "invalid secret"}, 401)
            return
        entries = body.get("entries") or []
        source_node = body.get("source_node", "")
        if not isinstance(entries, list):
            handler._json({"error": "entries must be a list"}, 400)
            return
        try:
            count = auth.ingest_remote_audit(entries, source_node=source_node)
            handler._json({"ok": True, "ingested": count})
        except Exception as e:
            handler._json({"error": str(e)}, 500)

    # --- Skill packages (new skill registry) ---
    elif path == "/api/portal/skill-pkgs/install":
        body = body or {}
        src_dir = body.get("path") or body.get("dir")
        reg = getattr(hub, "skill_registry", None)
        if not reg:
            handler._json({"error": "skill registry unavailable"}, 503)
            return
        if not src_dir:
            handler._json({"error": "missing 'path'"}, 400)
            return
        try:
            inst = reg.install_from_directory(src_dir)
            auth.audit("install_skill", actor=actor_name, role=user_role,
                       target=inst.id, ip=get_client_ip(handler))
            handler._json({"ok": True, "skill": inst.to_dict()})
        except Exception as e:
            handler._json({"error": str(e)}, 400)

    elif path.startswith("/api/portal/skill-pkgs/") and path.endswith("/uninstall"):
        import urllib.parse as _up
        sid = _up.unquote(path.split("/")[-2])
        reg = getattr(hub, "skill_registry", None)
        if not reg:
            handler._json({"error": "skill registry unavailable"}, 503)
            return
        ok = reg.uninstall(sid)
        auth.audit("uninstall_skill", actor=actor_name, role=user_role,
                   target=sid, ip=get_client_ip(handler))
        handler._json({"ok": ok})

    elif path.startswith("/api/portal/skill-pkgs/") and path.endswith("/grant"):
        import urllib.parse as _up
        sid = _up.unquote(path.split("/")[-2])
        body = body or {}
        agent_id = body.get("agent_id")
        reg = getattr(hub, "skill_registry", None)
        if not reg or not agent_id:
            handler._json({"error": "missing skill registry or agent_id"}, 400)
            return
        try:
            reg.grant(sid, agent_id)
            # 同步到 agent.granted_skills 持久化
            ag = hub.get_agent(agent_id)
            if ag is not None and sid not in ag.granted_skills:
                ag.granted_skills.append(sid)
                try:
                    hub.save_agents()
                except Exception:
                    pass
            auth.audit("grant_skill", actor=actor_name, role=user_role,
                       target=f"{sid}->{agent_id}", ip=get_client_ip(handler))
            handler._json({"ok": True})
        except Exception as e:
            handler._json({"error": str(e)}, 400)

    elif path.startswith("/api/portal/skill-pkgs/") and path.endswith("/revoke"):
        import urllib.parse as _up
        sid = _up.unquote(path.split("/")[-2])
        body = body or {}
        agent_id = body.get("agent_id")
        reg = getattr(hub, "skill_registry", None)
        if not reg or not agent_id:
            handler._json({"error": "missing skill registry or agent_id"}, 400)
            return
        reg.revoke(sid, agent_id)
        ag = hub.get_agent(agent_id)
        if ag is not None and sid in ag.granted_skills:
            ag.granted_skills.remove(sid)
            try:
                hub.save_agents()
            except Exception:
                pass
        auth.audit("revoke_skill", actor=actor_name, role=user_role,
                   target=f"{sid}->{agent_id}", ip=get_client_ip(handler))
        handler._json({"ok": True})

    elif path.startswith("/api/portal/skill-pkgs/") and path.endswith("/invoke"):
        import urllib.parse as _up
        sid = _up.unquote(path.split("/")[-2])
        body = body or {}
        agent_id = body.get("agent_id", "")
        inputs = body.get("inputs", {}) or {}
        reg = getattr(hub, "skill_registry", None)
        if not reg:
            handler._json({"error": "skill registry unavailable"}, 503)
            return
        try:
            result = reg.invoke(sid, agent_id, inputs)
            auth.audit("invoke_skill", actor=actor_name, role=user_role,
                       target=f"{sid}@{agent_id}", ip=get_client_ip(handler))
            handler._json({"ok": True, "result": result})
        except Exception as e:
            handler._json({"error": str(e)}, 400)

    else:
        if path.startswith("/api/"):
            handler._json({"error": f"Not found: {path}"}, status=404)
        else:
            handler.send_error(404)

def handle_delete(handler):
    """Main DELETE dispatcher."""
    path = urlparse(handler.path).path

    path = urlparse(handler.path).path
    hub = get_hub()

    if not require_auth(handler, ):
        return

    auth = get_auth()
    actor_name, user_role = get_auth_info(handler)

    # Node mode: block deletion of global resources. Only allow delete
    # on node-level resources: local agents, local nodes, local MCPs.
    if not is_hub_mode():
        _node_deletable_prefixes = (
            "/api/portal/agent/",
            "/api/portal/mcp/",
            "/api/hub/node/",  # unregister remote nodes on this node
            "/api/portal/scheduler/",
        )
        if not any(path.startswith(p) for p in _node_deletable_prefixes):
            handler._json({"error": "This delete operation requires Hub admin access"}, 403)
            return

    if path.startswith("/api/portal/agent/"):
        agent_id = path.split("/")[-1]
        ok = hub.remove_agent(agent_id)
        auth.audit("delete_agent", actor=actor_name, role=user_role, target=agent_id, ip=get_client_ip(handler), success=ok)
        handler._json({"ok": ok})

    elif path.startswith("/api/hub/node/"):
        node_id = path.split("/")[-1]
        hub.unregister_node(node_id)
        auth.audit("delete_node", actor=actor_name, role=user_role, target=node_id, ip=get_client_ip(handler))
        handler._json({"ok": True})

    elif path.startswith("/api/auth/tokens/"):
        token_id = path.split("/")[-1]
        auth.revoke_token(token_id)
        auth.audit("revoke_token", actor=actor_name, role=user_role, target=token_id, ip=get_client_ip(handler))
        handler._json({"ok": True})

    elif path.startswith("/api/portal/providers/"):
        provider_id = path.split("/")[-1]
        reg = get_registry()
        ok = reg.remove(provider_id)
        auth.audit("delete_provider", actor=actor_name, role=user_role, target=provider_id, ip=get_client_ip(handler))
        handler._json({"ok": ok})

    elif path.startswith("/api/portal/channels/"):
        channel_id = path.split("/")[-1]
        router = get_router()
        ok = router.remove_channel(channel_id)
        auth.audit("delete_channel", actor=actor_name, role=user_role, target=channel_id, ip=get_client_ip(handler))
        handler._json({"ok": ok})

    else:
        if path.startswith("/api/"):
            handler._json({"error": f"Not found: {path}"}, status=404)
        else:
            handler.send_error(404)